"""
utils/extractor_v4.py — DMP Extractor v4

Pipeline
--------
1. DocConverter   : document (PDF / DOCX) → flat List[TextBlock]
2. DMPTrimmer     : trim to DMP section only (start = first section-1/1.1 name,
                    end = "oświadczenia administracyjne")
3. LinearMatcher  : find subsection boundaries by name variants (forward-only)
4. ContentCleaner : strip formatting markers and skip-term lines
5. DMPExtractor   : orchestrate; produce JSON cache compatible with review.html
"""

import os
import re
import json
import uuid
import zipfile
import logging
from collections import Counter
from datetime import datetime
from typing import Dict, List, Optional, Set, Tuple

from docx import Document
import PyPDF2

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from pdf2image import convert_from_path
    import pytesseract
    if os.name == 'nt':
        for _p in [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]:
            if os.path.exists(_p):
                pytesseract.pytesseract.tesseract_cmd = _p
                break
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# DMP structure constants  (must match review.html expectations)
# ─────────────────────────────────────────────────────────────────────────────

SECTION_ORDER: List[str] = [
    "1.1", "1.2",
    "2.1", "2.2",
    "3.1", "3.2",
    "4.1", "4.2",
    "5.1", "5.2", "5.3", "5.4",
    "6.1", "6.2",
]

# Last subsection in each main section — their end boundary also accepts
# the next main section's header (in addition to the next subsection anchor).
LAST_IN_SECTION: Dict[str, str] = {
    "1.2": "2",
    "2.2": "3",
    "3.2": "4",
    "4.2": "5",
    "5.4": "6",
}

SECTION_TITLES: Dict[str, str] = {
    "1": "1. Data description and collection or re-use of existing data",
    "2": "2. Documentation and data quality",
    "3": "3. Storage and backup during the research process",
    "4": "4. Legal requirements, codes of conduct",
    "5": "5. Data sharing and long-term preservation",
    "6": "6. Data management responsibilities and resources",
}

SECTION_QUESTIONS: Dict[str, str] = {
    "1.1": "How will new data be collected or produced and/or how will existing data be re-used?",
    "1.2": "What data (for example the types, formats, and volumes) will be collected or produced?",
    "2.1": "What metadata and documentation (for example methodology or data collection and way of organising data) will accompany data?",
    "2.2": "What data quality control measures will be used?",
    "3.1": "How will data and metadata be stored and backed up during the research process?",
    "3.2": "How will data security and protection of sensitive data be taken care of during the research?",
    "4.1": "If personal data are processed, how will compliance with legislation on personal data and on data security be ensured?",
    "4.2": "How will other legal issues, such as intellectual property rights and ownership, be managed? What legislation is applicable?",
    "5.1": "How and when will data be shared? Are there possible restrictions to data sharing or embargo reasons?",
    "5.2": "How will data for preservation be selected, and where will data be preserved long-term (for example a data repository or archive)?",
    "5.3": "What methods or software tools will be needed to access and use the data?",
    "5.4": "How will the application of a unique and persistent identifier (such as a Digital Object Identifier (DOI)) to each data set be ensured?",
    "6.1": "Who (for example role, position, and institution) will be responsible for data management (i.e. the data steward)?",
    "6.2": "What resources (for example financial and time) will be dedicated to data management and ensuring the data will be FAIR (Findable, Accessible, Interoperable, Re-usable)?",
}

# ─────────────────────────────────────────────────────────────────────────────
# Text helpers
# ─────────────────────────────────────────────────────────────────────────────

_RE_FMT = re.compile(
    r'\b(?:BOLD|UNDERLINED|ITALIC|UNDERLINED_BOLD)\b:?\s*'
    r'|\[(?:BOLD|UNDERLINED|ITALIC|UNDERLINED_BOLD)\]'
    r'|\{(?:BOLD|UNDERLINED|ITALIC)\}',
    re.IGNORECASE,
)
_RE_SEC_NUM = re.compile(r'^\s*\d+\.\d+\.?\s*')  # "1.1 " or "1.1. " (trailing dot optional)
_RE_MAIN_NUM = re.compile(r'^\s*\d+\.\s*')
_RE_WS = re.compile(r'\s+')

_STOP: Set[str] = {
    # English
    'that', 'will', 'data', 'with', 'from', 'this', 'have', 'been', 'they',
    'what', 'when', 'where', 'which', 'such', 'used', 'example', 'also',
    # Polish with diacritics
    'danych', 'oraz', 'jest', 'jako', 'przez', 'przy', 'które', 'jakie',
    'jak', 'będą', 'każdego',
    # Polish without diacritics (after normalize_diacritics)
    'ktore', 'beda', 'kazdego',
}

_BUILTIN_NOISE: List = [
    re.compile(r'^\s*PLAN\s+ZARZĄDZANIA\s+DANYMI\b', re.IGNORECASE),
    re.compile(r'^\s*DATA\s+MANAGEMENT\s+PLAN\b', re.IGNORECASE),
    # Polish section titles — with OR without leading numeration
    re.compile(
        r'^\s*(?:\d+[\.\s]+)?(?:'
        r'Opis\s+danych.*'
        r'|Dokumentacja\s+i\s+jako[sś][cć].*'
        r'|Przechowywanie\s+i\s+tworzenie(?!\s+kopii\s+zapasowych\s+danych).*'
        r'|Wymogi\s+prawne.*'
        r'|Wymagania\s+prawne.*'
        r'|Udost[eę]pnianie\s+i\s+d[łl]ugotr.*'
        r'|Zadania\s+zwi[aą]zane.*'
        r'|Odpowiedzialno[sś][cć].*'
        r')$',
        re.IGNORECASE,
    ),
    # English section titles — with OR without leading numeration
    re.compile(
        r'^\s*(?:\d+[\.\s]+)?(?:'
        r'Data\s+description\s+and.*'
        r'|Documentation\s+and\s+data.*'
        r'|Storage\s+and\s+backup.*'
        r'|Legal\s+requirements.*'
        r'|Data\s+sharing\s+and.*'
        r'|Data\s+management\s+responsibilities.*'
        r')$',
        re.IGNORECASE,
    ),
    # OSF print page headers, e.g. "OSF, OPUS-31 Page 45 ID: 675502, 2026-06-07 22:40:36"
    re.compile(r'^\s*OSF\b.*\b(?:Page\s+\d+|ID:\s*\d{4,})', re.IGNORECASE),
]

# End-of-DMP pattern — triggers when a block matches "oświadczenia administracyjne"
# (Polish proposals) or "administrative declarations" (English proposals),
# in any spelling/diacritics variant.
_RE_END_DMP = re.compile(
    r'o[sś]wiadczeni[ae]\s+administracyjn|administrative\s+declarations?',
    re.IGNORECASE,
)

# Polish diacritic → ASCII map (for fuzzy matching)
_DIACRITIC_MAP = str.maketrans({
    'ą': 'a', 'ć': 'c', 'ę': 'e', 'ł': 'l', 'ń': 'n',
    'ó': 'o', 'ś': 's', 'ź': 'z', 'ż': 'z',
    'Ą': 'A', 'Ć': 'C', 'Ę': 'E', 'Ł': 'L', 'Ń': 'N',
    'Ó': 'O', 'Ś': 'S', 'Ź': 'Z', 'Ż': 'Z',
})


def normalize_diacritics(text: str) -> str:
    return text.translate(_DIACRITIC_MAP)


def normalize(text: str) -> str:
    text = _RE_FMT.sub(' ', text)
    text = _RE_SEC_NUM.sub('', text)
    text = _RE_MAIN_NUM.sub('', text)
    return _RE_WS.sub(' ', text).strip().lower()


def strip_formatting(text: str) -> str:
    return _RE_WS.sub(' ', _RE_FMT.sub('', text)).strip()


def token_overlap(query: str, candidate: str) -> float:
    """Fraction of query content tokens (≥4 chars, not in _STOP) found in candidate."""
    q_tokens = {w for w in query.split() if len(w) >= 4 and w not in _STOP}
    if not q_tokens:
        return 0.0
    c_tokens = set(candidate.split())
    return len(q_tokens & c_tokens) / len(q_tokens)


def _norm_for_match(text: str) -> str:
    """Normalize for anchor matching: lowercase, strip formatting/nums, strip diacritics."""
    return normalize_diacritics(normalize(text))


def _score_block(block_text: str, norm_names: List[str]) -> float:
    """Return best token_overlap score between block_text and any of norm_names."""
    if not norm_names:
        return 0.0
    c = _norm_for_match(block_text)
    return max(token_overlap(n, c) for n in norm_names)


# ─────────────────────────────────────────────────────────────────────────────
# Intermediate data model
# ─────────────────────────────────────────────────────────────────────────────

class TextBlock:
    __slots__ = ('text', 'is_bold', 'is_heading', 'source', 'page', 'is_hf')

    def __init__(
        self,
        text: str,
        is_bold: bool = False,
        is_heading: bool = False,
        source: str = 'paragraph',
        page: int = 0,
        is_hf: bool = False,
    ) -> None:
        self.text = text
        self.is_bold = is_bold
        self.is_heading = is_heading
        self.source = source
        self.page = page
        self.is_hf = is_hf


# ─────────────────────────────────────────────────────────────────────────────
# Document converter  (unchanged from v2)
# ─────────────────────────────────────────────────────────────────────────────

class DocConverter:
    """Converts DOCX or PDF to a flat, ordered list of TextBlock objects."""

    def convert(self, file_path: str) -> List[TextBlock]:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.docx':
            return self._from_docx(file_path)
        if ext == '.pdf':
            return self._from_pdf(file_path)
        raise ValueError(f"Unsupported format: {ext}")

    def _from_docx(self, path: str) -> List[TextBlock]:
        doc = Document(path)
        hf_set: Set[str] = set()
        for sec in doc.sections:
            for hf in (
                sec.header, sec.footer,
                sec.even_page_header, sec.even_page_footer,
                sec.first_page_header, sec.first_page_footer,
            ):
                if hf:
                    for para in hf.paragraphs:
                        t = para.text.strip()
                        if t:
                            hf_set.add(t.lower())

        def is_hf(t: str) -> bool:
            return t.strip().lower() in hf_set

        return self._traverse_body(doc, is_hf)

    def _traverse_body(self, doc: Document, is_hf) -> List[TextBlock]:
        from docx.oxml.ns import qn
        para_map = {p._element: p for p in doc.paragraphs}
        blocks: List[TextBlock] = []
        pos = 0

        for child in doc.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                para = para_map.get(child)
                if para is None:
                    continue
                text = para.text.strip()
                if not text:
                    continue
                bold = self._is_para_bold(para)
                heading = (
                    para.style.name.lower().startswith('heading')
                    or (bold and len(text) < 200)
                )
                source = (
                    'heading'
                    if para.style.name.lower().startswith('heading')
                    else 'paragraph'
                )
                blocks.append(TextBlock(text, bold, heading, source, pos, is_hf(text)))
                pos += 1

            elif tag == 'tbl':
                from docx.oxml.ns import qn as _qn
                for row_el in child.findall('.//' + _qn('w:tr')):
                    seen: Set[str] = set()
                    parts: List[str] = []
                    for cell in row_el.findall(_qn('w:tc')):
                        cell_words = []
                        for p_el in cell.findall('.//' + _qn('w:p')):
                            t = ''.join(
                                r.text or ''
                                for r in p_el.findall('.//' + _qn('w:t'))
                            ).strip()
                            if t:
                                cell_words.append(t)
                        cell_text = ' '.join(cell_words)
                        if cell_text and cell_text not in seen:
                            parts.append(cell_text)
                            seen.add(cell_text)
                    if parts:
                        text = ' | '.join(parts)
                        blocks.append(
                            TextBlock(text, False, False, 'table', pos, is_hf(text))
                        )
                        pos += 1

        return blocks

    @staticmethod
    def _is_para_bold(para) -> bool:
        runs = [r for r in para.runs if r.text.strip()]
        return bool(runs) and all(r.bold for r in runs)

    def _from_pdf(self, path: str) -> List[TextBlock]:
        pages_text = self._read_pdf_pages(path)
        hf_set = self._detect_pdf_hf(pages_text)
        blocks: List[TextBlock] = []
        for page_num, page_text in enumerate(pages_text):
            for line in page_text.split('\n'):
                text = line.strip()
                if len(text) < 3:
                    continue
                blocks.append(
                    TextBlock(text, False, False, 'paragraph', page_num,
                              text.lower() in hf_set)
                )
        return blocks

    def _read_pdf_pages(self, path: str) -> List[str]:
        pages: List[str] = []
        try:
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    pages.append(page.extract_text() or '')
        except Exception as exc:
            raise RuntimeError(f"PDF read error: {exc}") from exc

        if sum(len(t) for t in pages) < 100:
            if HAS_OCR:
                return self._ocr(path)
            raise RuntimeError(
                "PDF appears to be a scanned image but OCR (pytesseract + pdf2image) "
                "is not installed."
            )

        if self._is_text_malformed(pages):
            if HAS_PDFPLUMBER:
                logger.warning("PyPDF2 extracted malformed text. Trying pdfplumber...")
                try:
                    pages_pl = self._read_pdf_with_pdfplumber(path)
                    if not self._is_text_malformed(pages_pl):
                        logger.info("pdfplumber extracted text successfully")
                        return pages_pl
                    logger.warning("pdfplumber also gave malformed text")
                except Exception as e:
                    logger.warning(f"pdfplumber failed: {e}")

        return pages

    @staticmethod
    def _is_text_malformed(pages: List[str]) -> bool:
        if not pages:
            return False
        sample = ' '.join(pages[:min(5, len(pages))])
        if len(sample) < 100:
            return False
        long_sequences = re.findall(r'[a-z]{20,}', sample)
        camel_case = re.findall(r'[a-z][A-Z]', sample)
        camel_ratio = len(camel_case) / max(len(sample) / 100, 1)
        return len(long_sequences) >= 3 or camel_ratio > 1.5

    @staticmethod
    def _read_pdf_with_pdfplumber(path: str) -> List[str]:
        pages: List[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or '')
        return pages

    @staticmethod
    def _detect_pdf_hf(pages: List[str]) -> Set[str]:
        if len(pages) < 2:
            return set()
        counter: Counter = Counter()
        for page in pages:
            seen: Set[str] = set()
            for line in page.split('\n'):
                norm = line.strip().lower()
                if norm and norm not in seen:
                    counter[norm] += 1
                    seen.add(norm)
        threshold = max(2, len(pages) // 2)
        return {line for line, cnt in counter.items() if cnt >= threshold}

    @staticmethod
    def _ocr(path: str) -> List[str]:
        images = convert_from_path(path)
        return [pytesseract.image_to_string(img, lang='pol+eng') for img in images]


# ─────────────────────────────────────────────────────────────────────────────
# Content cleaner  (unchanged from v2)
# ─────────────────────────────────────────────────────────────────────────────

class ContentCleaner:
    def clean(self, blocks: List[TextBlock], skip_patterns: list) -> List[str]:
        result = []
        for block in blocks:
            if block.is_hf:
                continue
            text = strip_formatting(block.text)
            # Check noise BEFORE stripping numerations — patterns match "4. Legal requirements..."
            # as well as "Legal requirements..." (numeration-optional patterns).
            if any(p.search(text) for p in _BUILTIN_NOISE):
                continue
            text = _RE_SEC_NUM.sub('', text)
            text = _RE_MAIN_NUM.sub('', text)
            text = _RE_WS.sub(' ', text).strip()
            if not text:
                continue
            if any(p.search(text) for p in skip_patterns):
                continue
            result.append(text)
        return result


# ─────────────────────────────────────────────────────────────────────────────
# Skip-terms manager  (unchanged from v2)
# ─────────────────────────────────────────────────────────────────────────────

class SkipTermsManager:
    _DEFAULT_PATH = os.path.normpath(
        os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            '..', 'config', 'extraction_skip_terms.json',
        )
    )

    def __init__(self, path: Optional[str] = None) -> None:
        self.path = path or self._DEFAULT_PATH

    def load(self) -> List[str]:
        if not os.path.exists(self.path):
            return []
        with open(self.path, encoding='utf-8') as f:
            return json.load(f).get('terms', [])

    def save(self, terms: List[str]) -> None:
        os.makedirs(os.path.dirname(self.path), exist_ok=True)
        with open(self.path, 'w', encoding='utf-8') as f:
            json.dump({'terms': terms}, f, ensure_ascii=False, indent=2)

    def add(self, term: str) -> List[str]:
        terms = self.load()
        if term and term not in terms:
            terms.append(term)
            self.save(terms)
        return terms

    def remove(self, term: str) -> List[str]:
        terms = [t for t in self.load() if t != term]
        self.save(terms)
        return terms

    def compile(self) -> list:
        patterns = []
        for term in self.load():
            try:
                patterns.append(re.compile(term, re.IGNORECASE))
            except re.error:
                patterns.append(re.compile(re.escape(term), re.IGNORECASE))
        return patterns


# ─────────────────────────────────────────────────────────────────────────────
# Validation helpers  (unchanged from v2)
# ─────────────────────────────────────────────────────────────────────────────

def validate_docx_file(path: str) -> Tuple[bool, str]:
    if not os.path.exists(path):
        return False, 'File not found'
    if os.path.getsize(path) == 0:
        return False, 'File is empty'
    try:
        with zipfile.ZipFile(path) as z:
            for req in ('word/document.xml', '[Content_Types].xml'):
                if req not in z.namelist():
                    return False, f'Invalid DOCX: missing {req}'
    except zipfile.BadZipFile:
        return False, 'Not a valid DOCX / ZIP file'
    return True, 'OK'


def validate_pdf_file(path: str) -> Tuple[bool, str]:
    if not os.path.exists(path):
        return False, 'File not found'
    if os.path.getsize(path) == 0:
        return False, 'File is empty'
    with open(path, 'rb') as f:
        if f.read(4) != b'%PDF':
            return False, 'Not a valid PDF file'
    return True, 'OK'


# ─────────────────────────────────────────────────────────────────────────────
# Variants loader
# ─────────────────────────────────────────────────────────────────────────────

class VariantsLoader:
    """Loads subsection and section name variants from config/dmp_variants.json."""

    _DEFAULT_PATH = os.path.normpath(
        os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            '..', 'config', 'dmp_variants.json',
        )
    )

    def __init__(self, path: Optional[str] = None) -> None:
        self.path = path or self._DEFAULT_PATH

    def load(self) -> Tuple[Dict[str, List[str]], Dict[str, List[str]]]:
        """
        Returns:
            subsection_variants: {"1.1": [names...], ...}
            section_variants:    {"1": [names...], ...}
        """
        with open(self.path, encoding='utf-8') as f:
            data = json.load(f)

        subsection_variants: Dict[str, List[str]] = {}
        section_variants: Dict[str, List[str]] = {}

        for section in data:
            sid = section['id']
            section_variants[sid] = section.get('names', [])
            for sub in section.get('subsections', []):
                subsection_variants[sub['id']] = sub.get('names', [])

        return subsection_variants, section_variants


# ─────────────────────────────────────────────────────────────────────────────
# DMP Trimmer  (step 2 — independent of step 3)
# ─────────────────────────────────────────────────────────────────────────────

class DMPTrimmer:
    """
    Trims document blocks to the DMP section.

    Start: first block that matches any variant of section-1 names OR
           subsection-1.1 names (by token overlap). No "Plan zarządzania
           danymi" prefix is required — the match is on name content alone.

    End:   first block matching "oświadczenia administracyjne" (any spelling /
           diacritics variant). Everything from that block onward is removed.

    Both detections are independent of each other and of LinearMatcher.
    """

    HIGH = 0.55  # single-block score to accept immediately
    LOW = 0.45   # minimum score to accept with a 2-block window

    def trim(
        self,
        blocks: List[TextBlock],
        sec1_names: List[str],
        sub11_names: List[str],
    ) -> List[TextBlock]:
        all_start_names = sec1_names + sub11_names
        norm_start = [_norm_for_match(n) for n in all_start_names]

        start_idx = self._find_start(blocks, norm_start)
        end_idx = self._find_end(blocks)

        if start_idx is None:
            logger.warning("DMPTrimmer: DMP start not found — using full document")
            start_idx = 0
        else:
            logger.info("DMPTrimmer: DMP starts at block %d", start_idx)

        if end_idx is not None:
            logger.info("DMPTrimmer: DMP ends at block %d (oświadczenia administracyjne)", end_idx)
            return blocks[start_idx:end_idx]

        return blocks[start_idx:]

    def _find_start(
        self, blocks: List[TextBlock], norm_names: List[str]
    ) -> Optional[int]:
        """Return index of first block that looks like a DMP section-1 or 1.1 heading."""
        for i, blk in enumerate(blocks):
            if blk.is_hf:
                continue
            # Try single block first
            single = _score_block(blk.text, norm_names)
            if single >= self.HIGH:
                logger.debug("DMPTrimmer start: block %d score %.2f (single)", i, single)
                return i
            # Try 2-block window for cases where heading wraps
            if single >= self.LOW and i + 1 < len(blocks):
                window = blk.text + ' ' + blocks[i + 1].text
                score2 = _score_block(window, norm_names)
                if score2 >= self.HIGH:
                    logger.debug("DMPTrimmer start: block %d score %.2f (2-block)", i, score2)
                    return i
        return None

    @staticmethod
    def _find_end(blocks: List[TextBlock]) -> Optional[int]:
        """Return index of the block containing 'oświadczenia administracyjne'."""
        for i, blk in enumerate(blocks):
            if blk.is_hf:
                continue
            if _RE_END_DMP.search(blk.text):
                return i
        return None


# ─────────────────────────────────────────────────────────────────────────────
# Linear matcher  (step 3 — independent of step 2)
# ─────────────────────────────────────────────────────────────────────────────

class LinearMatcher:
    """
    Finds subsection boundaries by matching name variants in document order.

    Rules:
    - Sections are searched in SECTION_ORDER (not by number, but by name).
    - Search for section N starts where section N-1's anchor ended (cursor).
    - If a section anchor is not found, its start is None (empty section),
      and the cursor does not advance.
    - Section headers (2, 3, …) are searched independently (full document scan)
      and used as additional end-of-section boundaries for subsections that are
      last in their main section (1.2, 2.2, 3.2, 4.2, 5.4).
    """

    HIGH = 0.55
    LOW = 0.38
    MIN_FIRST = 0.38  # single-block pre-filter before trying windows

    def find_all(
        self,
        blocks: List[TextBlock],
        subsection_variants: Dict[str, List[str]],
        section_variants: Dict[str, List[str]],
    ) -> Tuple[Dict[str, Optional[Tuple[int, int]]], Dict[str, Optional[int]]]:
        """
        Returns:
            subsection_matches: {sid: (start_idx, win_size) or None}
            section_starts:     {sec_id: start_idx or None}  (section headers only)
        """
        # Phase 1: find subsections in forward order
        subsection_matches: Dict[str, Optional[Tuple[int, int]]] = {}
        cursor = 0
        for sid in SECTION_ORDER:
            names = subsection_variants.get(sid, [])
            result = self._find_anchor(blocks, names, cursor, sid=sid)
            subsection_matches[sid] = result
            if result is not None:
                cursor = result[0] + result[1]
                logger.info("LinearMatcher: %s → block %d (win %d)", sid, result[0], result[1])
            else:
                logger.info("LinearMatcher: %s → NOT FOUND (cursor stays at %d)", sid, cursor)

        # Phase 2: find section headers (independent full-document scan)
        section_starts: Dict[str, Optional[int]] = {}
        for sec_id, names in section_variants.items():
            result = self._find_anchor(blocks, names, 0)
            section_starts[sec_id] = result[0] if result else None
            if result:
                logger.debug("LinearMatcher: section %s header → block %d", sec_id, result[0])

        return subsection_matches, section_starts

    # Compiled pattern for numeration-based detection (e.g. "2.1.", "3.2 ", "5.4.")
    _RE_NUMERATION = re.compile(r'^\s*(\d+)\.(\d+)[.\s]')

    def _find_anchor(
        self,
        blocks: List[TextBlock],
        names: List[str],
        search_from: int,
        sid: Optional[str] = None,
    ) -> Optional[Tuple[int, int]]:
        """
        Find the first occurrence of any name variant at or after search_from.
        Returns (block_index, window_size) or None.

        Strategy 0 (highest priority): numeration-based detection.
          If the block starts with "X.Y." or "X.Y " matching the target sid,
          accept it immediately (no score calculation needed).

        Strategy 1: token overlap — single-block pre-filter (>= MIN_FIRST),
          then window of 1-3 blocks. Returns first HIGH hit; records first LOW
          as fallback.
        """
        if not names:
            return None

        norm_names = [_norm_for_match(n) for n in names]
        first_low: Optional[Tuple[int, int]] = None

        for i in range(search_from, len(blocks)):
            blk = blocks[i]
            if blk.is_hf:
                continue

            # Skip main section title blocks — they should not match as subsection anchors
            if any(p.search(blk.text) for p in _BUILTIN_NOISE):
                continue

            # Strategy 0: numeration prefix matching (e.g. "2.1. Metadata...")
            if sid is not None:
                m = self._RE_NUMERATION.match(blk.text)
                if m and f"{m.group(1)}.{m.group(2)}" == sid:
                    logger.debug("LinearMatcher: %s numeration match at block %d", sid, i)
                    return i, 1

            single = _score_block(blk.text, norm_names)
            if single < self.MIN_FIRST:
                continue

            if single >= self.HIGH:
                # High confidence on single block — accept immediately
                return i, 1

            if single >= self.LOW:
                # Single block already qualifies — prefer win=1 to avoid
                # accidentally consuming the next subsection's content block
                if first_low is None:
                    first_low = (i, 1)
                continue

            # win=1 is between MIN_FIRST and LOW — try extending the window
            best_score, best_win = single, 1
            for win in (2, 3):
                if i + win > len(blocks):
                    break
                window_text = ' '.join(blocks[j].text for j in range(i, i + win))
                s = _score_block(window_text, norm_names)
                if s > best_score:
                    best_score, best_win = s, win

            if best_score >= self.HIGH:
                return i, best_win
            if best_score >= self.LOW and first_low is None:
                first_low = (i, best_win)

        return first_low


# ─────────────────────────────────────────────────────────────────────────────
# Main extractor
# ─────────────────────────────────────────────────────────────────────────────

class DMPExtractor:
    """
    DMP section extractor v4.

    Public interface identical to v2:
        result = DMPExtractor().process_file(file_path, output_dir)
    """

    def __init__(self) -> None:
        self._converter = DocConverter()
        self._trimmer = DMPTrimmer()
        self._matcher = LinearMatcher()
        self._cleaner = ContentCleaner()
        self._skip_mgr = SkipTermsManager()
        self._variants_loader = VariantsLoader()

    def process_file(
        self,
        file_path: str,
        output_dir: str,
        progress_callback=None,
    ) -> dict:
        def cb(msg: str, pct: int) -> None:
            if progress_callback:
                progress_callback(msg, pct)

        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.docx':
                ok, msg = validate_docx_file(file_path)
            elif ext == '.pdf':
                ok, msg = validate_pdf_file(file_path)
            else:
                return {'success': False, 'message': f'Unsupported file type: {ext}'}
            if not ok:
                return {'success': False, 'message': msg}

            cb('Converting document to text blocks…', 10)
            blocks = self._converter.convert(file_path)
            logger.info('DocConverter: %d blocks from %s', len(blocks), file_path)

            cb('Loading section name variants…', 20)
            subsection_variants, section_variants = self._variants_loader.load()

            cb('Trimming to DMP section…', 30)
            sec1_names = section_variants.get('1', [])
            sub11_names = subsection_variants.get('1.1', [])
            trimmed = self._trimmer.trim(blocks, sec1_names, sub11_names)
            logger.info('DMPTrimmer: %d → %d blocks after trim', len(blocks), len(trimmed))

            cb('Locating section boundaries…', 50)
            subsection_matches, section_starts = self._matcher.find_all(
                trimmed, subsection_variants, section_variants
            )
            found = sum(1 for v in subsection_matches.values() if v is not None)
            logger.info('LinearMatcher: found %d / %d subsections', found, len(SECTION_ORDER))

            cb('Extracting and cleaning content…', 70)
            skip_patterns = self._skip_mgr.compile()
            cache = self._build_cache(trimmed, subsection_matches, section_starts, skip_patterns)

            cb('Saving cache…', 90)
            cache_id = str(uuid.uuid4())
            cache_dir = os.path.join(output_dir, 'cache')
            os.makedirs(cache_dir, exist_ok=True)
            cache_path = os.path.join(cache_dir, f'cache_{cache_id}.json')
            with open(cache_path, 'w', encoding='utf-8') as f:
                json.dump(cache, f, ensure_ascii=False, indent=2)

            filled = sum(
                1 for sid in SECTION_ORDER
                if cache.get(sid, {}).get('paragraphs')
            )
            cb('Done.', 100)

            return {
                'success': True,
                'filename': self._smart_filename(file_path),
                'cache_id': cache_id,
                'cache_file': f'cache_{cache_id}.json',
                'message': f'Extracted {filled} of {len(SECTION_ORDER)} sections',
            }

        except Exception as exc:
            logger.exception('process_file failed for %s', file_path)
            return {'success': False, 'message': str(exc)}

    def _build_cache(
        self,
        blocks: List[TextBlock],
        subsection_matches: Dict[str, Optional[Tuple[int, int]]],
        section_starts: Dict[str, Optional[int]],
        skip_patterns: list,
    ) -> dict:
        cache: dict = {}
        # When a subsection absorbs the rest of the document (because the
        # immediately-next subsection is not found), all subsequent subsections
        # must be empty — their anchors fall inside the absorbing section's range.
        absorbed = False

        for rank, sid in enumerate(SECTION_ORDER):
            match = subsection_matches.get(sid)
            if match is None or absorbed:
                cache[sid] = self._empty_section(sid)
                continue

            start_idx, win_size = match
            # Skip the anchor window (header/question text) — content starts
            # immediately after the subsection heading block(s).
            content_start = start_idx + win_size
            # Also skip any "header continuation" blocks:
            # - short fragments starting lowercase without sentence-ending
            #   punctuation (e.g. "towarzyszące danym" from a split heading), and
            # - short fragments ending with '?' — all 14 DMP questions end with
            #   '?', so a wrapped question's tail does too (e.g. "accompany data?",
            #   "FAIR (Findable, Accessible, Interoperable, Re-usable)?").
            while content_start < len(blocks):
                blk_text = blocks[content_start].text.strip()
                if not blk_text:
                    break
                is_lower_fragment = blk_text[0].islower() and len(blk_text) < 80 and blk_text[-1] not in '.?!:'
                is_question_tail = blk_text.endswith('?') and len(blk_text) < 120
                if is_lower_fragment or is_question_tail:
                    content_start += 1
                else:
                    break

            # Default end: last block
            content_end = len(blocks)

            # End = IMMEDIATELY NEXT subsection's start (not the next *found* one).
            # If the immediately-next subsection is not found, this section absorbs
            # the rest of the document and all subsequent sections become empty.
            if rank + 1 < len(SECTION_ORDER):
                next_sid = SECTION_ORDER[rank + 1]
                next_match = subsection_matches.get(next_sid)
                if next_match is not None:
                    content_end = next_match[0]
                else:
                    absorbed = True  # this section takes the rest; all after → empty

            # For subsections that are last in their main section, also
            # check if the next section's header appears before content_end.
            if sid in LAST_IN_SECTION:
                next_sec = LAST_IN_SECTION[sid]
                sec_start = section_starts.get(next_sec)
                if sec_start is not None and sec_start > start_idx:
                    content_end = min(content_end, sec_start)

            content_blocks = blocks[content_start:content_end]
            paragraphs = self._cleaner.clean(content_blocks, skip_patterns)
            tagged = [{'text': p, 'tags': [], 'title': ''} for p in paragraphs]
            main = sid.split('.')[0]

            cache[sid] = {
                'section': SECTION_TITLES[main],
                'question': SECTION_QUESTIONS[sid],
                'paragraphs': paragraphs,
                'tagged_paragraphs': tagged,
            }

            logger.debug(
                'Section %s: anchor %d-%d, content %d-%d (%d paragraphs)',
                sid, start_idx, start_idx + win_size - 1,
                content_start, content_end - 1, len(paragraphs),
            )

        # Any text before the first found section anchor → unconnected
        first_match = next(
            (subsection_matches[sid] for sid in SECTION_ORDER
             if subsection_matches.get(sid) is not None),
            None,
        )
        pre_lines: List[str] = []
        if first_match is not None and first_match[0] > 0:
            pre_lines = self._cleaner.clean(blocks[:first_match[0]], skip_patterns)

        cache['_unconnected_text'] = [
            {'text': t, 'type': 'no_section'} for t in pre_lines
        ]
        return cache

    @staticmethod
    def _empty_section(sid: str) -> dict:
        main = sid.split('.')[0]
        return {
            'section': SECTION_TITLES[main],
            'question': SECTION_QUESTIONS[sid],
            'paragraphs': [],
            'tagged_paragraphs': [],
        }

    @staticmethod
    def _smart_filename(file_path: str) -> str:
        base = os.path.splitext(os.path.basename(file_path))[0]
        safe = re.sub(r'[^\w\-]', '_', base)[:50]
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        return f'DMP_{safe}_{ts}'
