"""
utils/extractor_v2.py — DMP Extractor v2 (anchor-based)

Pipeline
--------
1. DocConverter   : document (PDF / DOCX) → flat List[TextBlock]
2. AnchorMatcher  : find block indices for 28 anchor texts (14 PL + 14 EN)
3. ContentCleaner : strip formatting markers and skip-term lines
4. DMPExtractor   : orchestrate; produce JSON cache compatible with review.html
"""

import os
import re
import json
import uuid
import zipfile
import logging
from collections import Counter
from datetime import datetime
from typing import Dict, List, Optional, Set, Tuple

from docx import Document
import PyPDF2

try:
    from pdf2image import convert_from_path
    import pytesseract
    if os.name == 'nt':
        for _p in [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]:
            if os.path.exists(_p):
                pytesseract.pytesseract.tesseract_cmd = _p
                break
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# DMP structure constants  (must match review.html expectations)
# ─────────────────────────────────────────────────────────────────────────────

SECTION_ORDER: List[str] = [
    "1.1", "1.2",
    "2.1", "2.2",
    "3.1", "3.2",
    "4.1", "4.2",
    "5.1", "5.2", "5.3", "5.4",
    "6.1", "6.2",
]

SECTION_TITLES: Dict[str, str] = {
    "1": "1. Data description and collection or re-use of existing data",
    "2": "2. Documentation and data quality",
    "3": "3. Storage and backup during the research process",
    "4": "4. Legal requirements, codes of conduct",
    "5": "5. Data sharing and long-term preservation",
    "6": "6. Data management responsibilities and resources",
}

SECTION_QUESTIONS: Dict[str, str] = {
    "1.1": "How will new data be collected or produced and/or how will existing data be re-used?",
    "1.2": "What data (for example the types, formats, and volumes) will be collected or produced?",
    "2.1": "What metadata and documentation (for example methodology or data collection and way of organising data) will accompany data?",
    "2.2": "What data quality control measures will be used?",
    "3.1": "How will data and metadata be stored and backed up during the research process?",
    "3.2": "How will data security and protection of sensitive data be taken care of during the research?",
    "4.1": "If personal data are processed, how will compliance with legislation on personal data and on data security be ensured?",
    "4.2": "How will other legal issues, such as intellectual property rights and ownership, be managed? What legislation is applicable?",
    "5.1": "How and when will data be shared? Are there possible restrictions to data sharing or embargo reasons?",
    "5.2": "How will data for preservation be selected, and where will data be preserved long-term (for example a data repository or archive)?",
    "5.3": "What methods or software tools will be needed to access and use the data?",
    "5.4": "How will the application of a unique and persistent identifier (such as a Digital Object Identifier (DOI)) to each data set be ensured?",
    "6.1": "Who (for example role, position, and institution) will be responsible for data management (i.e. the data steward)?",
    "6.2": "What resources (for example financial and time) will be dedicated to data management and ensuring the data will be FAIR (Findable, Accessible, Interoperable, Re-usable)?",
}

# ─────────────────────────────────────────────────────────────────────────────
# Text helpers
# ─────────────────────────────────────────────────────────────────────────────

_RE_FMT = re.compile(
    r'\b(?:BOLD|UNDERLINED|ITALIC|UNDERLINED_BOLD)\b:?\s*'
    r'|\[(?:BOLD|UNDERLINED|ITALIC|UNDERLINED_BOLD)\]'
    r'|\{(?:BOLD|UNDERLINED|ITALIC)\}',
    re.IGNORECASE,
)
_RE_SEC_NUM = re.compile(r'^\s*\d+\.\d+\s*')   # "1.1 " prefix
_RE_MAIN_NUM = re.compile(r'^\s*\d+\.\s*')      # "1. " prefix
_RE_WS = re.compile(r'\s+')

# Common words excluded from token-overlap coverage calculation
_STOP: Set[str] = {
    'that', 'will', 'data', 'with', 'from', 'this', 'have', 'been', 'they',
    'what', 'when', 'where', 'which', 'such', 'used', 'example', 'also',
    'danych', 'oraz', 'jest', 'jako', 'przez', 'przy', 'które', 'jakie',
    'jak', 'będą', 'każdego',
}

# Lines that are structural noise — always filtered, regardless of user skip terms.
# Matches: DMP/PZD header, main section titles (Polish + English numbered 1-6)
#
# Fix history:
# - 2026-05-26: Fixed regex for section 2 ("jako" → "jakość")
# - 2026-05-26: Extended patterns with optional suffixes for better matching
# - 2026-05-26: Added .* suffix to match any trailing text after main keywords
_BUILTIN_NOISE: List = [
    re.compile(r'^\s*PLAN\s+ZARZĄDZANIA\s+DANYMI\b', re.IGNORECASE),
    re.compile(r'^\s*DATA\s+MANAGEMENT\s+PLAN\b', re.IGNORECASE),
    # Polish section titles with flexible suffixes (matches anything after main keywords)
    re.compile(
        r'^\s*\d+[\.\s]+(?:'
        r'Opis\s+danych.*'                    # "oraz pozyskiwanie lub ponowne wykorzystanie..."
        r'|Dokumentacja\s+i\s+jakość.*'       # "danych"
        r'|Przechowywanie\s+i\s+tworzenie.*'  # "kopii zapasowych podczas badań"
        r'|Wymogi\s+prawne.*'                 # ", kodeks postępowania"
        r'|Udost[eę]pnianie\s+i\s+d[łl]ugotrwa[łl].*'  # "e przechowywanie danych"
        r'|Zadania\s+zwi[aą]zane.*'           # "z zarządzaniem danymi oraz zasoby"
        r')$',
        re.IGNORECASE,
    ),
    # English section titles with flexible suffixes
    re.compile(
        r'^\s*\d+[\.\s]+(?:'
        r'Data\s+description\s+and.*'         # "collection or re-use of existing data"
        r'|Documentation\s+and\s+data.*'      # "quality"
        r'|Storage\s+and\s+backup.*'          # "during the research process"
        r'|Legal\s+requirements.*'            # ", codes of conduct"
        r'|Data\s+sharing\s+and.*'            # "long-term preservation"
        r'|Data\s+management\s+responsibilities.*'  # "and resources"
        r')$',
        re.IGNORECASE,
    ),
]


def normalize(text: str) -> str:
    """Lowercase, strip section numbers and formatting markers."""
    text = _RE_FMT.sub(' ', text)
    text = _RE_SEC_NUM.sub('', text)
    text = _RE_MAIN_NUM.sub('', text)
    return _RE_WS.sub(' ', text).strip().lower()


def strip_formatting(text: str) -> str:
    """Remove only formatting markers; keep original case and punctuation."""
    return _RE_WS.sub(' ', _RE_FMT.sub('', text)).strip()


def token_overlap(query: str, candidate: str) -> float:
    """
    Fraction of query's content tokens (>= 4 chars, not in _STOP)
    that appear anywhere in the candidate string.
    """
    q_tokens = {w for w in query.split() if len(w) >= 4 and w not in _STOP}
    if not q_tokens:
        return 0.0
    c_tokens = set(candidate.split())
    return len(q_tokens & c_tokens) / len(q_tokens)


# ─────────────────────────────────────────────────────────────────────────────
# Intermediate data model
# ─────────────────────────────────────────────────────────────────────────────

class TextBlock:
    """One logical unit of text extracted from the source document."""
    __slots__ = ('text', 'is_bold', 'is_heading', 'source', 'page', 'is_hf')

    def __init__(
        self,
        text: str,
        is_bold: bool = False,
        is_heading: bool = False,
        source: str = 'paragraph',
        page: int = 0,
        is_hf: bool = False,
    ) -> None:
        self.text = text
        self.is_bold = is_bold
        self.is_heading = is_heading
        self.source = source    # 'paragraph' | 'table' | 'heading'
        self.page = page        # page number (PDF) or sequential position (DOCX)
        self.is_hf = is_hf     # True → identified as header or footer line


# ─────────────────────────────────────────────────────────────────────────────
# Document converter: file → List[TextBlock]
# ─────────────────────────────────────────────────────────────────────────────

class DocConverter:
    """Converts DOCX or PDF to a flat, ordered list of TextBlock objects."""

    def convert(self, file_path: str) -> List[TextBlock]:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.docx':
            return self._from_docx(file_path)
        if ext == '.pdf':
            return self._from_pdf(file_path)
        raise ValueError(f"Unsupported format: {ext}")

    # ── DOCX ─────────────────────────────────────────────────────────────────

    def _from_docx(self, path: str) -> List[TextBlock]:
        doc = Document(path)

        # Collect header / footer text so we can flag it later
        hf_set: Set[str] = set()
        for sec in doc.sections:
            for hf in (
                sec.header, sec.footer,
                sec.even_page_header, sec.even_page_footer,
                sec.first_page_header, sec.first_page_footer,
            ):
                if hf:
                    for para in hf.paragraphs:
                        t = para.text.strip()
                        if t:
                            hf_set.add(t.lower())

        def is_hf(t: str) -> bool:
            return t.strip().lower() in hf_set

        return self._traverse_body(doc, is_hf)

    def _traverse_body(self, doc: Document, is_hf) -> List[TextBlock]:
        """
        Walk the body element tree in document order.
        Each top-level <w:p> becomes one TextBlock.
        Each <w:tbl> row becomes one TextBlock (cells joined with ' | ').
        """
        from docx.oxml.ns import qn

        # Build lookup: lxml element → python-docx Paragraph
        # doc.paragraphs contains only top-level body paragraphs (not table cells)
        para_map = {p._element: p for p in doc.paragraphs}

        blocks: List[TextBlock] = []
        pos = 0

        for child in doc.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                para = para_map.get(child)
                if para is None:
                    continue
                text = para.text.strip()
                if not text:
                    continue
                bold = self._is_para_bold(para)
                heading = (
                    para.style.name.lower().startswith('heading')
                    or (bold and len(text) < 200)
                )
                source = (
                    'heading'
                    if para.style.name.lower().startswith('heading')
                    else 'paragraph'
                )
                blocks.append(TextBlock(text, bold, heading, source, pos, is_hf(text)))
                pos += 1

            elif tag == 'tbl':
                # Each table row → one block; duplicated merged-cell text is deduplicated
                for row_el in child.findall('.//' + qn('w:tr')):
                    seen: Set[str] = set()
                    parts: List[str] = []
                    for cell in row_el.findall(qn('w:tc')):
                        cell_words = []
                        for p_el in cell.findall('.//' + qn('w:p')):
                            t = ''.join(
                                r.text or ''
                                for r in p_el.findall('.//' + qn('w:t'))
                            ).strip()
                            if t:
                                cell_words.append(t)
                        cell_text = ' '.join(cell_words)
                        if cell_text and cell_text not in seen:
                            parts.append(cell_text)
                            seen.add(cell_text)
                    if parts:
                        text = ' | '.join(parts)
                        blocks.append(
                            TextBlock(text, False, False, 'table', pos, is_hf(text))
                        )
                        pos += 1

        return blocks

    @staticmethod
    def _is_para_bold(para) -> bool:
        runs = [r for r in para.runs if r.text.strip()]
        return bool(runs) and all(r.bold for r in runs)

    # ── PDF ──────────────────────────────────────────────────────────────────

    def _from_pdf(self, path: str) -> List[TextBlock]:
        pages_text = self._read_pdf_pages(path)
        hf_set = self._detect_pdf_hf(pages_text)

        blocks: List[TextBlock] = []
        for page_num, page_text in enumerate(pages_text):
            for line in page_text.split('\n'):
                text = line.strip()
                if len(text) < 3:
                    continue
                blocks.append(
                    TextBlock(text, False, False, 'paragraph', page_num,
                              text.lower() in hf_set)
                )
        return blocks

    def _read_pdf_pages(self, path: str) -> List[str]:
        pages: List[str] = []
        try:
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    pages.append(page.extract_text() or '')
        except Exception as exc:
            raise RuntimeError(f"PDF read error: {exc}") from exc

        if sum(len(t) for t in pages) < 100:
            if HAS_OCR:
                return self._ocr(path)
            raise RuntimeError(
                "PDF appears to be a scanned image but OCR (pytesseract + pdf2image) "
                "is not installed."
            )
        return pages

    @staticmethod
    def _detect_pdf_hf(pages: List[str]) -> Set[str]:
        """Return normalized lines that appear on more than half the pages."""
        if len(pages) < 2:
            return set()
        counter: Counter = Counter()
        for page in pages:
            seen: Set[str] = set()
            for line in page.split('\n'):
                norm = line.strip().lower()
                if norm and norm not in seen:
                    counter[norm] += 1
                    seen.add(norm)
        threshold = max(2, len(pages) // 2)
        return {line for line, cnt in counter.items() if cnt >= threshold}

    @staticmethod
    def _ocr(path: str) -> List[str]:
        images = convert_from_path(path)
        return [pytesseract.image_to_string(img, lang='pol+eng') for img in images]


# ─────────────────────────────────────────────────────────────────────────────
# Anchor matcher: find block indices for each DMP subsection
# ─────────────────────────────────────────────────────────────────────────────

class AnchorMatcher:
    """
    Searches for 28 anchor texts (14 PL + 14 EN) in a block list and returns
    the block index at which each subsection begins.
    """

    HIGH_THRESHOLD  = 0.55  # score at which we accept immediately (first match wins)
    LOW_THRESHOLD   = 0.35  # fallback minimum to accept any match
    MIN_FIRST_BLOCK = 0.20  # block[i] alone must cover ≥20% of anchor tokens;
                            # prevents content blocks from starting a wide window
                            # that only contains the anchor text further down

    def __init__(self, anchors_cfg: dict) -> None:
        self._anchors = anchors_cfg

    def find_boundaries(self, blocks: List[TextBlock]) -> Dict[str, Tuple[int, int]]:
        """
        Returns {section_id: (block_index, window_size)} in document order.
        - block_index: first block where anchor/question begins
        - window_size: number of blocks that contain the anchor text

        Each section is searched only in the portion of the document AFTER the
        previous section's anchor — forward-only search prevents a later duplicate
        of anchor text (e.g. a re-printed template at the end of the file) from
        stealing the boundary from the real occurrence.
        """
        result: Dict[str, Tuple[int, int]] = {}
        search_from = 0
        for sid in SECTION_ORDER:
            cfg = self._anchors.get(sid)
            if not cfg:
                continue
            texts = cfg.get('pl', []) + cfg.get('en', [])
            fingerprint = cfg.get('fingerprint_pl', []) + cfg.get('fingerprint_en', [])
            rel_idx, win_size, score = self._locate(blocks[search_from:], texts, fingerprint)
            if rel_idx >= 0:
                abs_idx = search_from + rel_idx
                result[sid] = (abs_idx, win_size)
                search_from = abs_idx + win_size  # Skip past entire matched window
                logger.debug('Anchor %s → blocks %d-%d (window=%d, score %.2f)',
                           sid, abs_idx, abs_idx + win_size - 1, win_size, score)
            else:
                logger.debug('Anchor %s → NOT FOUND', sid)
        return result

    def _locate(
        self,
        blocks: List[TextBlock],
        anchor_texts: List[str],
        fingerprint: List[str],
    ) -> Tuple[int, int, float]:
        """
        Find the FIRST position (left-to-right) whose token_overlap with any anchor
        reaches HIGH_THRESHOLD and return it immediately.  If no position reaches
        HIGH, return the first position that reached LOW.  This prevents a later
        high-scoring duplicate from stealing the boundary.

        Two guards ensure the anchor STARTS at block[i] rather than somewhere
        inside a wide window:
          1. Skip blocks that are structural noise (main section titles).
          2. Require block[i] alone to cover ≥ MIN_FIRST_BLOCK of anchor tokens.
             This prevents a content block followed 3 blocks later by the real
             question text from being accepted as an anchor start.

        Returns: (start_idx, matched_window_size, score)
        - start_idx: first block index where anchor begins
        - matched_window_size: number of blocks that contain the anchor text
        - score: match confidence score
        """
        norm_anchors = [normalize(a) for a in anchor_texts]
        first_low_idx, first_low_win, first_low_score = -1, 1, 0.0

        # Strategy 1: for each position i, try windows of 1/2/3/4 blocks.
        # Return the moment we hit HIGH; record first LOW as fallback.
        for i in range(len(blocks)):
            blk = blocks[i]

            # Skip header/footer lines
            if blk.is_hf:
                continue

            # Skip main section title lines (e.g. "3. Przechowywanie i tworzenie…")
            # These score highly against subsection anchors but are not the anchor.
            if any(p.search(blk.text) for p in _BUILTIN_NOISE):
                continue

            # The anchor must START in this block — compute single-block score first.
            # If block[i] has no anchor tokens, a window of 4 could still score high
            # by pulling in the actual question block further down, which would
            # consume the preceding content inside the "anchor window".
            norm_blk = normalize(blk.text)
            best_single = max(
                (token_overlap(na, norm_blk) for na in norm_anchors),
                default=0.0,
            )
            if best_single < self.MIN_FIRST_BLOCK:
                continue

            best_at_i = 0.0
            best_win = 1
            best_weight = 0.0  # score × anchor word-count — proxy for |matched tokens|
            for win_size in (1, 2, 3, 4):
                if i + win_size > len(blocks):
                    break
                window = ' '.join(
                    normalize(blocks[j].text) for j in range(i, i + win_size)
                )
                for na in norm_anchors:
                    s = token_overlap(na, window)
                    # weight = matched-token count proxy.  When two windows share the
                    # same ratio (e.g. a short 3-token anchor scores 1.0 on win=1
                    # and a long 9-token anchor also scores 1.0 on win=2), prefer
                    # the window that covers more anchor tokens absolutely.
                    # If adding extra blocks brings no extra anchor tokens, the weight
                    # stays the same and best_win is NOT increased (no bloat).
                    weight = s * len(na.split())
                    if s > best_at_i or (s == best_at_i and weight > best_weight):
                        best_at_i = s
                        best_win = win_size
                        best_weight = weight
            if best_at_i >= self.HIGH_THRESHOLD:
                return i, best_win, best_at_i
            if best_at_i >= self.LOW_THRESHOLD and first_low_idx < 0:
                first_low_idx, first_low_win, first_low_score = i, best_win, best_at_i

        # Strategy 2: fingerprint keyword coverage over a 6-block window.
        # Only runs when strategy 1 found nothing at all.
        if first_low_idx < 0 and fingerprint:
            fp_tokens = {w.lower() for w in fingerprint if len(w) >= 4}
            if fp_tokens:
                for i in range(len(blocks)):
                    window = ' '.join(
                        normalize(blocks[j].text)
                        for j in range(i, min(i + 6, len(blocks)))
                    )
                    window_tokens = set(window.split())
                    score = len(fp_tokens & window_tokens) / len(fp_tokens)
                    if score >= self.HIGH_THRESHOLD:
                        return i, 6, score  # Fingerprint uses 6-block window
                    if score >= self.LOW_THRESHOLD and first_low_idx < 0:
                        first_low_idx, first_low_win, first_low_score = i, 6, score

        if first_low_idx >= 0:
            return first_low_idx, first_low_win, first_low_score
        return -1, 1, 0.0


# ─────────────────────────────────────────────────────────────────────────────
# Content cleaner
# ─────────────────────────────────────────────────────────────────────────────

class ContentCleaner:
    """Converts TextBlock objects to cleaned display-ready text lines."""

    def clean(self, blocks: List[TextBlock], skip_patterns: list) -> List[str]:
        """
        Return cleaned non-empty text lines from *blocks*, skipping:
          - header / footer lines
          - subsection numerations (1.1, 2.2) and main section numbers (1., 2.)
          - structural noise (DMP title, main section titles)
          - lines matching any pattern in skip_patterns
          - lines that are empty after formatting removal
        """
        result = []
        for block in blocks:
            if block.is_hf:
                continue
            text = strip_formatting(block.text)
            # Remove subsection numerations (1.1, 2.2) and main section numbers (1., 2.)
            # These are identifiers only, NOT part of the content
            text = _RE_SEC_NUM.sub('', text)   # Remove "1.1 ", "2.2 " prefix
            text = _RE_MAIN_NUM.sub('', text)  # Remove "1. ", "2. " prefix
            text = _RE_WS.sub(' ', text).strip()
            if not text:
                continue
            if any(p.search(text) for p in _BUILTIN_NOISE):
                continue
            if any(p.search(text) for p in skip_patterns):
                continue
            result.append(text)
        return result


# ─────────────────────────────────────────────────────────────────────────────
# Skip-terms manager
# ─────────────────────────────────────────────────────────────────────────────

class SkipTermsManager:
    """CRUD for the user-editable skip-terms list (config/extraction_skip_terms.json)."""

    _DEFAULT_PATH = os.path.normpath(
        os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            '..', 'config', 'extraction_skip_terms.json',
        )
    )

    def __init__(self, path: Optional[str] = None) -> None:
        self.path = path or self._DEFAULT_PATH

    def load(self) -> List[str]:
        if not os.path.exists(self.path):
            return []
        with open(self.path, encoding='utf-8') as f:
            return json.load(f).get('terms', [])

    def save(self, terms: List[str]) -> None:
        os.makedirs(os.path.dirname(self.path), exist_ok=True)
        with open(self.path, 'w', encoding='utf-8') as f:
            json.dump({'terms': terms}, f, ensure_ascii=False, indent=2)

    def add(self, term: str) -> List[str]:
        terms = self.load()
        if term and term not in terms:
            terms.append(term)
            self.save(terms)
        return terms

    def remove(self, term: str) -> List[str]:
        terms = [t for t in self.load() if t != term]
        self.save(terms)
        return terms

    def compile(self) -> list:
        """Return a list of compiled regex objects for all skip terms."""
        patterns = []
        for term in self.load():
            try:
                patterns.append(re.compile(term, re.IGNORECASE))
            except re.error:
                patterns.append(re.compile(re.escape(term), re.IGNORECASE))
        return patterns


# ─────────────────────────────────────────────────────────────────────────────
# Validation helpers  (also imported by app.py)
# ─────────────────────────────────────────────────────────────────────────────

def validate_docx_file(path: str) -> Tuple[bool, str]:
    if not os.path.exists(path):
        return False, 'File not found'
    if os.path.getsize(path) == 0:
        return False, 'File is empty'
    try:
        with zipfile.ZipFile(path) as z:
            for req in ('word/document.xml', '[Content_Types].xml'):
                if req not in z.namelist():
                    return False, f'Invalid DOCX: missing {req}'
    except zipfile.BadZipFile:
        return False, 'Not a valid DOCX / ZIP file'
    return True, 'OK'


def validate_pdf_file(path: str) -> Tuple[bool, str]:
    if not os.path.exists(path):
        return False, 'File not found'
    if os.path.getsize(path) == 0:
        return False, 'File is empty'
    with open(path, 'rb') as f:
        if f.read(4) != b'%PDF':
            return False, 'Not a valid PDF file'
    return True, 'OK'


# ─────────────────────────────────────────────────────────────────────────────
# Main extractor  (drop-in replacement for the old DMPExtractor)
# ─────────────────────────────────────────────────────────────────────────────

class DMPExtractor:
    """
    Anchor-based DMP section extractor.

    Public interface identical to the old extractor:
        result = DMPExtractor().process_file(
            file_path, output_dir, progress_callback=None
        )

    Returns:
        {'success': True, 'filename': str, 'cache_id': str, 'message': str}
        {'success': False, 'message': str}
    """

    def __init__(self) -> None:
        self._anchors_cfg = self._load_anchors()
        self._converter = DocConverter()
        self._matcher = AnchorMatcher(self._anchors_cfg.get('sections', {}))
        self._cleaner = ContentCleaner()
        self._skip_mgr = SkipTermsManager()

    # ── Public API ────────────────────────────────────────────────────────────

    def process_file(
        self,
        file_path: str,
        output_dir: str,
        progress_callback=None,
    ) -> dict:
        def cb(msg: str, pct: int) -> None:
            if progress_callback:
                progress_callback(msg, pct)

        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.docx':
                ok, msg = validate_docx_file(file_path)
            elif ext == '.pdf':
                ok, msg = validate_pdf_file(file_path)
            else:
                return {'success': False, 'message': f'Unsupported file type: {ext}'}
            if not ok:
                return {'success': False, 'message': msg}

            cb('Converting document to text blocks…', 10)
            blocks = self._converter.convert(file_path)
            logger.info(
                'DocConverter produced %d blocks from %s', len(blocks), file_path
            )

            cb('Locating section boundaries…', 30)
            boundaries = self._matcher.find_boundaries(blocks)
            logger.info(
                'Found %d / %d section boundaries',
                len(boundaries), len(SECTION_ORDER),
            )

            cb('Extracting content between boundaries…', 60)
            skip_patterns = self._skip_mgr.compile()
            cache = self._build_cache(blocks, boundaries, skip_patterns)

            cb('Saving cache…', 85)
            cache_id = str(uuid.uuid4())
            cache_dir = os.path.join(output_dir, 'cache')
            os.makedirs(cache_dir, exist_ok=True)
            cache_path = os.path.join(cache_dir, f'cache_{cache_id}.json')
            with open(cache_path, 'w', encoding='utf-8') as f:
                json.dump(cache, f, ensure_ascii=False, indent=2)

            filled = sum(
                1 for sid in SECTION_ORDER
                if cache.get(sid, {}).get('paragraphs')
            )
            cb('Done.', 100)

            return {
                'success': True,
                'filename': self._smart_filename(file_path),
                'cache_id': cache_id,
                'cache_file': f'cache_{cache_id}.json',
                'message': f'Extracted {filled} of {len(SECTION_ORDER)} sections',
            }

        except Exception as exc:
            logger.exception('process_file failed for %s', file_path)
            return {'success': False, 'message': str(exc)}

    # ── Internal ──────────────────────────────────────────────────────────────

    def _build_cache(
        self,
        blocks: List[TextBlock],
        boundaries: Dict[str, Tuple[int, int]],
        skip_patterns: list,
    ) -> dict:
        """
        Slice *blocks* into sections using *boundaries*, clean each slice,
        and assemble the JSON cache structure expected by review.html.

        boundaries format: {section_id: (start_idx, window_size)}
        - start_idx: first block of the anchor/question
        - window_size: number of blocks occupied by the anchor/question
        """
        cache: dict = {}

        if not boundaries:
            # Nothing found — put everything as unconnected text
            lines = self._cleaner.clean(blocks, skip_patterns)
            for sid in SECTION_ORDER:
                cache[sid] = self._empty_section(sid)
            cache['_unconnected_text'] = [
                {'text': t, 'type': 'no_section'} for t in lines
            ]
            return cache

        # Sort by start_idx (first element of tuple)
        ordered = sorted(boundaries.items(), key=lambda x: x[1][0])

        # Text before the first found anchor → unconnected
        first_start_idx, _ = ordered[0][1]
        pre_lines = self._cleaner.clean(blocks[:first_start_idx], skip_patterns)

        # Slice each found section
        for rank, (sid, (start_idx, win_size)) in enumerate(ordered):
            # Content begins AFTER the entire anchor window
            content_start = start_idx + win_size

            # Content ends at the start of next anchor (or end of document)
            if rank + 1 < len(ordered):
                next_start_idx, _ = ordered[rank + 1][1]
                content_end = next_start_idx
            else:
                content_end = len(blocks)

            # Extract content blocks (excludes ALL anchor blocks)
            content_blocks = blocks[content_start:content_end]
            paragraphs = self._cleaner.clean(content_blocks, skip_patterns)
            tagged = [{'text': p, 'tags': [], 'title': ''} for p in paragraphs]
            main = sid.split('.')[0]
            cache[sid] = {
                'section': SECTION_TITLES[main],
                'question': SECTION_QUESTIONS[sid],
                'paragraphs': paragraphs,
                'tagged_paragraphs': tagged,
            }

            logger.debug(
                'Section %s: anchor blocks %d-%d, content blocks %d-%d (%d paragraphs)',
                sid, start_idx, start_idx + win_size - 1,
                content_start, content_end - 1, len(paragraphs)
            )

        # Fill any sections that were not found
        for sid in SECTION_ORDER:
            if sid not in cache:
                cache[sid] = self._empty_section(sid)

        cache['_unconnected_text'] = [
            {'text': t, 'type': 'no_section'} for t in pre_lines
        ]
        return cache

    @staticmethod
    def _empty_section(sid: str) -> dict:
        main = sid.split('.')[0]
        return {
            'section': SECTION_TITLES[main],
            'question': SECTION_QUESTIONS[sid],
            'paragraphs': [],
            'tagged_paragraphs': [],
        }

    def _load_anchors(self) -> dict:
        path = os.path.normpath(
            os.path.join(
                os.path.dirname(os.path.abspath(__file__)),
                '..', 'config', 'dmp_anchors.json',
            )
        )
        with open(path, encoding='utf-8') as f:
            return json.load(f)

    @staticmethod
    def _smart_filename(file_path: str) -> str:
        base = os.path.splitext(os.path.basename(file_path))[0]
        safe = re.sub(r'[^\w\-]', '_', base)[:50]
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        return f'DMP_{safe}_{ts}'
