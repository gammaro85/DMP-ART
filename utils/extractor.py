# utils/extractor.py - Enhanced version with improved DOCX and PDF processing
import os
import re
import json
import uuid
import zipfile
from docx import Document
from datetime import datetime
import PyPDF2
import logging

# Optional OCR dependencies
try:
    from pdf2image import convert_from_path
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


def validate_docx_file(file_path):
    """
    Enhanced DOCX file validation.

    Args:
        file_path: Path to the DOCX file

    Returns:
        tuple: (is_valid: bool, message: str)
    """
    try:
        if not os.path.exists(file_path):
            return False, "File does not exist"

        if not file_path.lower().endswith('.docx'):
            return False, "File is not a DOCX file"

        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return False, "File is empty"

        if file_size > 16 * 1024 * 1024:  # 16MB limit
            return False, "File is too large (max 16MB)"

        # Check if it's a valid ZIP file (DOCX is ZIP-based)
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                file_list = zip_file.namelist()
                required_files = ['word/document.xml', '[Content_Types].xml']

                for required_file in required_files:
                    if required_file not in file_list:
                        return False, f"Invalid DOCX structure: missing {required_file}"

        except zipfile.BadZipFile:
            return False, "File is not a valid ZIP archive"
        except Exception as e:
            return False, f"ZIP validation error: {str(e)}"

        # Try to load with python-docx
        try:
            doc = Document(file_path)
            table_count = len(doc.tables)

            # Check for minimum content
            has_content = any(p.text.strip() for p in doc.paragraphs) or table_count > 0

            if not has_content:
                return False, "Document appears to be empty or contains no readable content"

        except Exception as e:
            return False, f"Document processing error: {str(e)}"

        return True, "File is valid"

    except Exception as e:
        return False, f"Unexpected validation error: {str(e)}"


class DMPExtractor:
    def __init__(self, debug_mode=False):
        """
        Initialize DMP Extractor

        Args:
            debug_mode (bool): Enable verbose logging for debugging (default: False)
        """
        self.debug_mode = debug_mode

        # Define start and end markers for extraction
        self.start_marks = [
            "DATA MANAGEMENT PLAN",
            "DATA MANAGEMENT PLAN [in English]",
            "PLAN ZARZĄDZANIA DANYMI",
            ]
        self.end_marks = [
            "ADMINISTRATIVE DECLARATIONS",
            "OŚWIADCZENIA ADMINISTRACYJNE"
        ]

        # Performance optimization: Pre-compile regex patterns
        self._compile_regex_patterns()
        
        # Bilingual section mapping
        self.section_mapping = {
            "Opis danych oraz pozyskiwanie": "1. Data description and collection or re-use of existing data",
            "Dokumentacja i jakość danych": "2. Documentation and data quality",
            "Przechowywanie i tworzenie kopii zapasowych": "3. Storage and backup during the research process", 
            "Wymogi prawne, kodeks postępowania": "4. Legal requirements, codes of conduct",
            "Udostępnianie i długotrwałe przechowywanie": "5. Data sharing and long-term preservation",
            "Zadania związane z zarządzaniem danymi": "6. Data management responsibilities and resources"
        }
        
        # Bilingual subsection mapping with exact Polish subsections
        raw_subsection_mapping = {
            # 1. Data description section
            "Sposób pozyskiwania i opracowywania nowych danych i/lub ponownego wykorzystania dostępnych danych": 
                "How will new data be collected or produced and/or how will existing data be re-used?",
            "Pozyskiwane lub opracowywane dane (np. rodzaj, format, ilość)": 
                "What data (for example the types, formats, and volumes) will be collected or produced?",
            
            # 2. Documentation and data quality
            "Metadane i dokumenty (np. metodologia lub pozyskiwanie danych oraz sposób porządkowania danych) towarzyszące danym": 
                "What metadata and documentation (for example methodology or data collection and way of organising data) will accompany data?",
            "Stosowane środki kontroli jakości danych": 
                "What data quality control measures will be used?",
            
            # 3. Storage and backup
            "Przechowywanie i tworzenie kopii zapasowych danych i metadanych podczas badań": 
                "How will data and metadata be stored and backed up during the research process?",
            "Sposób zapewnienia bezpieczeństwa danych oraz ochrony danych wrażliwych podczas badań": 
                "How will data security and protection of sensitive data be taken care of during the research?",
            
            # 4. Legal requirements
            "Sposób zapewnienia zgodności z przepisami dotyczącymi danych osobowych i bezpieczeństwa danych w przypadku przetwarzania danych osobowych": 
                "If personal data are processed, how will compliance with legislation on personal data and on data security be ensured?",
            "Sposób zarządzania innymi kwestiami prawnymi, np. prawami własności intelektualnej lub własnością. Obowiązujące przepisy": 
                "How will other legal issues, such as intelectual property rights and ownership, be managed? What legislation is applicable?",
            
            # 5. Data sharing and preservation
            "Sposób i termin udostępnienia danych. Ewentualne ograniczenia w udostępnianiu danych lub przyczyny embarga": 
                "How and when will data be shared? Are there possible restrictions to data sharing or embargo reasons?",
            "Sposób wyboru danych przeznaczonych do przechowania oraz miejsce długotrwałego przechowywania danych (np. repozytorium lub archiwum danych)": 
                "How will data for preservation be selected, and where will data be preserved long-term (for example a data repository or archive)?",
            "Metody lub narzędzia programowe umożliwiające dostęp do danych i korzystanie z danych": 
                "What methods or software tools will be needed to access and use the data?",
            "Sposób zapewniający stosowanie unikalnego i trwałego identyfikatora (np. cyfrowego identyfikatora obiektu (DOI)) dla każdego zestawu danych": 
                "How will the application of a unique and persistent identifier (such us a Digital Object Identifier (DOI)) to each data set be ensured?",
            
            # 6. Data management responsibilities
            "Osoba (np. funkcja, stanowisko i instytucja) odpowiedzialna za zarządzanie danymi (np. data steward)": 
                "Who (for example role, position, and institution) will be responsible for data management (i.e the data steward)?",
            "Środki (np. finansowe i czasowe) przeznaczone do zarządzania danymi i zapewnienia możliwości odnalezienia, dostępu, interoperacyjności i ponownego wykorzystania danych": 
                "What resources (for example financial and time) will be dedicated to data management and ensuring the data will be FAIR (Findable, Accessible, Interoperable, Re-usable)?"
        }
        
        # Create normalized versions of subsection mapping
        self.normalized_subsection_mapping = {}
        self.subsection_mapping = raw_subsection_mapping.copy()
        
        for polish, english in raw_subsection_mapping.items():
            # Remove trailing colon if present
            normalized_polish = polish[:-1].strip() if polish.endswith(':') else polish.strip()
            # Convert to lowercase for case-insensitive matching
            normalized_polish = normalized_polish.lower()
            self.normalized_subsection_mapping[normalized_polish] = english
            
            # Also add versions with/without trailing colon to main mapping
            if not polish.endswith(':'):
                self.subsection_mapping[polish + ':'] = english
            else:
                self.subsection_mapping[polish[:-1].strip()] = english
        
        # Define the structure of DMP sections and questions
        self.dmp_structure = {
            "1. Data description and collection or re-use of existing data": [
                "How will new data be collected or produced and/or how will existing data be re-used?",
                "What data (for example the types, formats, and volumes) will be collected or produced?"
            ],
            "2. Documentation and data quality": [
                "What metadata and documentation (for example methodology or data collection and way of organising data) will accompany data?",
                "What data quality control measures will be used?"
            ],
            "3. Storage and backup during the research process": [
                "How will data and metadata be stored and backed up during the research process?",
                "How will data security and protection of sensitive data be taken care of during the research?"
            ],
            "4. Legal requirements, codes of conduct": [
                "If personal data are processed, how will compliance with legislation on personal data and on data security be ensured?",
                "How will other legal issues, such as intelectual property rights and ownership, be managed? What legislation is applicable?"
            ],
            "5. Data sharing and long-term preservation": [
                "How and when will data be shared? Are there possible restrictions to data sharing or embargo reasons?",
                "How will data for preservation be selected, and where will data be preserved long-term (for example a data repository or archive)?",
                "What methods or software tools will be needed to access and use the data?",
                "How will the application of a unique and persistent identifier (such us a Digital Object Identifier (DOI)) to each data set be ensured?"
            ],
            "6. Data management responsibilities and resources": [
                "Who (for example role, position, and institution) will be responsible for data management (i.e the data steward)?",
                "What resources (for example financial and time) will be dedicated to data management and ensuring the data will be FAIR (Findable, Accessible, Interoperable, Re-usable)?"
            ]
        }
        
        # Map section numbers to IDs for the review interface
        self.section_ids = {
            "1.1": "How will new data be collected or produced and/or how will existing data be re-used?",
            "1.2": "What data (for example the types, formats, and volumes) will be collected or produced?",
            "2.1": "What metadata and documentation (for example methodology or data collection and way of organising data) will accompany data?",
            "2.2": "What data quality control measures will be used?",
            "3.1": "How will data and metadata be stored and backed up during the research process?",
            "3.2": "How will data security and protection of sensitive data be taken care of during the research?",
            "4.1": "If personal data are processed, how will compliance with legislation on personal data and on data security be ensured?",
            "4.2": "How will other legal issues, such as intelectual property rights and ownership, be managed? What legislation is applicable?",
            "5.1": "How and when will data be shared? Are there possible restrictions to data sharing or embargo reasons?",
            "5.2": "How will data for preservation be selected, and where will data be preserved long-term (for example a data repository or archive)?",
            "5.3": "What methods or software tools will be needed to access and use the data?",
            "5.4": "How will the application of a unique and persistent identifier (such us a Digital Object Identifier (DOI)) to each data set be ensured?",
            "6.1": "Who (for example role, position, and institution) will be responsible for data management (i.e the data steward)?",
            "6.2": "What resources (for example financial and time) will be dedicated to data management and ensuring the data will be FAIR (Findable, Accessible, Interoperable, Re-usable)?"
        }
        
        # Define key phrases to identify and tag in paragraphs
        # Key phrases functionality removed

        # Performance optimization: Text similarity cache
        self._similarity_cache = {}

        # Performance optimization: Pre-compute subsection word index
        self._subsection_word_index = self._build_subsection_word_index()

    def _build_subsection_word_index(self):
        """
        Pre-compute word sets for each subsection for faster matching

        This optimization reduces the complexity of subsection detection from O(n*m)
        to O(n) by computing the word sets once during initialization instead of
        on every detection attempt.

        Returns:
            dict: Mapping of subsection text to set of significant words
        """
        index = {}
        common_words = {'data', 'will', 'used', 'such', 'example', 'how', 'what', 'are'}

        for section, subsections in self.dmp_structure.items():
            for subsection in subsections:
                # Extract significant words (length > 3, not common words)
                words = set(
                    word.lower()
                    for word in subsection.split()
                    if len(word) > 3 and word.lower() not in common_words
                )
                index[subsection] = words

        self._log_debug(f"Built subsection word index with {len(index)} entries")
        return index

    def _compile_regex_patterns(self):
        """Pre-compile regex patterns for better performance"""
        # Skip patterns - compiled once at initialization
        self.skip_patterns_compiled = [
            re.compile(r"Strona \d+", re.IGNORECASE),
            re.compile(r"Page \d+", re.IGNORECASE),
            re.compile(r"ID:\s*\d+", re.IGNORECASE),
            re.compile(r"\[wydruk roboczy\]", re.IGNORECASE),
            re.compile(r"WZÓR", re.IGNORECASE),
            re.compile(r"W Z Ó R", re.IGNORECASE),
            re.compile(r"OSF,", re.IGNORECASE),
            re.compile(r"^\d+$"),
            re.compile(r"^\+[-=]+\+$"),
            re.compile(r"^\|[\s\|]*\|$"),
            re.compile(r"Dół formularza", re.IGNORECASE),
            re.compile(r"Początek formularza", re.IGNORECASE),
            re.compile(r"^\s*Dół\s+formularza\s*$", re.IGNORECASE),
            re.compile(r"^\s*Początek\s+formularza\s*$", re.IGNORECASE),
        ]

        # PDF-specific patterns
        self.pdf_skip_patterns_compiled = [
            re.compile(r"wydruk roboczy", re.IGNORECASE),
            re.compile(r"Strona \d+ z \d+", re.IGNORECASE),
            re.compile(r"TAK\s*NIE\s*$", re.IGNORECASE),
            re.compile(r"^\s*[✓✗×]\s*$"),
            re.compile(r"^\s*\[\s*[Xx]?\s*\]\s*$"),
            re.compile(r"^\s*_{3,}\s*$"),
            re.compile(r"^\.{3,}$"),
            re.compile(r"^\s*data\s*:\s*$", re.IGNORECASE),
            re.compile(r"^\s*podpis\s*:\s*$", re.IGNORECASE),
            re.compile(r"OSF,?\s*OPUS-\d+\s*Strona\s+\d+\s*ID:\s*\d+,?\s*\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}", re.IGNORECASE),
            re.compile(r"OSF,?\s*OPUS-\d+\s*Strona", re.IGNORECASE),
            re.compile(r"\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}"),
        ]

        # Header/footer component patterns
        self.header_component_patterns = {
            'osf': re.compile(r'OSF,?\s*', re.IGNORECASE),
            'opus': re.compile(r'OPUS-\d+', re.IGNORECASE),
            'page': re.compile(r'Strona\s+\d+', re.IGNORECASE),
            'id': re.compile(r'ID:\s*\d+', re.IGNORECASE),
            'date': re.compile(r'\d{4}-\d{2}-\d{2}'),
            'time': re.compile(r'\d{2}:\d{2}:\d{2}'),
            'timestamp': re.compile(r'\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}'),
        }

        # Other useful patterns
        self.numbered_section_pattern = re.compile(r'^\s*(\d+)\.\s*(.*?)$')
        self.markup_patterns = {
            'underline': re.compile(r'\[([^]]+)\]\{\.underline\}'),
            'bold': re.compile(r'\*\*([^*]+)\*\*'),
            'underline2': re.compile(r'__([^_]+)__'),
            'mark': re.compile(r'\{\.mark\}'),
        }

    def _log_debug(self, message):
        """Log debug messages only if debug_mode is enabled"""
        if self.debug_mode:
            print(message)

    def validate_docx_file(self, file_path):
        """Validate DOCX file integrity"""
        try:
            if not os.path.exists(file_path):
                return False, "File does not exist"
            
            if not file_path.lower().endswith('.docx'):
                return False, "File is not a DOCX file"
            
            # Check if it's a valid ZIP file (DOCX is ZIP-based)
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    file_list = zip_file.namelist()
                    if 'word/document.xml' not in file_list:
                        return False, "Invalid DOCX structure: missing document.xml"
            except zipfile.BadZipFile:
                return False, "File is not a valid ZIP archive"
            
            # Try to load with python-docx
            doc = Document(file_path)
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            if paragraph_count == 0 and table_count == 0:
                return False, "Document appears to be empty"
            
            return True, "File is valid"
            
        except Exception as e:
            return False, f"Validation error: {str(e)}"
    
    def clean_table_delimiters(self, text):
        """Remove table formatting artifacts"""
        # Remove table border characters
        text = re.sub(r'\+[-=]+\+', '', text)
        text = re.sub(r'\|[\s]*\|', '', text)
        text = re.sub(r'^\||\|$', '', text, flags=re.MULTILINE)
        
        # Clean up excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        return text.strip()
    
    def extract_table_content(self, doc):
        """Extract content from table structures in DOCX files"""
        table_content = []
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and not self.should_skip_text(cell_text):
                        # Clean table delimiters
                        cell_text = self.clean_table_delimiters(cell_text)
                        
                        if not cell_text:
                            continue
                        
                        # Check for formatting in cell runs
                        is_bold = any(run.bold for paragraph in cell.paragraphs for run in paragraph.runs if run.bold)
                        is_underlined = any(run.underline for paragraph in cell.paragraphs for run in paragraph.runs if run.underline)
                        
                        if is_bold and is_underlined:
                            table_content.append(f"UNDERLINED_BOLD:{cell_text}")
                        elif is_bold:
                            table_content.append(f"BOLD:{cell_text}")
                        elif is_underlined:
                            table_content.append(f"UNDERLINED:{cell_text}")
                        else:
                            table_content.append(cell_text)
        
        return table_content
    
    def _report_progress(self, callback, message, progress):
        """
        Safely report progress through callback if provided

        Args:
            callback: Progress callback function (message, progress) or None
            message (str): Human-readable progress message
            progress (int): Progress percentage (0-100)
        """
        if callback and callable(callback):
            try:
                callback(message, progress)
            except Exception as e:
                self._log_debug(f"Progress callback error: {str(e)}")

    def process_file(self, file_path, output_dir, progress_callback=None):
        """
        Process a file and extract DMP content based on file type

        Args:
            file_path (str): Path to the input file (PDF or DOCX)
            output_dir (str): Directory to save output files
            progress_callback (callable, optional): Function(message, progress) for progress updates

        Returns:
            dict: Processing result with success status and file information
        """
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == '.pdf':
            return self.process_pdf(file_path, output_dir, progress_callback)
        elif file_extension == '.docx':
            return self.process_docx(file_path, output_dir, progress_callback)
        else:
            return {
                "success": False,
                "message": f"Unsupported file type: {file_extension}"
            }
    
    def should_skip_text(self, text, is_pdf=False):
        """Determine if text should be skipped (headers, footers, etc.)
        Optimized version using pre-compiled regex patterns"""

        # Check basic patterns first using pre-compiled patterns
        for pattern in self.skip_patterns_compiled:
            if pattern.search(text) is not None:
                return True

        # Additional PDF-specific patterns
        if is_pdf:
            for pattern in self.pdf_skip_patterns_compiled:
                if pattern.search(text) is not None:
                    return True

            # Special handling for complex grant application headers/footers
            return self._is_grant_header_footer(text)

        return False
    
    def _is_grant_header_footer(self, text):
        """Detect complex grant application header/footer patterns with variable elements
        Optimized version using pre-compiled patterns"""
        # Clean the text for analysis
        clean_text = re.sub(r'\s+', ' ', text.strip())

        # Check for very short text that's clearly a header (< 15 chars with OSF or ID)
        if len(clean_text) < 15:
            short_check = re.compile(r'(OSF|ID:|Strona)', re.IGNORECASE)
            if short_check.search(clean_text):
                return True

        # If text is very long (>120 chars), it's probably real content with some metadata appended
        # Only consider it a header if it has strong header indicators
        if len(clean_text) > 120:
            # Only skip if it has OSF AND OPUS AND Strona (very specific header pattern)
            has_osf = bool(self.header_component_patterns['osf'].search(clean_text))
            has_opus = bool(self.header_component_patterns['opus'].search(clean_text))
            has_page = bool(self.header_component_patterns['page'].search(clean_text))

            if has_osf and has_opus and has_page:
                self._log_debug(f"Detected header in long text: '{clean_text[:80]}...'")
                return True
            else:
                # Probably real content with some numbers
                return False

        # Count how many components are present using pre-compiled patterns
        component_matches = 0
        matched_components = []

        for name, pattern in self.header_component_patterns.items():
            if pattern.search(clean_text):
                component_matches += 1
                matched_components.append(name)

        # If we have 3 or more strong components, it's likely a header/footer
        if component_matches >= 3:
            self._log_debug(f"Detected grant header/footer: '{clean_text[:80]}...' (matched: {matched_components})")
            return True

        # Additional check for specific patterns that are clearly headers/footers
        header_indicators = [
            re.compile(r'(OSF|OPUS).{0,50}(Strona|ID).{0,50}\d{4}-\d{2}-\d{2}', re.IGNORECASE),
            re.compile(r'ID:\s*\d+.{0,30}\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}', re.IGNORECASE),
            re.compile(r'OPUS-\d+.{0,50}ID:\s*\d+', re.IGNORECASE),
            re.compile(r'^.{0,100}(OSF|OPUS|ID:|Strona).{0,100}(OSF|OPUS|ID:|Strona).{0,100}$', re.IGNORECASE),
        ]

        for pattern in header_indicators:
            if pattern.search(clean_text):
                self._log_debug(f"Detected header via indicator pattern: '{clean_text}'")
                return True

        return False
    
    def test_skip_patterns(self):
        """Test method to verify skip patterns work correctly"""
        test_cases = [
            "OSF, OPUS-29 Strona 41 ID: 651313, 2025-06-09 11:29:38",
            "OSF OPUS-30 Strona 1 ID: 123456 2024-12-01 09:15:22",
            "OPUS-15 ID: 987654 Strona 25",
            "Regular content that should not be skipped",
            "Some research data about experiments"
        ]
        
        self._log_debug("Testing skip patterns:")
        for test_case in test_cases:
            should_skip = self.should_skip_text(test_case, is_pdf=True)
            self._log_debug(f"'{test_case}' -> Skip: {should_skip}")
        
        return True
    
    def extract_author_name(self, text):
        """Extract author name from the document text"""
        # Common patterns for author name in grant applications
        patterns = [
            r"dr\s+(?:inż\.|hab\.)?\s+([\w\s-]+)",  # Polish academic titles
            r"Principal\s+Investigator[:\s]+([\w\s-]+)",  # English PI designation
            r"Kierownik\s+projektu[:\s]+([\w\s-]+)"  # Polish PI designation
        ]

        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1).strip()

        return None

    def extract_metadata(self, text, doc=None, filename=None):
        """
        Extract metadata from DMP document

        Returns:
            dict with keys:
                - researcher_surname: str
                - researcher_firstname: str
                - competition_name: str (OPUS, PRELUDIUM, etc.)
                - competition_edition: str (number)
                - creation_date: str (DD-MM-YY format)
                - filename_original: str
        """
        metadata = {
            'researcher_surname': None,
            'researcher_firstname': None,
            'competition_name': None,
            'competition_edition': None,
            'creation_date': None,
            'filename_original': filename
        }

        # Extract from filename first (most reliable)
        if filename:
            metadata.update(self._extract_from_filename(filename))

        # Extract from DOCX properties (if available)
        if doc is not None:
            try:
                from docx import Document
                if isinstance(doc, Document):
                    metadata.update(self._extract_from_docx_properties(doc))
            except:
                pass

        # Extract from document text content
        metadata.update(self._extract_from_text_content(text))

        # Format creation date
        if metadata.get('creation_date'):
            metadata['creation_date'] = self._format_date(metadata['creation_date'])

        return metadata

    def _extract_from_filename(self, filename):
        """Extract metadata from filename patterns"""
        metadata = {}

        # Competition patterns in filename
        comp_patterns = [
            r'(OPUS|opus)[\s_-]*(\d+)?',
            r'(PRELUDIUM|preludium|Preludium)[\s_-]*(\d+)?',
            r'(SONATA|sonata|Sonata)[\s_-]*(\d+)?',
            r'(SYMFONIA|symfonia|Symfonia)[\s_-]*(\d+)?',
            r'(MAESTRO|maestro|Maestro)[\s_-]*(\d+)?',
            r'(HARMONIA|harmonia|Harmonia)[\s_-]*(\d+)?',
            r'(MINIATURA|miniatura|Miniatura)[\s_-]*(\d+)?'
        ]

        for pattern in comp_patterns:
            match = re.search(pattern, filename, re.IGNORECASE)
            if match:
                metadata['competition_name'] = match.group(1).upper()
                if match.group(2):
                    metadata['competition_edition'] = match.group(2)
                break

        # Researcher name patterns (Surname-Name or Name_Surname)
        name_patterns = [
            r'([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)[-_]([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)',  # Polish
            r'([A-Z][a-z]+)[-_]([A-Z][a-z]+)'  # English
        ]

        for pattern in name_patterns:
            match = re.search(pattern, filename)
            if match:
                # Assume first is surname, second is firstname
                metadata['researcher_surname'] = match.group(1)
                metadata['researcher_firstname'] = match.group(2)
                break

        return metadata

    def _extract_from_docx_properties(self, doc):
        """Extract metadata from DOCX document properties"""
        metadata = {}

        try:
            # Author from document properties
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                author = doc.core_properties.author.strip()
                # Try to split into firstname and surname
                parts = author.split()
                if len(parts) >= 2:
                    metadata['researcher_firstname'] = parts[0]
                    metadata['researcher_surname'] = parts[-1]
                elif len(parts) == 1:
                    metadata['researcher_surname'] = parts[0]

            # Creation date from document properties
            if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                metadata['creation_date'] = doc.core_properties.created

        except Exception as e:
            self._log_debug(f"Warning: Could not extract DOCX properties: {e}")

        return metadata

    def _extract_from_text_content(self, text):
        """Extract metadata from document text content"""
        metadata = {}

        # Competition patterns in text
        comp_patterns = [
            r'(?:konkurs|competition|grant)[\s:]*(?:NCN[\s:]*)?(OPUS|PRELUDIUM|SONATA|SYMFONIA|MAESTRO|HARMONIA|MINIATURA)[\s-]*(\d+)?',
            r'(OPUS|PRELUDIUM|SONATA|SYMFONIA|MAESTRO|HARMONIA|MINIATURA)[\s-]*(\d+)',
            r'ID:\s*\d+\s*,?\s*(OPUS|PRELUDIUM|SONATA|SYMFONIA|MAESTRO|HARMONIA|MINIATURA)[\s-]*(\d+)?'
        ]

        for pattern in comp_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                if not metadata.get('competition_name'):
                    metadata['competition_name'] = match.group(1).upper()
                if match.lastindex >= 2 and match.group(2) and not metadata.get('competition_edition'):
                    metadata['competition_edition'] = match.group(2)

        # Researcher name patterns
        name_patterns = [
            # Polish patterns
            r'Kierownik\s+projektu[:\s]+(?:dr\.?|prof\.?)?\s*(?:hab\.?|inż\.?)?\s*([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)\s+([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)',
            r'Wykonawca[:\s]+(?:dr\.?|prof\.?)?\s*(?:hab\.?|inż\.?)?\s*([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)\s+([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)',
            # English patterns
            r'Principal\s+Investigator[:\s]+(?:Dr\.?|Prof\.?)?\s*([A-Z][a-z]+)\s+([A-Z][a-z]+)',
            r'Researcher[:\s]+(?:Dr\.?|Prof\.?)?\s*([A-Z][a-z]+)\s+([A-Z][a-z]+)',
            # Name patterns in headers
            r'(?:^|\n)([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]{2,})\s+([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]{2,})\s*(?:\n|$)'
        ]

        for pattern in name_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE)
            for match in matches:
                if not metadata.get('researcher_firstname') and not metadata.get('researcher_surname'):
                    # First match wins - assume "Firstname Surname" order
                    metadata['researcher_firstname'] = match.group(1)
                    metadata['researcher_surname'] = match.group(2)
                    break

        # Date patterns (DD-MM-YYYY, DD.MM.YYYY, YYYY-MM-DD)
        date_patterns = [
            r'(\d{2})[-.](\d{2})[-.](\d{4})',  # DD-MM-YYYY or DD.MM.YYYY
            r'(\d{4})[-.](\d{2})[-.](\d{2})'   # YYYY-MM-DD
        ]

        for pattern in date_patterns:
            match = re.search(pattern, text)
            if match and not metadata.get('creation_date'):
                metadata['creation_date'] = match.group(0)
                break

        return metadata

    def _format_date(self, date_value):
        """Format date to DD-MM-YY"""
        from datetime import datetime

        try:
            # If it's already a datetime object
            if isinstance(date_value, datetime):
                return date_value.strftime('%d-%m-%y')

            # If it's a string, try to parse it
            if isinstance(date_value, str):
                # Try different formats
                for fmt in ['%d-%m-%Y', '%d.%m.%Y', '%Y-%m-%d', '%d/%m/%Y']:
                    try:
                        dt = datetime.strptime(date_value, fmt)
                        return dt.strftime('%d-%m-%y')
                    except:
                        continue

                # If parsing failed, return as-is
                return date_value
        except:
            pass

        return date_value

    def _extract_pdf_with_ocr(self, pdf_path):
        """
        Extract text from scanned PDF using OCR

        Args:
            pdf_path (str): Path to PDF file

        Returns:
            str: Extracted text or None if OCR failed
        """
        if not HAS_OCR:
            self._log_debug("OCR dependencies not available (pdf2image, pytesseract)")
            return None

        try:
            self._log_debug("Attempting OCR extraction for scanned PDF...")

            # Convert PDF to images
            images = convert_from_path(pdf_path, dpi=300)
            self._log_debug(f"Converted PDF to {len(images)} images")

            # Extract text from each page
            extracted_text = []
            for i, image in enumerate(images):
                self._log_debug(f"Processing page {i+1}/{len(images)} with OCR...")
                # Use Polish and English language packs
                text = pytesseract.image_to_string(image, lang='pol+eng')
                extracted_text.append(text)

            full_text = '\n\n'.join(extracted_text)
            self._log_debug(f"OCR extracted {len(full_text)} characters total")

            return full_text

        except Exception as e:
            self._log_debug(f"OCR extraction failed: {str(e)}")
            return None

    def _is_scanned_pdf(self, pdf_path):
        """
        Check if PDF is scanned (no extractable text)

        Args:
            pdf_path (str): Path to PDF file

        Returns:
            bool: True if PDF appears to be scanned, False otherwise
        """
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Check first 3 pages for text
                pages_to_check = min(3, len(pdf_reader.pages))
                total_chars = 0

                for i in range(pages_to_check):
                    text = pdf_reader.pages[i].extract_text()
                    total_chars += len(text.strip())

                # If very little text (< 50 chars per page on average), likely scanned
                avg_chars_per_page = total_chars / pages_to_check if pages_to_check > 0 else 0
                is_scanned = avg_chars_per_page < 50

                if is_scanned:
                    self._log_debug(f"PDF appears to be scanned ({avg_chars_per_page:.0f} chars/page)")
                else:
                    self._log_debug(f"PDF has extractable text ({avg_chars_per_page:.0f} chars/page)")

                return is_scanned

        except Exception as e:
            self._log_debug(f"Error checking if PDF is scanned: {str(e)}")
            return False

    def generate_smart_filename(self, metadata, file_type="DMP", extension=".docx"):
        """
        Generate intelligent filename from metadata

        Format: {type}_{Surname}_{FirstInitial}_{Competition}_{Edition}_{DDMMYY}.{ext}
        Example: DMP_Kowalski_J_OPUS_25_161125.docx

        Args:
            metadata: dict with extracted metadata
            file_type: prefix (default: "DMP")
            extension: file extension (default: ".docx")

        Returns:
            str: Generated filename
        """
        from datetime import datetime

        parts = [file_type]

        # Add researcher name
        if metadata.get('researcher_surname'):
            surname = metadata['researcher_surname']
            # Clean and capitalize
            surname = re.sub(r'[^a-zA-ZąćęłńóśźżĄĆĘŁŃÓŚŹŻ]', '', surname)
            parts.append(surname)

            # Add first initial if available
            if metadata.get('researcher_firstname'):
                firstname = metadata['researcher_firstname']
                first_initial = firstname[0].upper() if firstname else ''
                if first_initial:
                    parts.append(first_initial)

        # Add competition name
        if metadata.get('competition_name'):
            comp = metadata['competition_name'].upper()
            parts.append(comp)

            # Add edition if available
            if metadata.get('competition_edition'):
                parts.append(metadata['competition_edition'])

        # Add date
        if metadata.get('creation_date'):
            # Try to parse and format as DDMMYY
            date_str = metadata['creation_date']
            try:
                # If it's already in DD-MM-YY format
                if isinstance(date_str, str) and '-' in date_str:
                    parts_date = date_str.split('-')
                    if len(parts_date) == 3:
                        # DD-MM-YY format
                        date_formatted = ''.join(parts_date)  # DDMMYY
                        parts.append(date_formatted)
                elif isinstance(date_str, datetime):
                    parts.append(date_str.strftime('%d%m%y'))
            except:
                pass

        # If no date from metadata, use current date
        if not metadata.get('creation_date'):
            parts.append(datetime.now().strftime('%d%m%y'))

        # Join parts with underscore
        filename = '_'.join(parts)

        # Add extension
        if not extension.startswith('.'):
            extension = '.' + extension
        filename += extension

        # Clean filename (remove invalid characters)
        filename = re.sub(r'[\\/*?:"<>|]', '_', filename)

        return filename

    def clean_markup(self, text):
        """Remove common markup from text"""
        # Remove underline markup
        text = re.sub(r'\[([^]]+)\]\{\.underline\}', r'\1', text)
        # Remove bold markup
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        # Remove other common markup
        text = re.sub(r'__([^_]+)__', r'\1', text)
        # Remove mark formatting
        text = re.sub(r'\{\.mark\}', '', text)
        # Clean table delimiters
        text = self.clean_table_delimiters(text)
        return text
    
    def extract_formatted_text(self, paragraph):
        """Extract text with formatting information from a DOCX paragraph"""
        text = paragraph.text.strip()
        if not text:
            return text
            
        is_underlined = False
        is_bold = False
        
        # Check if any runs in the paragraph are underlined or bold
        for run in paragraph.runs:
            if run.underline:
                is_underlined = True
            if run.bold:
                is_bold = True
        
        # Clean the text
        text = self.clean_markup(text)
        
        # Return text with formatting info
        if is_underlined and is_bold:
            return f"UNDERLINED_BOLD:{text}"
        elif is_underlined:
            return f"UNDERLINED:{text}"
        elif is_bold:
            return f"BOLD:{text}"
        else:
            return text
    
    # Key phrases identification removed
    
    def process_paragraph(self, paragraph):
        """Process a paragraph to extract key information"""
        # Clean any markup in the paragraph
        clean_text = self.clean_markup(paragraph)
        tags = []  # Key phrases functionality removed

        # Find lead sentence (first sentence or sentence with most key phrases)
        sentences = re.split(r'(?<=[.!?])\s+', clean_text)

        if not sentences:
            return {
                "text": clean_text,
                "tags": tags,
                "title": None
            }

        # Use first sentence as title if it's reasonably short
        title = sentences[0] if len(sentences[0]) < 100 else None

        return {
            "text": clean_text,
            "tags": tags,
            "title": title
        }

    def clean_extracted_paragraphs(self, paragraphs):
        """
        Remove form delimiters and artifacts from final extracted paragraphs

        This fixes the content pollution issue where form markers like
        "Dół formularza" and "Początek formularza" appear in extracted content.

        Args:
            paragraphs (list): List of extracted paragraph strings

        Returns:
            list: Cleaned paragraphs with form delimiters removed
        """
        # Known form delimiters and artifacts to remove
        skip_phrases = {
            "Dół formularza",
            "Początek formularza",
            "dół formularza",
            "początek formularza",
        }

        cleaned = []
        for para in paragraphs:
            if not para:
                continue

            # Remove formatting prefixes for comparison
            clean_para = para
            for prefix in ["BOLD:", "UNDERLINED:", "UNDERLINED_BOLD:"]:
                if clean_para.startswith(prefix):
                    clean_para = clean_para[len(prefix):].strip()

            # Skip exact matches with form delimiters
            if clean_para.strip() in skip_phrases:
                self._log_debug(f"Removing form delimiter: '{para}'")
                continue

            # Skip if the paragraph is only a form delimiter
            if any(delimiter in clean_para for delimiter in skip_phrases):
                # If it's ONLY the delimiter (possibly with whitespace), skip it
                remaining = clean_para
                for delimiter in skip_phrases:
                    remaining = remaining.replace(delimiter, '')
                if not remaining.strip():
                    self._log_debug(f"Removing paragraph containing only delimiter: '{para}'")
                    continue

            cleaned.append(para)

        return cleaned

    def calculate_extraction_confidence(self, paragraphs, detection_method=None):
        """
        Calculate confidence score for extracted content

        Args:
            paragraphs (list): List of extracted paragraphs
            detection_method (str): Method used for detection (e.g., 'direct_match', 'fuzzy_match')

        Returns:
            float: Confidence score between 0.0 and 1.0
        """
        if not paragraphs:
            return 0.0

        confidence = 0.0

        # Factor 1: Presence of content (40% weight)
        if len(paragraphs) > 0:
            confidence += 0.4

        # Factor 2: Content length indicator (30% weight)
        total_length = sum(len(p) for p in paragraphs)
        avg_length = total_length / len(paragraphs) if paragraphs else 0

        if avg_length > 100:
            confidence += 0.3
        elif avg_length > 50:
            confidence += 0.2
        elif avg_length > 20:
            confidence += 0.1

        # Factor 3: Detection method quality (30% weight)
        method_scores = {
            'direct_section_match': 0.3,
            'direct_subsection_match': 0.3,
            'numbered_section': 0.25,
            'fuzzy_match': 0.15,
            'word_match': 0.1,
            'buffered': 0.05,
            None: 0.1  # Default
        }
        confidence += method_scores.get(detection_method, 0.1)

        # Cap at 1.0
        return min(confidence, 1.0)

    def map_section_to_id(self, section, subsection):
        """Map a section and subsection to a section ID"""
        # Try to find the matching ID
        for section_id, question in self.section_ids.items():
            if question == subsection:
                return section_id
        
        # Fallback: try to infer from section number
        section_num = re.match(r'(\d+)\.', section)
        if section_num:
            num = section_num.group(1)
            for section_id in self.section_ids.keys():
                if section_id.startswith(f"{num}."):
                    return section_id
        
        # If no match found, return a generated ID
        return f"section_{hash(section + subsection) % 10000}"
    
    def detect_section_from_text(self, text, is_pdf=False):
        """Detect section from text content, supporting multiple languages and PDF forms"""
        # CRITICAL FIX: Strip formatting prefixes FIRST before any other processing
        original_text = text

        # Remove formatting prefixes (BOLD:, UNDERLINED:, etc.)
        formatting_prefixes = ["BOLD:", "UNDERLINED:", "UNDERLINED_BOLD:"]
        for prefix in formatting_prefixes:
            if text.startswith(prefix):
                text = text[len(prefix):].strip()
                self._log_debug(f"Stripped formatting prefix '{prefix}' from: '{original_text[:60]}...'")
                break

        # Clean the text of any markup
        text = self.clean_markup(text)
        
        # Enhanced section patterns for PDFs
        if is_pdf:
            # Look for form-style section headers
            form_patterns = [
                r"PLAN\s+ZARZĄDZANIA\s+DANYMI",
                r"DATA\s+MANAGEMENT\s+PLAN",
                r"Opis\s+danych\s+oraz\s+pozyskiwanie",
                r"Dokumentacja\s+i\s+jakość\s+danych",
                r"Przechowywanie\s+i\s+tworzenie\s+kopii",
                r"Wymogi\s+prawne",
                r"Udostępnianie\s+i\s+długotrwałe",
                r"Zadania\s+związane\s+z\s+zarządzaniem"
            ]
            
            for pattern in form_patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    # Map to corresponding English section
                    for polish, english in self.section_mapping.items():
                        if re.search(pattern.replace("\\s+", "\\s*"), polish, re.IGNORECASE):
                            self._log_debug(f"PDF form section detected: '{text}' -> '{english}'")
                            return english
        
        # Try numbered section (e.g. "1. Section title")
        section_match = re.match(r'^\s*(\d+)\.\s*(.*?)$', text)
        if section_match:
            section_num = section_match.group(1)
            section_title = section_match.group(2).strip()
            
            self._log_debug(f"Numbered section detected: {section_num}. {section_title}")
            
            # Try to find matching section in dmp_structure
            for section in self.dmp_structure:
                if section.startswith(f"{section_num}."):
                    self._log_debug(f"Matched to DMP structure: {section}")
                    return section
            
            # Try to find in section mapping if not found directly
            for polish, english in self.section_mapping.items():
                if self._text_similarity(polish.lower(), section_title.lower()) > 0.5:
                    self._log_debug(f"Matched via section mapping: {polish} -> {english}")
                    return english
        
        # Try direct section title matching (formatting prefixes already stripped above)
        for polish, english in self.section_mapping.items():
            if self._text_similarity(polish.lower(), text.lower()) > 0.6:
                self._log_debug(f"Direct title match: {polish} -> {english}")
                return english
        
        return None
    
    def _text_similarity(self, text1, text2):
        """Calculate simple text similarity based on word overlap
        Optimized with caching for repeated calculations"""
        # Create cache key (order matters for similarity, so use tuple)
        cache_key = (text1, text2)

        # Check cache first
        if cache_key in self._similarity_cache:
            return self._similarity_cache[cache_key]

        # Calculate similarity
        words1 = set(word.lower() for word in text1.split() if len(word) > 2)
        words2 = set(word.lower() for word in text2.split() if len(word) > 2)

        if not words1 or not words2:
            similarity = 0.0
        else:
            intersection = words1.intersection(words2)
            union = words1.union(words2)
            similarity = len(intersection) / len(union) if union else 0.0

        # Store in cache (limit cache size to prevent memory issues)
        if len(self._similarity_cache) < 1000:
            self._similarity_cache[cache_key] = similarity

        return similarity
    
    def detect_subsection_from_text(self, text, current_section, is_pdf=False):
        """Detect subsection from text content, supporting multiple languages and partial matching"""
        # CRITICAL FIX: Strip formatting prefixes FIRST
        original_text = text

        # Remove formatting prefixes (BOLD:, UNDERLINED:, etc.)
        formatting_prefixes = ["BOLD:", "UNDERLINED:", "UNDERLINED_BOLD:"]
        for prefix in formatting_prefixes:
            if text.startswith(prefix):
                text = text[len(prefix):].strip()
                break

        # Clean the text of any markup
        text = self.clean_markup(text)

        # For PDFs, check if subsection header is embedded in the middle of text
        # This happens when PDF text extraction concatenates lines
        if is_pdf and len(text) > 100:
            # Try to find subsection patterns in the middle of the text
            for polish, english in self.subsection_mapping.items():
                if current_section and english in self.dmp_structure.get(current_section, []):
                    # Check if the Polish header appears anywhere in the text
                    if polish.lower() in text.lower():
                        self._log_debug(f"Found embedded subsection header: '{polish}' in '{text[:80]}...'")
                        return english

        # Check if this is an underlined text which is often a subsection header
        is_underlined = text.startswith("UNDERLINED:")
        if is_underlined:
            text = text.replace("UNDERLINED:", "").strip()

        # Check if this is bold text
        if text.startswith("BOLD:"):
            text = text.replace("BOLD:", "").strip()
        
        # Normalize text - remove trailing colon and convert to lowercase
        normalized_text = text.lower()
        if normalized_text.endswith(':'):
            normalized_text = normalized_text[:-1].strip()
        
        # Debug output
        self._log_debug(f"Trying to match subsection: '{text}' (normalized: '{normalized_text}', section: {current_section})")
        
        if not current_section:
            self._log_debug("No current section, skipping subsection detection")
            return None
        
        # 1. Check for direct match with English subsections (case insensitive)
        for subsection in self.dmp_structure.get(current_section, []):
            if self._text_similarity(text.lower(), subsection.lower()) > 0.8:
                self._log_debug(f"High similarity English match: '{text}' ~ '{subsection}'")
                return subsection
        
        # 2. Try enhanced Polish subsection matching with similarity scoring
        best_match = None
        best_score = 0.0
        
        for polish, english in self.subsection_mapping.items():
            if current_section and english in self.dmp_structure.get(current_section, []):
                # Calculate similarity scores
                similarity = self._text_similarity(normalized_text, polish.lower())
                
                # Boost score for exact matches or strong partial matches
                if normalized_text == polish.lower():
                    similarity = 1.0
                elif normalized_text in polish.lower() or polish.lower() in normalized_text:
                    similarity = max(similarity, 0.7)
                
                if similarity > best_score and similarity > 0.3:
                    best_score = similarity
                    best_match = english
                    self._log_debug(f"Polish mapping candidate: '{text}' ~ '{polish}' -> '{english}' (score: {similarity:.2f})")
        
        # ENHANCED: Lower threshold from 0.5 to 0.4 for better matching
        if best_match and best_score > 0.4:
            self._log_debug(f"Best Polish match: '{text}' -> '{best_match}' (score: {best_score:.2f})")
            return best_match
        
        # 3. Enhanced word-based matching for formatted text (OPTIMIZED with pre-computed index)
        # IMPORTANT: Only apply to likely headers (short text or formatted), not long paragraphs (likely content)
        if ((is_underlined or original_text.endswith(':') or
            original_text.startswith("BOLD:")) and len(text) < 200) or (20 < len(text) < 150):

            best_word_match = None
            max_match_ratio = 0

            # Pre-compute line words once
            polish_common = {'data', 'będą', 'które', 'oraz', 'jako', 'dane'}
            line_words = set(word.lower() for word in text.split()
                           if len(word) > 3 and word.lower() not in polish_common)

            if not line_words:
                # Skip if no significant words in input text
                pass
            else:
                for subsection in self.dmp_structure.get(current_section, []):
                    # PERFORMANCE: Use pre-computed word index instead of recalculating
                    subsection_words = self._subsection_word_index.get(subsection, set())

                    if not subsection_words:
                        continue

                    # Count matching words
                matching_words = len(subsection_words.intersection(line_words))
                match_ratio = matching_words / max(len(subsection_words), 1)

                # Require at least 3 matching words AND higher threshold for stricter matching
                if matching_words >= 3 and match_ratio > max_match_ratio:
                    max_match_ratio = match_ratio
                    best_word_match = subsection
                    self._log_debug(f"Word match candidate: '{text}' ~ '{subsection}' ({matching_words} words, {match_ratio:.2f} ratio)")

            if best_word_match and max_match_ratio > 0.15:  # Lower threshold for better extraction
                self._log_debug(f"Best word match: '{text}' -> '{best_word_match}' (ratio: {max_match_ratio:.2f})")
                return best_word_match
        
        # 4. PDF-specific subsection detection for form fields
        if is_pdf and len(text) > 10:
            # Look for characteristic Polish question patterns
            question_indicators = [
                r"sposób.*?danych",
                r"jak.*?będą",
                r"jakie.*?dane",
                r"gdzie.*?przechowywane",
                r"kto.*?odpowiedzialny",
                r"środki.*?przeznaczone"
            ]
            
            for indicator in question_indicators:
                if re.search(indicator, normalized_text, re.IGNORECASE):
                    # Try to match with subsections in current section
                    for subsection in self.dmp_structure.get(current_section, []):
                        if self._text_similarity(normalized_text, subsection.lower()) > 0.2:
                            self._log_debug(f"PDF question pattern match: '{text}' -> '{subsection}'")
                            return subsection
        
        self._log_debug(f"No subsection match found for: '{text}' (best score: {best_score:.2f})")
        return None
    
    def _split_embedded_headers(self, line):
        """Split lines that contain embedded subsection headers (common in PDFs)

        Returns a list of split lines, or [line] if no splitting needed.
        """
        # Check if line contains Polish subsection headers
        for polish_header in self.subsection_mapping.keys():
            # Look for the header in the middle of the line
            lower_line = line.lower()
            lower_header = polish_header.lower()

            # Remove colons for matching
            if lower_header.endswith(':'):
                lower_header = lower_header[:-1]

            pos = lower_line.find(lower_header)
            if pos > 10:  # Header is embedded (not at start, some content before)
                # Split the line at the header
                before = line[:pos].strip()
                header_and_after = line[pos:].strip()

                self._log_debug(f"Splitting embedded header: '{line[:60]}...' into 2 parts")
                return [before, header_and_after]

        # No embedded header found
        return [line]

    def extract_pdf_table_content(self, text_lines):
        """Extract and structure table-like content from PDF text lines"""
        table_content = []
        current_table = []
        in_table = False

        for line in text_lines:
            line = line.strip()

            # Skip empty lines
            if not line:
                if in_table and current_table:
                    # End of table, process it
                    processed_table = self._process_table_rows(current_table)
                    table_content.extend(processed_table)
                    current_table = []
                    in_table = False
                continue

            # IMPORTANT: Filter out headers/footers before any processing
            if self.should_skip_text(line, is_pdf=True):
                continue

            # Check if line contains an embedded subsection header (concatenated by PDF extraction)
            # If so, split it into separate lines
            split_lines = self._split_embedded_headers(line)
            if len(split_lines) > 1:
                # Add all split parts instead of the original line
                for split_line in split_lines:
                    if split_line.strip():
                        table_content.append(split_line.strip())
                continue
            
            # Detect table patterns
            is_table_line = (
                # Multiple columns separated by spaces (3+ spaces between words)
                len(re.findall(r'\s{3,}', line)) > 1 or
                # Currency/number patterns (budget tables)
                re.search(r'\d+[,.]?\d*\s+(PLN|EUR|USD|\d+)', line) or
                # Aligned data with consistent spacing
                re.search(r'^[^\s]+\s{5,}[^\s]+\s{5,}[^\s]+', line) or
                # Form fields with underscores or dots
                re.search(r'_{3,}|\.{3,}', line)
            )
            
            if is_table_line:
                if not in_table:
                    in_table = True
                current_table.append(line)
            else:
                if in_table and current_table:
                    # Process completed table
                    processed_table = self._process_table_rows(current_table)
                    table_content.extend(processed_table)
                    current_table = []
                    in_table = False

                # Add non-table content normally (already filtered above)
                table_content.append(line)
        
        # Process any remaining table
        if current_table:
            processed_table = self._process_table_rows(current_table)
            table_content.extend(processed_table)
        
        return table_content
    
    def _process_table_rows(self, table_rows):
        """Process table rows to extract meaningful content"""
        if not table_rows:
            return []
        
        processed = []
        for row in table_rows:
            # Clean up table formatting
            clean_row = re.sub(r'\s{2,}', ' | ', row)  # Replace multiple spaces with separator
            clean_row = re.sub(r'[_\.]{3,}', '[field]', clean_row)  # Replace field markers
            
            # Skip rows that are mostly formatting
            if not re.search(r'[a-zA-ZąćęłńóśźżĄĆĘŁŃÓŚŹŻ]', clean_row):
                continue
                
            processed.append(clean_row)
        
        return processed
    
    def improve_content_assignment(self, all_content, is_pdf=False):
        """Improved content assignment logic with better section/subsection detection"""
        # Initialize structures
        section_content = {}
        tagged_content = {}
        unconnected_text = []
        
        for section in self.dmp_structure:
            section_content[section] = {}
            tagged_content[section] = {}
            for subsection in self.dmp_structure[section]:
                section_content[section][subsection] = []
                tagged_content[section][subsection] = []
        
        current_section = None
        current_subsection = None
        content_buffer = []  # Buffer for content without clear assignment
        
        self._log_debug(f"Processing {len(all_content)} content items...")
        
        for i, content_item in enumerate(all_content):
            if not content_item.strip():
                continue
            
            self._log_debug(f"\n--- Processing item {i+1}: '{content_item[:100]}...'")
            
            # Try to detect section
            detected_section = self.detect_section_from_text(content_item, is_pdf=is_pdf)
            if detected_section:
                # Flush buffer to previous section/subsection
                if current_section and current_subsection and content_buffer:
                    self._log_debug(f"Flushing {len(content_buffer)} buffered items to {current_section} -> {current_subsection}")
                    for buffered_content in content_buffer:
                        self._assign_content_safely(section_content, tagged_content, 
                                                  current_section, current_subsection, buffered_content)
                    content_buffer = []
                
                current_section = detected_section
                current_subsection = None
                self._log_debug(f"Section changed to: {current_section}")
                continue
            
            # Try to detect subsection
            if current_section:
                detected_subsection = self.detect_subsection_from_text(content_item, current_section, is_pdf=is_pdf)

                # Only change subsection if it's different from the current one
                if detected_subsection and detected_subsection != current_subsection:
                    # Flush buffer to previous subsection OR first subsection if none set yet
                    if content_buffer:
                        if current_subsection:
                            # Flush to previous subsection
                            self._log_debug(f"Flushing {len(content_buffer)} buffered items to {current_section} -> {current_subsection}")
                            target_subsection = current_subsection
                        else:
                            # No previous subsection - flush to FIRST subsection of current section
                            target_subsection = self.dmp_structure[current_section][0]
                            self._log_debug(f"Flushing {len(content_buffer)} buffered items to FIRST subsection: {current_section} -> {target_subsection}")

                        for buffered_content in content_buffer:
                            self._assign_content_safely(section_content, tagged_content,
                                                      current_section, target_subsection, buffered_content)
                        content_buffer = []

                    current_subsection = detected_subsection
                    self._log_debug(f"Subsection changed to: {current_subsection}")
                    continue
                elif detected_subsection and detected_subsection == current_subsection:
                    # Same subsection detected - this is probably content, not a header
                    # Don't continue, let it fall through to content assignment
                    self._log_debug(f"Ignoring re-detection of same subsection: {current_subsection}")
                    pass
            
            # Handle regular content
            if current_section and current_subsection:
                # Direct assignment
                self._log_debug(f"Assigning content to {current_section} -> {current_subsection}")
                self._assign_content_safely(section_content, tagged_content, 
                                          current_section, current_subsection, content_item)
            elif current_section:
                # Buffer content until we find a subsection
                self._log_debug(f"Buffering content for section {current_section}")
                content_buffer.append(content_item)
            else:
                # No section identified yet
                self._log_debug("Adding to unconnected text (no section)")
                unconnected_text.append({"text": content_item, "type": "no_section"})
        
        # Handle remaining buffered content
        if content_buffer:
            if current_section and current_subsection:
                self._log_debug(f"Final flush: {len(content_buffer)} items to {current_section} -> {current_subsection}")
                for buffered_content in content_buffer:
                    self._assign_content_safely(section_content, tagged_content, 
                                              current_section, current_subsection, buffered_content)
            else:
                self._log_debug(f"Moving {len(content_buffer)} buffered items to unconnected")
                for buffered_content in content_buffer:
                    unconnected_text.append({"text": buffered_content, "type": "buffered"})
        
        return section_content, tagged_content, unconnected_text
    
    def _assign_content_safely(self, section_content, tagged_content, section, subsection, content):
        """Safely assign content to section/subsection with error handling"""
        try:
            if section not in section_content:
                self._log_debug(f"Warning: Section '{section}' not in structure")
                return
            if subsection not in section_content[section]:
                self._log_debug(f"Warning: Subsection '{subsection}' not in section '{section}'")
                return
            
            section_content[section][subsection].append(content)
            tagged_content[section][subsection].append(self.process_paragraph(content))
            self._log_debug(f"Successfully assigned content (length: {len(content)})")
            
        except Exception as e:
            self._log_debug(f"Error assigning content: {str(e)}")
            # Don't fail completely, just log the error
    
    def process_docx(self, docx_path, output_dir, progress_callback=None):
        """Process a DOCX file and extract DMP content with enhanced table support"""
        try:
            self._log_debug(f"Processing DOCX: {docx_path}")
            self._report_progress(progress_callback, "Starting DOCX processing...", 0)

            # Validate the DOCX file first
            self._report_progress(progress_callback, "Validating DOCX file...", 5)
            is_valid, validation_message = validate_docx_file(docx_path)
            if not is_valid:
                return {
                    "success": False,
                    "message": f"DOCX validation failed: {validation_message}"
                }
            
            # Create a new Word document for output
            output_doc = Document()
            
            # Load the input document
            self._report_progress(progress_callback, "Loading DOCX document...", 10)
            doc = Document(docx_path)

            # Extract text from both paragraphs and tables
            self._report_progress(progress_callback, "Extracting paragraphs and tables...", 15)
            formatted_paragraphs = []

            # Process paragraphs
            for paragraph in doc.paragraphs:
                formatted_text = self.extract_formatted_text(paragraph)
                if formatted_text.strip():  # Only add non-empty paragraphs
                    formatted_paragraphs.append(formatted_text)

            # Process tables
            table_content = self.extract_table_content(doc)
            formatted_paragraphs.extend(table_content)
            self._report_progress(progress_callback, "Content extraction complete", 25)
            
            # Join paragraphs for author detection and searching for markers
            all_text = "\n".join([p.replace("UNDERLINED:", "").replace("BOLD:", "").replace("UNDERLINED_BOLD:", "")
                                 for p in formatted_paragraphs])

            # Extract metadata from document
            filename = os.path.basename(docx_path)
            metadata = self.extract_metadata(all_text, doc=doc, filename=filename)
            self._log_debug(f"Metadata extracted: {metadata}")

            # Extract author name (legacy compatibility)
            author_name = self.extract_author_name(all_text)
            self._log_debug(f"Author detected: {author_name}")
            
            # Find start and end positions in the formatted paragraphs list
            start_idx = -1
            end_idx = len(formatted_paragraphs)
            
            # Find start marker
            for i, para in enumerate(formatted_paragraphs):
                clean_para = para.replace("UNDERLINED:", "").replace("BOLD:", "").replace("UNDERLINED_BOLD:", "").strip()
                for mark in self.start_marks:
                    if mark in clean_para:
                        start_idx = i + 1  # Start from the next paragraph
                        self._log_debug(f"Found start mark '{mark}' at paragraph {i}")
                        break
                if start_idx != -1:
                    break
            
            if start_idx == -1:
                # ENHANCED FALLBACK: Try multiple strategies to find DMP content
                self._log_debug("Standard markers not found, trying fallback detection...")

                # Strategy 1: Look for "1." pattern (numbered section 1)
                for i, para in enumerate(formatted_paragraphs):
                    clean_para = para.replace("UNDERLINED:", "").replace("BOLD:", "").replace("UNDERLINED_BOLD:", "").strip()
                    # More flexible pattern - just look for "1." anywhere near the start
                    if re.search(r'^\s*1\.\s*', clean_para):
                        start_idx = i
                        self._log_debug(f"Fallback Strategy 1: Found '1.' pattern at paragraph {i}")
                        break

                # Strategy 2: Look for Polish section 1 keywords
                if start_idx == -1:
                    for i, para in enumerate(formatted_paragraphs):
                        clean_para = para.replace("UNDERLINED:", "").replace("BOLD:", "").replace("UNDERLINED_BOLD:", "").strip().lower()
                        if ("opis danych" in clean_para and "pozyskiwanie" in clean_para):
                            start_idx = i
                            self._log_debug(f"Fallback Strategy 2: Found Polish section 1 keywords at paragraph {i}")
                            break

                # Strategy 3: Look for English section 1 keywords
                if start_idx == -1:
                    for i, para in enumerate(formatted_paragraphs):
                        clean_para = para.replace("UNDERLINED:", "").replace("BOLD:", "").replace("UNDERLINED_BOLD:", "").strip().lower()
                        if ("data description" in clean_para or "data collection" in clean_para):
                            start_idx = i
                            self._log_debug(f"Fallback Strategy 3: Found English section 1 keywords at paragraph {i}")
                            break

                # Strategy 4: Look for any section header pattern
                if start_idx == -1:
                    for i, para in enumerate(formatted_paragraphs):
                        clean_para = para.replace("UNDERLINED:", "").replace("BOLD:", "").replace("UNDERLINED_BOLD:", "").strip()
                        if self.detect_section_from_text(clean_para) is not None:
                            start_idx = i
                            self._log_debug(f"Fallback Strategy 4: Detected section header at paragraph {i}: '{clean_para[:60]}...'")
                            break
            
            if start_idx == -1:
                return {
                    "success": False,
                    "message": "Could not find the start marker or section 1 in the document."
                }
            
            # Find end marker
            for i in range(start_idx, len(formatted_paragraphs)):
                clean_para = formatted_paragraphs[i].replace("UNDERLINED:", "").replace("BOLD:", "").replace("UNDERLINED_BOLD:", "").strip()
                for mark in self.end_marks:
                    if mark in clean_para:
                        end_idx = i
                        self._log_debug(f"Found end mark '{mark}' at paragraph {i}")
                        break
                if i == end_idx:
                    break
            
            # Extract DMP content (paragraphs between start and end markers)
            dmp_paragraphs = formatted_paragraphs[start_idx:end_idx]
            self._log_debug(f"Extracted {len(dmp_paragraphs)} paragraphs of DMP content")
            
            # Create document
            output_doc.add_heading("DATA MANAGEMENT PLAN", level=0)
            
            # Use improved content assignment logic
            self._log_debug("Using improved content assignment for DOCX processing...")
            
            # Filter out meaningful content (skip formatting markers, headers, etc.)
            meaningful_content = []
            for para in dmp_paragraphs:
                if not para.strip():
                    continue
                    
                clean_para = para
                if "BOLD:" in para or "UNDERLINED:" in para or "UNDERLINED_BOLD:" in para:
                    clean_para = para.replace("BOLD:", "").replace("UNDERLINED:", "").replace("UNDERLINED_BOLD:", "").strip()
                
                # Skip paragraphs that should be skipped
                if self.should_skip_text(clean_para, is_pdf=False):
                    continue
                    
                # Only add substantial content
                if len(clean_para) > 5:
                    meaningful_content.append(para)  # Keep original formatting for detection
            
            # Use improved assignment logic
            self._report_progress(progress_callback, "Analyzing content and assigning to sections...", 40)
            section_content, tagged_content, unconnected_text = self.improve_content_assignment(
                meaningful_content, is_pdf=False
            )
            self._report_progress(progress_callback, "Content assignment complete", 60)

            # Add content to document
            self._report_progress(progress_callback, "Building output document...", 65)
            for section in self.dmp_structure:
                output_doc.add_heading(section, level=1)
                
                for subsection in self.dmp_structure[section]:
                    output_doc.add_heading(subsection, level=2)
                    
                    # Safely get content
                    content = []
                    try:
                        content = section_content[section][subsection]
                    except KeyError:
                        self._log_debug(f"Warning: Missing content for {section} - {subsection}")
                    
                    if content:
                        for text in content:
                            output_doc.add_paragraph(text)
                    else:
                        # Add blank paragraph for empty content
                        output_doc.add_paragraph("")
            
            # Create review structure
            self._report_progress(progress_callback, "Creating review structure with confidence scores...", 75)
            review_structure = {}

            for section in self.dmp_structure:
                for subsection in self.dmp_structure[section]:
                    section_id = self.map_section_to_id(section, subsection)

                    # Safely get paragraphs and tagged paragraphs
                    paragraphs = []
                    tagged_paragraphs = []

                    try:
                        paragraphs = section_content[section][subsection]
                    except KeyError:
                        self._log_debug(f"Warning: Missing paragraphs for {section} - {subsection}")

                    try:
                        tagged_paragraphs = tagged_content[section][subsection]
                    except KeyError:
                        self._log_debug(f"Warning: Missing tagged paragraphs for {section} - {subsection}")

                    # CRITICAL FIX: Clean extracted paragraphs to remove form delimiters
                    paragraphs = self.clean_extracted_paragraphs(paragraphs)

                    # Calculate extraction confidence score
                    confidence = self.calculate_extraction_confidence(
                        paragraphs,
                        detection_method='direct_subsection_match' if paragraphs else None
                    )

                    # Add to review structure with confidence metrics
                    review_structure[section_id] = {
                        "section": section,
                        "question": subsection,
                        "paragraphs": paragraphs,
                        "tagged_paragraphs": tagged_paragraphs,
                        "confidence": round(confidence, 2),
                        "extraction_method": "DOCX processing"
                    }
            
            # Generate smart output filename from metadata
            output_filename = self.generate_smart_filename(metadata, file_type="DMP", extension=".docx")
            self._log_debug(f"Generated smart filename: {output_filename}")

            # Create subdirectories for organized file storage
            dmp_dir = os.path.join(output_dir, 'dmp')
            cache_dir = os.path.join(output_dir, 'cache')
            os.makedirs(dmp_dir, exist_ok=True)
            os.makedirs(cache_dir, exist_ok=True)

            # Save the document to dmp folder
            self._report_progress(progress_callback, "Saving output document...", 85)
            output_path = os.path.join(dmp_dir, output_filename)
            output_doc.save(output_path)

            # Add unconnected text to review structure if present
            if unconnected_text:
                review_structure["_unconnected_text"] = unconnected_text
                self._log_debug(f"Added {len(unconnected_text)} unconnected text items to review structure")

            # Fill empty sections with placeholder text for complete extraction
            empty_count = 0
            for section_id in ['1.1', '1.2', '2.1', '2.2', '3.1', '3.2', '4.1', '4.2', '5.1', '5.2', '5.3', '6.1', '6.2']:
                if section_id in review_structure:
                    paras = review_structure[section_id].get('paragraphs', [])
                    if not paras or len(paras) == 0:
                        # Add placeholder for empty sections
                        placeholder = "Not answered in the source document."
                        review_structure[section_id]['paragraphs'] = [placeholder]
                        review_structure[section_id]['tagged_paragraphs'] = [{
                            'text': placeholder,
                            'tags': [],
                            'title': None
                        }]
                        empty_count += 1

            if empty_count > 0:
                print(f"Added placeholder text to {empty_count} empty section(s)")

            # Add metadata to review structure
            review_structure["_metadata"] = metadata
            self._log_debug(f"Added metadata to review structure")

            # Save review structure as JSON to cache folder
            self._report_progress(progress_callback, "Generating cache file...", 90)
            cache_id = str(uuid.uuid4())
            cache_filename = f"cache_{cache_id}.json"
            cache_path = os.path.join(cache_dir, cache_filename)

            with open(cache_path, 'w', encoding='utf-8') as f:
                json.dump(review_structure, f, ensure_ascii=False, indent=2)

            self._report_progress(progress_callback, "DOCX processing complete!", 100)

            return {
                "success": True,
                "filename": output_filename,
                "path": output_path,
                "cache_id": cache_id,
                "cache_file": cache_filename,
                "message": "DMP successfully extracted from DOCX with table support"
            }
            
        except Exception as e:
            import traceback
            traceback_str = traceback.format_exc()
            self._log_debug(f"Error processing DOCX: {str(e)}")
            print(traceback_str)
            return {
                "success": False,
                "message": f"Error processing DOCX: {str(e)}"
            }
    
    def process_pdf(self, pdf_path, output_dir, progress_callback=None):
        """Process a PDF and extract DMP content"""
        try:
            self._log_debug(f"Processing PDF: {pdf_path}")
            self._report_progress(progress_callback, "Starting PDF processing...", 0)

            # Create a new Word document
            doc = Document()

            # Read the PDF
            self._report_progress(progress_callback, "Opening PDF file...", 5)
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Extract first few pages text for author detection
                first_pages_text = ""
                for i in range(min(3, len(reader.pages))):
                    first_pages_text += reader.pages[i].extract_text() + "\n"

                author_name = self.extract_author_name(first_pages_text)
                self._log_debug(f"Author detected: {author_name}")

                # Find the DMP section
                all_text = ""
                all_pages_text = []

                # Extract text from all pages
                self._report_progress(progress_callback, "Extracting text from PDF pages...", 15)
                total_pages = len(reader.pages)
                for i, page in enumerate(reader.pages):
                    # Report progress for each 10% of pages processed
                    if total_pages > 10 and i % (total_pages // 10) == 0:
                        page_progress = 15 + int((i / total_pages) * 15)  # Progress from 15% to 30%
                        self._report_progress(progress_callback, f"Processing page {i+1}/{total_pages}...", page_progress)

                    page_text = page.extract_text()
                    all_pages_text.append(page_text)
                    all_text += page_text + "\n\n"

                    # Print info about start/end marks found
                    for mark in self.start_marks:
                        if mark in page_text:
                            self._log_debug(f"Found start mark '{mark}' on page {i+1}")

                    for mark in self.end_marks:
                        if mark in page_text:
                            self._log_debug(f"Found end mark '{mark}' on page {i+1}")

                self._report_progress(progress_callback, "Text extraction complete", 30)

                # Check if PDF is scanned and use OCR if needed
                if self._is_scanned_pdf(pdf_path):
                    self._report_progress(progress_callback, "Scanned PDF detected, running OCR...", 35)
                    ocr_text = self._extract_pdf_with_ocr(pdf_path)
                    if ocr_text:
                        self._log_debug("Using OCR extracted text instead of standard extraction")
                        all_text = ocr_text
                    else:
                        return {
                            "success": False,
                            "message": "PDF appears to be scanned but OCR extraction failed. Install pytesseract and pdf2image for OCR support."
                        }

                # Extract metadata from PDF
                filename = os.path.basename(pdf_path)
                metadata = self.extract_metadata(all_text, doc=None, filename=filename)
                self._log_debug(f"Metadata extracted: {metadata}")

                # Find start and end positions
                start_pos = -1
                end_pos = len(all_text)
                
                for mark in self.start_marks:
                    pos = all_text.find(mark)
                    if pos != -1:
                        start_pos = pos + len(mark)
                        self._log_debug(f"Found start mark at position {pos}")
                        break
                
                if start_pos == -1:
                    # ENHANCED FALLBACK for PDFs: Try multiple strategies
                    self._log_debug("Standard PDF markers not found, trying fallback detection...")

                    # Strategy 1: Look for "1." pattern
                    match = re.search(r'1\.\s*[\w\s]+', all_text, re.IGNORECASE)
                    if match:
                        start_pos = match.start()
                        self._log_debug(f"Fallback Strategy 1: Found '1.' pattern at position {start_pos}")

                    # Strategy 2: Look for Polish section 1 keywords
                    if start_pos == -1:
                        match = re.search(r'opis\s+danych.*?pozyskiwanie', all_text, re.IGNORECASE | re.DOTALL)
                        if match:
                            start_pos = match.start()
                            self._log_debug(f"Fallback Strategy 2: Found Polish section 1 keywords at position {start_pos}")

                    # Strategy 3: Look for English section 1 keywords
                    if start_pos == -1:
                        match = re.search(r'data\s+description|data\s+collection', all_text, re.IGNORECASE)
                        if match:
                            start_pos = match.start()
                            self._log_debug(f"Fallback Strategy 3: Found English section 1 keywords at position {start_pos}")

                    # If still not found, fail
                    if start_pos == -1:
                        return {
                            "success": False,
                            "message": "Could not find the start marker or section 1 in the document."
                        }
                
                for mark in self.end_marks:
                    pos = all_text.find(mark, start_pos)
                    if pos != -1 and pos < end_pos:
                        end_pos = pos
                        self._log_debug(f"Found end mark at position {pos}")
                
                # Extract DMP content
                dmp_text = all_text[start_pos:end_pos]
                self._log_debug(f"Extracted {len(dmp_text)} characters of DMP content")
                
                # Create document
                doc.add_heading("DATA MANAGEMENT PLAN", level=0)
                
                # Process the content by section
                section_content = {}
                tagged_content = {}
                
                for section in self.dmp_structure:
                    section_content[section] = {}
                    tagged_content[section] = {}
                    for subsection in self.dmp_structure[section]:
                        section_content[section][subsection] = []
                        tagged_content[section][subsection] = []
                
                # Split the content into lines and use improved table extraction
                self._report_progress(progress_callback, "Processing PDF content structure...", 45)
                lines = dmp_text.split("\n")
                self._log_debug(f"Extracted {len(lines)} lines from PDF")

                # Use PDF table extraction to better structure content
                structured_content = self.extract_pdf_table_content(lines)
                self._log_debug(f"After table processing: {len(structured_content)} content items")

                # Use improved content assignment logic
                self._report_progress(progress_callback, "Analyzing content and assigning to sections...", 50)
                section_content, tagged_content, unconnected_text = self.improve_content_assignment(
                    structured_content, is_pdf=True
                )
                self._report_progress(progress_callback, "Content assignment complete", 65)

                # Add content to document
                self._report_progress(progress_callback, "Building output document...", 70)
                for section in self.dmp_structure:
                    doc.add_heading(section, level=1)
                    
                    for subsection in self.dmp_structure[section]:
                        doc.add_heading(subsection, level=2)
                        
                        # Safely get content from the section/subsection
                        content = []
                        try:
                            content = section_content[section][subsection]
                        except KeyError:
                            self._log_debug(f"Warning: Missing content for {section} - {subsection}")
                        
                        if content:
                            for text in content:
                                doc.add_paragraph(text)
                        else:
                            # Add blank paragraph for empty content
                            doc.add_paragraph("")
                
         # Create a structured representation for the review interface
                self._report_progress(progress_callback, "Creating review structure with confidence scores...", 75)
                review_structure = {}

                for section in self.dmp_structure:
                    for subsection in self.dmp_structure[section]:
                        section_id = self.map_section_to_id(section, subsection)

                        # Safely get paragraphs
                        paragraphs = []
                        tagged_paragraphs = []

                        try:
                            paragraphs = section_content[section][subsection]
                        except KeyError:
                            self._log_debug(f"Warning: Missing paragraphs for {section} - {subsection}")

                        try:
                            tagged_paragraphs = tagged_content[section][subsection]
                        except KeyError:
                            self._log_debug(f"Warning: Missing tagged paragraphs for {section} - {subsection}")

                        # CRITICAL FIX: Clean extracted paragraphs to remove form delimiters
                        paragraphs = self.clean_extracted_paragraphs(paragraphs)

                        # Calculate extraction confidence score
                        confidence = self.calculate_extraction_confidence(
                            paragraphs,
                            detection_method='direct_subsection_match' if paragraphs else None
                        )

                        # Add to review structure with confidence metrics
                        review_structure[section_id] = {
                            "section": section,
                            "question": subsection,
                            "paragraphs": paragraphs,
                            "tagged_paragraphs": tagged_paragraphs,
                            "confidence": round(confidence, 2),
                            "extraction_method": "PDF processing"
                        }
                
                # Generate smart output filename from metadata
                output_filename = self.generate_smart_filename(metadata, file_type="DMP", extension=".docx")
                self._log_debug(f"Generated smart filename: {output_filename}")

                # Create subdirectories for organized file storage
                dmp_dir = os.path.join(output_dir, 'dmp')
                cache_dir = os.path.join(output_dir, 'cache')
                os.makedirs(dmp_dir, exist_ok=True)
                os.makedirs(cache_dir, exist_ok=True)

                # Save the document to dmp folder
                self._report_progress(progress_callback, "Saving output document...", 85)
                output_path = os.path.join(dmp_dir, output_filename)
                doc.save(output_path)

                # Add unconnected text to review structure if present
                if unconnected_text:
                    review_structure["_unconnected_text"] = unconnected_text
                    self._log_debug(f"Added {len(unconnected_text)} unconnected text items to PDF review structure")

                # Fill empty sections with placeholder text for complete extraction
                empty_count = 0
                for section_id in ['1.1', '1.2', '2.1', '2.2', '3.1', '3.2', '4.1', '4.2', '5.1', '5.2', '5.3', '6.1', '6.2']:
                    if section_id in review_structure:
                        paras = review_structure[section_id].get('paragraphs', [])
                        if not paras:
                            # Add placeholder for empty sections
                            placeholder = "Not answered in the source document."
                            review_structure[section_id]['paragraphs'] = [placeholder]
                            review_structure[section_id]['tagged_paragraphs'] = [{
                                'text': placeholder,
                                'tags': [],
                                'title': None
                            }]
                            empty_count += 1

                if empty_count > 0:
                    print(f"Added placeholder text to {empty_count} empty PDF section(s)")

                # Add metadata to review structure
                review_structure["_metadata"] = metadata
                self._log_debug(f"Added metadata to PDF review structure")

                # Save review structure as JSON to cache folder
                self._report_progress(progress_callback, "Generating cache file...", 90)
                cache_id = str(uuid.uuid4())
                cache_filename = f"cache_{cache_id}.json"
                cache_path = os.path.join(cache_dir, cache_filename)

                with open(cache_path, 'w', encoding='utf-8') as f:
                    json.dump(review_structure, f, ensure_ascii=False, indent=2)

                self._report_progress(progress_callback, "PDF processing complete!", 100)

                return {
                    "success": True,
                    "filename": output_filename,
                    "path": output_path,
                    "cache_id": cache_id,
                    "cache_file": cache_filename,
                    "message": "DMP successfully extracted from PDF"
                }
                
        except Exception as e:
            import traceback
            traceback_str = traceback.format_exc()
            self._log_debug(f"Error processing PDF: {str(e)}")
            print(traceback_str)
            return {
                "success": False,
                "message": f"Error processing PDF: {str(e)}"
            }
    
    def get_section_ids(self):
        """Return mapping of section IDs to questions"""
        return self.section_ids
    
    # Key phrases functionality removed