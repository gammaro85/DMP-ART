# app.py - Enhanced Flask application with About page and AI Module
from flask import Flask, render_template, request, send_file, jsonify, redirect, url_for, Response, stream_with_context, send_from_directory
import os
import json
import time
import threading
import shutil
import zipfile
import uuid
import re
from datetime import datetime
from werkzeug.utils import secure_filename
from utils.extractor_v4 import DMPExtractor, SkipTermsManager
from utils.ai_module import AIReviewAssistant
# Comments are now managed through JSON files in config/ directory

# Global progress state for real-time SSE updates
progress_state = {}
progress_lock = threading.Lock()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['CACHE_FOLDER'] = 'outputs/cache'
app.config['DMP_FOLDER'] = 'outputs/dmp'
app.config['REVIEWS_FOLDER'] = 'outputs/reviews'
app.config['ARCHIVES_FOLDER'] = 'outputs/archives'
app.config['SESSIONS_FOLDER'] = 'outputs/sessions'
app.config['ACTIVE_SESSIONS_FOLDER'] = 'outputs/sessions/active'
app.config['SESSION_ARCHIVE_FOLDER'] = 'outputs/sessions/archive'
app.config['FEEDBACK_TEMPLATES_PATH'] = os.path.join('config', 'feedback_templates.json')
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx'}
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB default; overridden by config/settings.json

SECTION_IDS = ['1.1', '1.2', '2.1', '2.2', '3.1', '3.2',
               '4.1', '4.2', '5.1', '5.2', '5.3', '5.4', '6.1', '6.2']

CATEGORY_SYSTEM_FILES = {
    'dmp_structure.json', 'quick_comments.json', 'category_comments.json',
    'ai_config.json', 'knowledge_base.json', 'extraction_rules.json',
    'dmp_anchors.json', 'extraction_skip_terms.json', 'settings.json'
}

CATEGORY_VARIANT_SUFFIXES = ('_pl_stare', '_en_stare', '_pl', '_en')

# Load persisted general settings
_GENERAL_SETTINGS_PATH = os.path.join('config', 'settings.json')
# Extractor selection — currently only 'v4' is available.
# To add a new extractor: import it above, add its key here, and handle it in the upload route.
EXTRACTOR_NAME = 'v4'  # active extractor identifier
try:
    if os.path.exists(_GENERAL_SETTINGS_PATH):
        with open(_GENERAL_SETTINGS_PATH, 'r', encoding='utf-8') as _f:
            _saved = json.load(_f)
        if 'max_upload_mb' in _saved:
            app.config['MAX_CONTENT_LENGTH'] = int(_saved['max_upload_mb']) * 1024 * 1024
        if 'extractor_name' in _saved:
            EXTRACTOR_NAME = _saved['extractor_name']
except Exception:
    pass  # Fall back to default if file is corrupt

# Create necessary directories
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)
os.makedirs(app.config['CACHE_FOLDER'], exist_ok=True)
os.makedirs(app.config['DMP_FOLDER'], exist_ok=True)
os.makedirs(app.config['REVIEWS_FOLDER'], exist_ok=True)
os.makedirs(app.config['ARCHIVES_FOLDER'], exist_ok=True)
os.makedirs(app.config['ACTIVE_SESSIONS_FOLDER'], exist_ok=True)
os.makedirs(app.config['SESSION_ARCHIVE_FOLDER'], exist_ok=True)

# Initialize AI Module
ai_assistant = AIReviewAssistant()

# Default templates for feedback (loaded once at startup)
DEFAULT_FEEDBACK_TEMPLATES = {
    "1.1": "The data collection methodology needs more detail. Please specify the exact sources and collection methods.",
    "1.2": "Please provide more specific information about data formats and expected volumes.",
    "2.1": "The metadata standards should be clearly specified. Consider using established standards in your field.",
    "2.2": "More rigorous quality control measures should be implemented. Consider validation procedures.",
    "3.1": "Your backup strategy needs improvement. Consider redundant storage solutions.",
    "3.2": "The security measures seem inadequate. Please detail encryption methods and access controls.",
    "4.1": "Your GDPR compliance plan needs more detail. Specify consent procedures and data minimization strategies.",
    "4.2": "Intellectual property considerations are unclear. Please specify licensing arrangements.",
    "5.1": "Data sharing timeline is vague. Please provide specific milestones for data publication.",
    "5.2": "Long-term preservation strategy needs more detail. Specify repository selection criteria.",
    "5.3": "Software documentation is insufficient. Please list all required tools and versions.",
    "5.4": "Your DOI implementation plan lacks detail. Specify exactly how and when identifiers will be assigned.",
    "6.1": "Data stewardship responsibilities are unclear. Please designate specific roles and responsibilities.",
    "6.2": "Resource allocation for data management seems insufficient. Consider budgeting for dedicated staff time."
}


def _load_json_file(file_path, default=None):
    if not os.path.exists(file_path):
        return {} if default is None else default

    with open(file_path, 'r', encoding='utf-8') as file_handle:
        return json.load(file_handle)


def _write_json_file(file_path, data):
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    with open(file_path, 'w', encoding='utf-8') as file_handle:
        json.dump(data, file_handle, ensure_ascii=False, indent=2)


def _get_feedback_templates_path():
    return app.config.get('FEEDBACK_TEMPLATES_PATH', os.path.join('config', 'feedback_templates.json'))


def _get_cache_path(cache_id):
    return os.path.join(app.config['CACHE_FOLDER'], f"cache_{cache_id}.json")


def _get_active_session_paths(cache_id):
    session_dir = os.path.join(app.config['ACTIVE_SESSIONS_FOLDER'], cache_id)

    return {
        'session_dir': session_dir,
        'dmp_path': os.path.join(session_dir, 'dmp_plan.json'),
        'feedback_path': os.path.join(session_dir, 'feedback.json'),
        'metadata_path': os.path.join(session_dir, 'metadata.json'),
        'review_export_path': os.path.join(session_dir, 'review_export.json')
    }


def _find_session_source_upload(session_dir):
    if not os.path.isdir(session_dir):
        return None, None

    for entry in os.listdir(session_dir):
        if not entry.startswith('source_upload'):
            continue

        source_path = os.path.join(session_dir, entry)
        if os.path.isfile(source_path):
            return source_path, entry

    return None, None


def _store_session_source_upload(session_dir, source_file_path, original_filename=''):
    if not source_file_path or not os.path.exists(source_file_path):
        return None

    _, previous_name = _find_session_source_upload(session_dir)
    if previous_name:
        previous_path = os.path.join(session_dir, previous_name)
        if os.path.abspath(previous_path) != os.path.abspath(source_file_path):
            os.remove(previous_path)

    extension = os.path.splitext(original_filename or source_file_path)[1].lower()
    stored_filename = f"source_upload{extension}" if extension else 'source_upload'
    stored_path = os.path.join(session_dir, stored_filename)

    if os.path.abspath(source_file_path) != os.path.abspath(stored_path):
        shutil.copy2(source_file_path, stored_path)

    return stored_filename


def _load_cache_data(cache_id):
    cache_path = _get_cache_path(cache_id)
    if not os.path.exists(cache_path):
        raise FileNotFoundError('Cache file not found')

    return _load_json_file(cache_path), cache_path


def _build_dmp_plan(cache_id, cache_data):
    dmp_plan = {
        'cache_id': cache_id,
        'sections': {}
    }

    for section_id in SECTION_IDS:
        if section_id in cache_data:
            section_info = cache_data[section_id]
            dmp_plan['sections'][section_id] = {
                'section': section_info.get('section', ''),
                'question': section_info.get('question', ''),
                'content': '\n'.join(section_info.get('paragraphs', [])),
                'tagged_paragraphs': section_info.get('tagged_paragraphs', [])
            }

    return dmp_plan


def _build_session_metadata(cache_id, extracted_metadata, session_created_at, status='active', existing_session_name=''):
    return {
        'cache_id': cache_id,
        'status': status,
        'session_created_at': session_created_at,
        'last_updated': datetime.now().isoformat(),
        'session_name': existing_session_name,
        'researcher_surname': extracted_metadata.get('researcher_surname', ''),
        'researcher_firstname': extracted_metadata.get('researcher_firstname', ''),
        'competition_name': extracted_metadata.get('competition_name', ''),
        'competition_edition': extracted_metadata.get('competition_edition', ''),
        'creation_date': extracted_metadata.get('creation_date', ''),
        'filename_original': extracted_metadata.get('filename_original', ''),
        'source_cache_file': f"cache_{cache_id}.json",
        'dmp_file': 'dmp_plan.json',
        'feedback_file': 'feedback.json',
        'review_export_file': 'review_export.json'
    }


def _ensure_active_session(cache_id, feedback_data=None, compiled_feedback=None, source_file_path=None, original_filename=''):
    cache_data, cache_path = _load_cache_data(cache_id)
    paths = _get_active_session_paths(cache_id)

    os.makedirs(paths['session_dir'], exist_ok=True)

    dmp_plan = _build_dmp_plan(cache_id, cache_data)
    _write_json_file(paths['dmp_path'], dmp_plan)

    existing_metadata = _load_json_file(paths['metadata_path'], {})
    session_created_at = existing_metadata.get('session_created_at', datetime.now().isoformat())
    existing_session_name = existing_metadata.get('session_name', '')
    extracted_metadata = cache_data.get('_metadata', {})
    metadata_json = _build_session_metadata(cache_id, extracted_metadata, session_created_at, existing_session_name=existing_session_name)
    metadata_json['session_folder'] = paths['session_dir']
    metadata_json['source_cache_path'] = cache_path

    existing_source_path, existing_source_file = _find_session_source_upload(paths['session_dir'])
    if existing_source_path and existing_source_file:
        metadata_json['source_upload_file'] = existing_source_file
        metadata_json['source_upload_name'] = existing_metadata.get('source_upload_name') or existing_metadata.get('filename_original', '')

    stored_source_file = _store_session_source_upload(paths['session_dir'], source_file_path, original_filename)
    if stored_source_file:
        metadata_json['source_upload_file'] = stored_source_file
        metadata_json['source_upload_name'] = original_filename or os.path.basename(source_file_path)

    _write_json_file(paths['metadata_path'], metadata_json)

    feedback_json = _load_json_file(paths['feedback_path'], {
        'cache_id': cache_id,
        'sections': {},
        'compiled_feedback': '',
        'last_saved': None
    })

    if feedback_data is not None:
        feedback_json['sections'] = feedback_data
        feedback_json['last_saved'] = datetime.now().isoformat()

    if compiled_feedback is not None:
        feedback_json['compiled_feedback'] = compiled_feedback
        feedback_json['last_saved'] = datetime.now().isoformat()

    feedback_json['cache_id'] = cache_id
    _write_json_file(paths['feedback_path'], feedback_json)

    return {
        'paths': paths,
        'cache_data': cache_data,
        'dmp_plan': dmp_plan,
        'feedback': feedback_json,
        'metadata': metadata_json
    }


def _iter_archive_roots():
    roots = [app.config['SESSION_ARCHIVE_FOLDER'], app.config['ARCHIVES_FOLDER']]
    seen = set()

    for root in roots:
        normalized = os.path.normpath(root)
        if normalized in seen:
            continue
        seen.add(normalized)
        yield root


def _find_archive_path(archive_id):
    for root in _iter_archive_roots():
        archive_path = os.path.join(root, archive_id)
        if os.path.exists(archive_path):
            return archive_path
    return None

def load_dmp_templates():
    """
    Load DMP templates from dmp_structure.json (single source of truth).
    Combines questions from JSON with default feedback templates.
    """
    templates = {}

    try:
        with open('config/dmp_structure.json', 'r', encoding='utf-8') as f:
            dmp_structure = json.load(f)

        for section in dmp_structure.get('structure', []):
            for subsection in section.get('subsections', []):
                section_id = subsection.get('id')
                question = subsection.get('question', '')

                templates[section_id] = {
                    'question': question,
                    'template': DEFAULT_FEEDBACK_TEMPLATES.get(section_id, '')
                }

    except Exception as e:
        print(f"Warning: Could not load dmp_structure.json: {e}")
        # Fallback to defaults if file not found
        for section_id, template in DEFAULT_FEEDBACK_TEMPLATES.items():
            templates[section_id] = {
                'question': f"Section {section_id}",
                'template': template
            }

    template_overrides = _load_json_file(_get_feedback_templates_path(), {})
    if isinstance(template_overrides, dict):
        for section_id, template_text in template_overrides.items():
            if section_id in templates and isinstance(template_text, str):
                templates[section_id]['template'] = template_text

    return templates

# Load DMP templates at startup (single source of truth from dmp_structure.json)
DMP_TEMPLATES = load_dmp_templates()

# Template categories will be managed through template editor


def split_category_variant(category_name):
    """Return category base name and optional language suffix."""
    for suffix in CATEGORY_VARIANT_SUFFIXES:
        if category_name.endswith(suffix):
            return category_name[:-len(suffix)], suffix
    return category_name, None


def collect_category_base_names(config_dir, skip_files=None):
    """Collect unique category base names from config files."""
    base_categories = set()
    skip_files = set(skip_files or ())

    if not os.path.exists(config_dir):
        return base_categories

    for filename in os.listdir(config_dir):
        if not filename.endswith('.json'):
            continue
        if filename in skip_files:
            continue
        if 'backup' in filename.lower():
            continue

        file_base = filename[:-5]
        base_name, _ = split_category_variant(file_base)
        if base_name:
            base_categories.add(base_name)

    return base_categories


def resolve_category_file(config_dir, category_base, lang='pl'):
    """Resolve the preferred category file for the requested language."""
    candidates = [
        f'{category_base}_{lang}.json',
        f'{category_base}.json',
        f'{category_base}_{lang}_stare.json',
    ]

    for candidate in candidates:
        file_path = os.path.join(config_dir, candidate)
        if os.path.exists(file_path):
            return candidate, file_path

    return None, None

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def _resolve_generated_file_path(filename):
    safe_filename = secure_filename(filename or '')
    if not safe_filename:
        return None

    candidates = [
        os.path.join(app.config['OUTPUT_FOLDER'], safe_filename),
        os.path.join(app.config['DMP_FOLDER'], safe_filename),
    ]

    for candidate in candidates:
        if os.path.exists(candidate):
            return candidate

    return None

def validate_docx_file(file_path):
    """Enhanced DOCX file validation"""
    try:
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        if not file_path.lower().endswith('.docx'):
            return False, "File is not a DOCX file"
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return False, "File is empty"
        
        if file_size > 16 * 1024 * 1024:  # 16MB limit
            return False, "File is too large (max 16MB)"
        
        # Check if it's a valid ZIP file (DOCX is ZIP-based)
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                file_list = zip_file.namelist()
                required_files = ['word/document.xml', '[Content_Types].xml']
                
                for required_file in required_files:
                    if required_file not in file_list:
                        return False, f"Invalid DOCX structure: missing {required_file}"
                        
        except zipfile.BadZipFile:
            return False, "File is not a valid ZIP archive"
        except Exception as e:
            return False, f"ZIP validation error: {str(e)}"
        
        # Try to load with python-docx
        try:
            from docx import Document
            doc = Document(file_path)
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            # Check for minimum content
            has_content = any(p.text.strip() for p in doc.paragraphs) or table_count > 0
            
            if not has_content:
                return False, "Document appears to be empty or contains no readable content"
                
        except Exception as e:
            return False, f"Document processing error: {str(e)}"
        
        return True, "File is valid"
        
    except Exception as e:
        return False, f"Validation error: {str(e)}"

def validate_pdf_file(file_path):
    """Enhanced PDF file validation"""
    try:
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        if not file_path.lower().endswith('.pdf'):
            return False, "File is not a PDF file"
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return False, "File is empty"
        
        if file_size > 16 * 1024 * 1024:  # 16MB limit
            return False, "File is too large (max 16MB)"
        
        # Try to read with PyPDF2
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                page_count = len(reader.pages)
                
                if page_count == 0:
                    return False, "PDF contains no pages"
                
                # Try to extract text from first page
                try:
                    first_page_text = reader.pages[0].extract_text()
                    if not first_page_text.strip():
                        return False, "PDF appears to contain no readable text"
                except Exception:
                    return False, "Cannot extract text from PDF"
                    
        except Exception as e:
            return False, f"PDF processing error: {str(e)}"
        
        return True, "File is valid"
        
    except Exception as e:
        return False, f"Validation error: {str(e)}"

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/documentation')
def documentation():
    """Documentation page with features and technical information"""
    return render_template('documentation.html')

@app.route('/test_categories_page')
def test_categories_page():
    """Test page for debugging category loading"""
    return render_template('test_categories.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({
            'success': False,
            'message': 'No file part'
        })

    file = request.files['file']

    if not file or not file.filename:
        return jsonify({
            'success': False,
            'message': 'No selected file'
        })

    # Generate unique session ID for progress tracking
    session_id = str(uuid.uuid4())

    # Initialize progress state
    with progress_lock:
        progress_state[session_id] = {
            'message': 'Starting upload...',
            'progress': 0,
            'status': 'processing'
        }

    file_path = None  # Ensure file_path is always defined
    if file and allowed_file(file.filename):
        try:
            filename = secure_filename(file.filename or "")
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            # Update progress: File saved
            with progress_lock:
                progress_state[session_id].update({
                    'message': 'File uploaded, validating...',
                    'progress': 5
                })

            # Enhanced file validation based on type
            if filename.lower().endswith('.docx'):
                is_valid, validation_message = validate_docx_file(file_path)
                if not is_valid:
                    try:
                        os.remove(file_path)
                    except:
                        pass

                    # Mark progress as error
                    with progress_lock:
                        progress_state[session_id].update({
                            'message': f'Validation failed: {validation_message}',
                            'progress': 0,
                            'status': 'error'
                        })

                    return jsonify({
                        'success': False,
                        'message': f'DOCX validation failed: {validation_message}',
                        'session_id': session_id
                    })
            elif filename.lower().endswith('.pdf'):
                is_valid, validation_message = validate_pdf_file(file_path)
                if not is_valid:
                    try:
                        os.remove(file_path)
                    except:
                        pass

                    # Mark progress as error
                    with progress_lock:
                        progress_state[session_id].update({
                            'message': f'Validation failed: {validation_message}',
                            'progress': 0,
                            'status': 'error'
                        })

                    return jsonify({
                        'success': False,
                        'message': f'PDF validation failed: {validation_message}',
                        'session_id': session_id
                    })

            # Process the file with real-time progress callback
            def progress_callback(message, progress):
                """Update global progress state for SSE streaming"""
                timestamp = datetime.now().strftime('%H:%M:%S')
                print(f"[{timestamp}] Processing: {progress}% - {message}")

                with progress_lock:
                    if session_id in progress_state:
                        progress_state[session_id].update({
                            'message': message,
                            'progress': progress,
                            'status': 'processing'
                        })

            # Extractor selection — extend EXTRACTOR_NAME handling here when adding new extractors
            extractor = DMPExtractor()

            result = extractor.process_file(
                file_path,
                app.config['OUTPUT_FOLDER'],
                progress_callback=progress_callback
            )

            if result['success']:
                # Mark progress as complete
                cache_id = result.get('cache_id', '')
                if cache_id:
                    try:
                        _ensure_active_session(cache_id, source_file_path=file_path, original_filename=filename)
                    except Exception as e:
                        print(f"Warning: Could not initialize active session history: {str(e)}")
                redirect_url = url_for('review_dmp', filename=result['filename'], cache_id=cache_id)

                # Clean up the uploaded file after preserving the original in the session bundle.
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Warning: Could not remove uploaded file: {str(e)}")

                with progress_lock:
                    progress_state[session_id].update({
                        'message': 'Processing complete!',
                        'progress': 100,
                        'status': 'complete',
                        'redirect': redirect_url
                    })

                return jsonify({
                    'success': True,
                    'redirect': redirect_url,
                    'message': result.get('message', 'File processed successfully'),
                    'session_id': session_id
                })
            else:
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Warning: Could not remove uploaded file: {str(e)}")

                # Mark progress as error
                with progress_lock:
                    progress_state[session_id].update({
                        'message': result.get('message', 'Processing failed'),
                        'progress': 0,
                        'status': 'error'
                    })

                return jsonify({
                    **result,
                    'session_id': session_id
                })

        except Exception as e:
            import traceback
            traceback_str = traceback.format_exc()
            print(f"Error processing file: {str(e)}")
            print(traceback_str)

            # Mark progress as error
            with progress_lock:
                if session_id in progress_state:
                    progress_state[session_id].update({
                        'message': f'Error: {str(e)}',
                        'progress': 0,
                        'status': 'error'
                    })

            # Clean up uploaded file in case of error
            try:
                if file_path is not None:
                    os.remove(file_path)
            except Exception:
                pass

            return jsonify({
                'success': False,
                'message': f'Error processing file: {str(e)}',
                'session_id': session_id
            })

    # Invalid file format
    with progress_lock:
        progress_state[session_id].update({
            'message': 'Invalid file format',
            'progress': 0,
            'status': 'error'
        })

    return jsonify({
        'success': False,
        'message': 'Invalid file format. Only PDF and DOCX files are allowed.',
        'session_id': session_id
    })

@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_path = _resolve_generated_file_path(filename)
        if not file_path:
            return "File not found", 404
        
        return send_file(
            file_path,
            as_attachment=True
        )
    except Exception as e:
        return f"Error downloading file: {str(e)}", 500

@app.route('/download-original/<cache_id>')
def download_original_file(cache_id):
    try:
        session_dir = os.path.join(app.config['ACTIVE_SESSIONS_FOLDER'], cache_id)
        source_path, stored_name = _find_session_source_upload(session_dir)
        if not source_path:
            return "Original source file not found", 404

        metadata = _load_json_file(os.path.join(session_dir, 'metadata.json'), {})
        download_name = metadata.get('source_upload_name') or metadata.get('filename_original') or stored_name

        return send_file(
            source_path,
            as_attachment=True,
            download_name=download_name
        )
    except Exception as e:
        return f"Error downloading original file: {str(e)}", 500

@app.route('/review/<filename>')
def review_dmp(filename):
    # File existence is no longer required — all content comes from the cache.
    # Keep the lookup only so the download link works if a DMP DOCX was created.
    file_path = _resolve_generated_file_path(filename)
    # No file is fine — cache is the source of truth
    
    cache_id = request.args.get('cache_id', '')
    
    extracted_content = {}
    extraction_info = {}
    unconnected_text = []
    has_original_source = False
    
    if cache_id:
        cache_path = os.path.join(app.config['CACHE_FOLDER'], f"cache_{cache_id}.json")
        source_path, _ = _find_session_source_upload(os.path.join(app.config['ACTIVE_SESSIONS_FOLDER'], cache_id))
        has_original_source = source_path is not None
        if os.path.exists(cache_path):
            try:
                with open(cache_path, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
                    
                    if cache_data is not None and isinstance(cache_data, dict):
                        if "_unconnected_text" in cache_data:
                            unconnected_text = cache_data["_unconnected_text"]
                            del cache_data["_unconnected_text"]
                        extracted_content = cache_data
                        extraction_info = {
                            'total_sections': len([k for k in extracted_content.keys() if k.startswith(('1.', '2.', '3.', '4.', '5.', '6.'))]),
                            'sections_with_content': len([k for k, v in extracted_content.items() if v.get('paragraphs') and len(v['paragraphs']) > 0]),
                            'extraction_method': 'Enhanced DOCX processing with table support' if filename.lower().endswith('.docx') else 'PDF text extraction'
                        }
            except Exception as e:
                print(f"Error loading extracted content: {str(e)}")
    
    return render_template('review.html', 
                           filename=filename,
                           templates=DMP_TEMPLATES,
                           extracted_content=extracted_content,
                           extraction_info=extraction_info,
                           unconnected_text=unconnected_text,
                           cache_id=cache_id,
                           has_original_source=has_original_source)

@app.route('/save_templates', methods=['POST'])
def save_templates():
    try:
        data = request.json or {}
        global DMP_TEMPLATES

        if not isinstance(data, dict):
            return jsonify({
                'success': False,
                'message': 'Invalid template payload'
            }), 400
        
        # Update the templates with the new data
        updated_count = 0
        for key, value in data.items():
            if key in DMP_TEMPLATES and isinstance(value, str):
                DMP_TEMPLATES[key]['template'] = value
                updated_count += 1

        _write_json_file(
            _get_feedback_templates_path(),
            {section_id: template_data.get('template', '') for section_id, template_data in DMP_TEMPLATES.items()}
        )
        
        return jsonify({
            'success': True,
            'message': 'Templates saved successfully',
            'updated': updated_count
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error saving templates: {str(e)}'
        })

# Removed /save_comments endpoint

@app.route('/load_dmp_structure', methods=['GET'])
def load_dmp_structure():
    """Load DMP structure"""
    try:
        structure_path = os.path.join('config', 'dmp_structure.json')

        if not os.path.exists(structure_path):
            return jsonify({
                'success': True,
                'structure': []
            })

        with open(structure_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if data is None:
                data = {}

        return jsonify({
            'success': True,
            'structure': data.get('structure', [])
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error loading DMP structure: {str(e)}'
        })

@app.route('/save_dmp_structure', methods=['POST'])
def save_dmp_structure():
    try:
        data = request.json

        # Save DMP structure to a file
        structure_path = os.path.join('config', 'dmp_structure.json')
        os.makedirs(os.path.dirname(structure_path), exist_ok=True)

        with open(structure_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        return jsonify({
            'success': True,
            'message': 'DMP structure saved successfully'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error saving DMP structure: {str(e)}'
        })

# Old template_editor route removed - now handled by settings_redirect()

@app.route('/save_feedback', methods=['POST'])
def save_feedback():
    try:
        data = request.json or {}
        
        filename = data.get('filename', '')
        feedback = data.get('feedback', '')
        cache_id = data.get('cache_id', '')
        feedback_data = data.get('feedbackData', {})
        
        if not filename or not feedback:
            return jsonify({
                'success': False,
                'message': 'Missing filename or feedback text'
            })
        
        feedback_filename = f"feedback_{os.path.splitext(filename)[0]}.txt"
        feedback_path = os.path.join(app.config['REVIEWS_FOLDER'], feedback_filename)
        
        with open(feedback_path, 'w', encoding='utf-8') as f:
            f.write(feedback)

        if cache_id:
            _ensure_active_session(
                cache_id,
                feedback_data=feedback_data,
                compiled_feedback=feedback
            )
        
        return jsonify({
            'success': True,
            'filename': feedback_filename,
            'path': feedback_path,
            'message': 'Feedback saved successfully'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error saving feedback: {str(e)}'
        })

@app.route('/export_json', methods=['POST'])
def export_json():
    """Export review with metadata as structured JSON"""
    try:
        data = request.json or {}

        cache_id = data.get('cache_id', '')
        feedback_data = data.get('feedback', {})

        if not cache_id:
            return jsonify({
                'success': False,
                'message': 'Missing cache_id'
            })

        try:
            session_bundle = _ensure_active_session(cache_id, feedback_data=feedback_data)
        except FileNotFoundError:
            return jsonify({
                'success': False,
                'message': 'Cache file not found'
            })

        cache_data = session_bundle['cache_data']

        # Extract metadata
        metadata = cache_data.get('_metadata', {})

        # Build structured export
        export_data = {
            'metadata': {
                'researcher_surname': metadata.get('researcher_surname'),
                'researcher_firstname': metadata.get('researcher_firstname'),
                'competition_name': metadata.get('competition_name'),
                'competition_edition': metadata.get('competition_edition'),
                'creation_date': metadata.get('creation_date'),
                'review_date': datetime.now().strftime('%d-%m-%y'),
                'filename_original': metadata.get('filename_original'),
                'cache_id': cache_id,
                'dmp_cache_file': f"cache_{cache_id}.json"
            },
            'dmp_content': {},
            'review_feedback': {}
        }

        # Add DMP content for each section
        for section_id in ['1.1', '1.2', '2.1', '2.2', '3.1', '3.2',
                          '4.1', '4.2', '5.1', '5.2', '5.3', '5.4', '6.1', '6.2']:
            if section_id in cache_data:
                section_info = cache_data[section_id]
                export_data['dmp_content'][section_id] = {
                    'section': section_info.get('section', ''),
                    'question': section_info.get('question', ''),
                    'content': '\n'.join(section_info.get('paragraphs', []))
                }

                # Add review feedback if provided
                if section_id in feedback_data:
                    export_data['review_feedback'][section_id] = feedback_data[section_id]

        # Generate filename
        if metadata.get('researcher_surname'):
            json_filename = f"Review_{metadata['researcher_surname']}"
            if metadata.get('researcher_firstname'):
                json_filename += f"_{metadata['researcher_firstname'][0]}"
            if metadata.get('competition_name'):
                json_filename += f"_{metadata['competition_name']}"
            if metadata.get('competition_edition'):
                json_filename += f"_{metadata['competition_edition']}"
        else:
            json_filename = f"Review_{cache_id[:8]}"

        json_filename += f"_{datetime.now().strftime('%d%m%y')}.json"
        json_path = os.path.join(app.config['REVIEWS_FOLDER'], json_filename)

        # Save JSON file
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)

        _write_json_file(session_bundle['paths']['review_export_path'], export_data)

        return jsonify({
            'success': True,
            'filename': json_filename,
            'path': json_path,
            'data': export_data,
            'message': 'JSON exported successfully'
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'Error exporting JSON: {str(e)}'
        })

# ===========================================
# ARCHIVE SYSTEM - Session Management
# ===========================================

@app.route('/api/archive-session', methods=['POST'])
def archive_session():
    """
    Archive a review session with DMP plan and feedback.
    Creates a folder with: dmp_plan.json, feedback.json, metadata.json
    """
    try:
        data = request.json or {}
        cache_id = data.get('cache_id', '')
        feedback_data = data.get('feedbackData', {})       # Raw section data
        compiled_feedback = data.get('feedback', '')       # Compiled text report
        meta_override = data.get('meta', {})               # User-supplied metadata from form

        if not cache_id:
            return jsonify({
                'success': False,
                'message': 'Missing cache_id'
            })

        try:
            session_bundle = _ensure_active_session(
                cache_id,
                feedback_data=feedback_data,
                compiled_feedback=compiled_feedback
            )
        except FileNotFoundError:
            return jsonify({
                'success': False,
                'message': 'Cache file not found'
            })

        active_paths = session_bundle['paths']
        metadata = session_bundle['metadata']

        # Create archive folder with timestamp and cache_id
        now = datetime.now()
        timestamp = now.strftime('%Y%m%d_%H%M%S')
        archive_id = f"{timestamp}_{cache_id[:8]}"
        archive_folder = os.path.join(app.config['SESSION_ARCHIVE_FOLDER'], archive_id)
        os.makedirs(archive_folder, exist_ok=True)

        # Prepare metadata — user-supplied values override extracted (always empty) ones
        metadata_json = dict(metadata)
        metadata_json.update({
            'archive_id': archive_id,
            'timestamp': timestamp,
            'archived_date': now.isoformat(),
            'creation_date': now.strftime('%d.%m.%Y %H:%M'),
            'status': 'archived',
            'archive_folder': archive_folder
        })
        if meta_override.get('researcher_surname'):
            metadata_json['researcher_surname'] = meta_override['researcher_surname']
        if meta_override.get('researcher_firstname'):
            metadata_json['researcher_firstname'] = meta_override['researcher_firstname']
        if meta_override.get('competition_name'):
            metadata_json['competition_name'] = meta_override['competition_name']
        if meta_override.get('competition_edition'):
            metadata_json['competition_edition'] = meta_override['competition_edition']
        if meta_override.get('session_name'):
            metadata_json['session_name'] = meta_override['session_name']

        shutil.copy2(active_paths['dmp_path'], os.path.join(archive_folder, 'dmp_plan.json'))
        shutil.copy2(active_paths['feedback_path'], os.path.join(archive_folder, 'feedback.json'))

        if os.path.exists(active_paths['review_export_path']):
            shutil.copy2(active_paths['review_export_path'], os.path.join(archive_folder, 'review_export.json'))

        source_upload_path, source_upload_name = _find_session_source_upload(active_paths['session_dir'])
        if source_upload_path and source_upload_name:
            shutil.copy2(source_upload_path, os.path.join(archive_folder, source_upload_name))

        _write_json_file(os.path.join(archive_folder, 'metadata.json'), metadata_json)

        preserved_metadata = dict(session_bundle['metadata'])
        preserved_metadata['last_archived_at'] = metadata_json['archived_date']
        preserved_metadata['last_archive_id'] = archive_id
        preserved_metadata['preserved_after_archive'] = True
        _write_json_file(active_paths['metadata_path'], preserved_metadata)

        return jsonify({
            'success': True,
            'archive_id': archive_id,
            'message': 'Session archived successfully'
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'Error archiving session: {str(e)}'
        })

@app.route('/api/get-archived-sessions', methods=['GET'])
def get_archived_sessions():
    """Get list of all archived sessions with metadata"""
    try:
        archives = []
        seen_archive_ids = set()

        for archives_folder in _iter_archive_roots():
            if not os.path.exists(archives_folder):
                continue

            for archive_id in os.listdir(archives_folder):
                archive_path = os.path.join(archives_folder, archive_id)

                if not os.path.isdir(archive_path) or archive_id in seen_archive_ids:
                    continue

                metadata_path = os.path.join(archive_path, 'metadata.json')

                if os.path.exists(metadata_path):
                    metadata = _load_json_file(metadata_path, {})
                    metadata.setdefault('archive_id', archive_id)
                    metadata.setdefault('archive_folder', archive_path)
                    archives.append(metadata)
                    seen_archive_ids.add(archive_id)

        # Sort by archived_date descending (newest first)
        archives.sort(key=lambda x: x.get('archived_date', ''), reverse=True)

        return jsonify({
            'success': True,
            'archives': archives
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error loading archived sessions: {str(e)}'
        })

@app.route('/api/get-active-sessions', methods=['POST'])
def get_active_sessions():
    """
    Get list of cache files that have NOT been archived yet.
    Receives session IDs from localStorage and checks which ones exist in cache.
    """
    try:
        data = request.json or {}
        session_ids = data.get('session_ids', [])

        active_sessions = []
        active_folder = app.config['ACTIVE_SESSIONS_FOLDER']

        if session_ids:
            candidate_ids = session_ids
        else:
            candidate_ids = [
                entry for entry in os.listdir(active_folder)
                if os.path.isdir(os.path.join(active_folder, entry))
            ] if os.path.exists(active_folder) else []

        for session_id in candidate_ids:
            try:
                session_bundle = _ensure_active_session(session_id)
                metadata = session_bundle['metadata']
            except FileNotFoundError:
                continue

            active_sessions.append({
                'session_id': session_id,
                'filename': metadata.get('filename_original', 'Unknown'),
                'session_name': metadata.get('session_name', ''),
                'researcher_surname': metadata.get('researcher_surname', ''),
                'researcher_firstname': metadata.get('researcher_firstname', ''),
                'creation_date': metadata.get('creation_date', ''),
                'last_updated': metadata.get('last_updated', '')
            })

        active_sessions.sort(key=lambda item: item.get('last_updated', ''), reverse=True)

        return jsonify({
            'success': True,
            'active_sessions': active_sessions
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error loading active sessions: {str(e)}'
        })

@app.route('/api/delete-archived-session/<archive_id>', methods=['DELETE'])
def delete_archived_session(archive_id):
    """Delete an archived session folder"""
    try:
        archive_path = _find_archive_path(archive_id)

        if not archive_path:
            return jsonify({
                'success': False,
                'message': 'Archive not found'
            })

        # Delete all files in folder
        shutil.rmtree(archive_path)

        return jsonify({
            'success': True,
            'message': 'Archive deleted successfully'
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error deleting archive: {str(e)}'
        })

@app.route('/api/restore-archived-session/<archive_id>', methods=['GET'])
def restore_archived_session(archive_id):
    """Load archived session data for viewing/restoring"""
    try:
        archive_path = _find_archive_path(archive_id)

        if not archive_path:
            return jsonify({
                'success': False,
                'message': 'Archive not found'
            })

        for required_file in ['metadata.json', 'dmp_plan.json', 'feedback.json']:
            if not os.path.exists(os.path.join(archive_path, required_file)):
                return jsonify({'success': False, 'message': f'Plik archiwum niekompletny: brak {required_file}'})

        with open(os.path.join(archive_path, 'metadata.json'), 'r', encoding='utf-8') as f:
            metadata = json.load(f)

        with open(os.path.join(archive_path, 'dmp_plan.json'), 'r', encoding='utf-8') as f:
            dmp_plan = json.load(f)

        with open(os.path.join(archive_path, 'feedback.json'), 'r', encoding='utf-8') as f:
            feedback = json.load(f)

        return jsonify({
            'success': True,
            'metadata': metadata,
            'dmp_plan': dmp_plan,
            'feedback': feedback
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error restoring archive: {str(e)}'
        })

@app.route('/api/rename-session', methods=['POST'])
def rename_session():
    """Rename a session (active or archived) by updating session_name in metadata.json"""
    try:
        data = request.json or {}
        session_id = data.get('session_id', '')   # cache_id for active, archive_id for archived
        session_type = data.get('session_type', 'archive')   # 'active' or 'archive'
        session_name = data.get('session_name', '').strip()

        if not session_id:
            return jsonify({'success': False, 'message': 'Missing session_id'})

        if session_type == 'active':
            paths = _get_active_session_paths(session_id)
            metadata_path = paths['metadata_path']
        else:
            archive_path = _find_archive_path(session_id)
            if not archive_path:
                return jsonify({'success': False, 'message': 'Archive not found'})
            metadata_path = os.path.join(archive_path, 'metadata.json')

        if not os.path.exists(metadata_path):
            return jsonify({'success': False, 'message': 'Metadata not found'})

        metadata = _load_json_file(metadata_path, {})
        metadata['session_name'] = session_name
        metadata['last_updated'] = datetime.now().isoformat()
        _write_json_file(metadata_path, metadata)

        return jsonify({'success': True, 'message': 'Session renamed successfully'})

    except Exception as e:
        return jsonify({'success': False, 'message': f'Error renaming session: {str(e)}'})


@app.route('/save_category', methods=['POST'])
def save_category():
    """Save category with its comments"""
    try:
        data = request.json or {}
        file = data.get('file')
        category_data = data.get('data', {})
        lang = data.get('lang', 'pl')

        if not file:
            return jsonify({
                'success': False,
                'message': 'File name is required'
            })

        # Resolve the actual file path using the same logic as load
        config_dir = 'config'
        filename, file_path = resolve_category_file(config_dir, file, lang)

        # If file doesn't exist, create new one with language suffix
        if not file_path:
            filename = f'{file}_{lang}.json'
            file_path = os.path.join(config_dir, filename)

        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(category_data, f, indent=2, ensure_ascii=False)

        return jsonify({
            'success': True,
            'message': f'Category "{filename}" saved successfully'
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error saving category: {str(e)}'
        })

@app.route('/load_categories', methods=['GET'])
def load_categories():
    """Load all categories and their comments from individual JSON files.

    Supports language-specific files via ?lang=pl|en parameter.
    For each category base name, prefers {base}_{lang}.json over {base}.json.
    Language-variant files (ending in _pl or _en) are not shown as separate categories.
    """
    try:
        config_dir = 'config'
        categories = {}
        lang = request.args.get('lang', 'pl')

        skip_files = CATEGORY_SYSTEM_FILES

        if os.path.exists(config_dir):
            base_categories = collect_category_base_names(config_dir, skip_files)

            # Second pass: for each base category, load language-specific file or fallback
            for file_base in sorted(base_categories):
                _, file_path = resolve_category_file(config_dir, file_base, lang)
                if not file_path:
                    continue

                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)

                        # Check if data is a valid category (dict with section keys)
                        if data and isinstance(data, dict):
                            has_section_keys = any(
                                key.replace('.', '').replace('GENERAL', '').isdigit() or key == 'GENERAL'
                                for key in data.keys()
                                if not key.startswith('_')
                            )

                            if has_section_keys:
                                display_name = format_category_name(file_base)
                                categories[display_name] = data

                except Exception as e:
                    print(f"Error loading category file {file_path}: {str(e)}")
                    continue

        return jsonify({
            'success': True,
            'categories': categories
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error loading categories: {str(e)}'
        })

@app.route('/load_category_comments', methods=['GET'])
def load_category_comments():
    """Load category-specific comments for feedback sections"""
    try:
        category_comments_path = os.path.join('config', 'category_comments.json')
        
        if not os.path.exists(category_comments_path):
            return jsonify({
                'success': True,
                'category_comments': {}
            })
        
        with open(category_comments_path, 'r', encoding='utf-8') as f:
            category_comments = json.load(f)
            if category_comments is None:
                category_comments = {}
        
        return jsonify({
            'success': True,
            'category_comments': category_comments
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error loading category comments: {str(e)}'
        })

@app.route('/save_category_comments', methods=['POST'])
def save_category_comments():
    """Save category-specific comments for feedback sections"""
    try:
        data = request.json or {}
        category_comments = data.get('category_comments', {})
        
        category_comments_path = os.path.join('config', 'category_comments.json')
        os.makedirs(os.path.dirname(category_comments_path), exist_ok=True)
        
        with open(category_comments_path, 'w', encoding='utf-8') as f:
            json.dump(category_comments, f, indent=2, ensure_ascii=False)
        
        return jsonify({
            'success': True,
            'message': 'Category comments saved successfully'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error saving category comments: {str(e)}'
        })

@app.route('/api/discover-categories', methods=['GET'])
def discover_categories():
    """
    Discover all category JSON files in config/ directory dynamically.

    Returns JSON list of categories (excluding dmp_structure and quick_comments).
    This enables the template editor to work with any number of categories
    without hardcoding.
    """
    try:
        config_dir = 'config'
        categories = []

        if not os.path.exists(config_dir):
            return jsonify({
                'success': False,
                'message': 'Config directory not found'
            }), 404

        for category_name in collect_category_base_names(config_dir, CATEGORY_SYSTEM_FILES):
            categories.append({
                'id': category_name,
                'filename': f'{category_name}.json',
                'display_name': format_category_name(category_name)
            })

        # Sort alphabetically
        categories.sort(key=lambda x: x['display_name'])

        return jsonify({
            'success': True,
            'categories': categories,
            'count': len(categories)
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/load-category/<category_id>', methods=['GET'])
def load_single_category(category_id):
    """
    Load a single category JSON file by ID.

    Used by Template Editor to load category content for editing.
    Returns the full category data structure with all sections and comments.
    """
    try:
        # Validate category_id to prevent directory traversal
        if '..' in category_id or '/' in category_id or '\\' in category_id:
            return jsonify({
                'success': False,
                'message': 'Invalid category ID'
            }), 400

        # Build file path, with language-specific fallback
        config_dir = 'config'
        lang = request.args.get('lang', 'pl')
        base_filename = f"{category_id}.json"
        filename, file_path = resolve_category_file(config_dir, category_id, lang)

        # Check if file exists
        if not file_path or not os.path.exists(file_path):
            return jsonify({
                'success': False,
                'message': f'Category file not found: {base_filename}'
            }), 404

        # Skip system files
        if base_filename in CATEGORY_SYSTEM_FILES:
            return jsonify({
                'success': False,
                'message': 'Cannot load system configuration files'
            }), 403

        # Load and return the category data
        with open(file_path, 'r', encoding='utf-8') as f:
            category_data = json.load(f)

        return jsonify({
            'success': True,
            'category_id': category_id,
            'display_name': format_category_name(category_id),
            'data': category_data
        })

    except json.JSONDecodeError as e:
        return jsonify({
            'success': False,
            'message': f'Invalid JSON in category file: {str(e)}'
        }), 500
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error loading category: {str(e)}'
        }), 500


def format_category_name(category_id):
    """
    Format category ID for display.

    Examples:
        'for_newbies' -> 'For Newbies'
        'missing_info' -> 'Missing Info'
        'ready_to_use' -> 'Ready to Use'
        'my_custom_category' -> 'My Custom Category'
    """
    # Special cases for existing categories (handle both old and new naming)
    special_names = {
        'newcomer': 'Newcomer',
        'mising': 'Missing Info',
        'ready': 'Ready to Use',
        'for_newbies': 'For Newbies',
        'missing_info': 'Missing Info',
        'ready_to_use': 'Ready to Use'
    }

    if category_id in special_names:
        return special_names[category_id]

    # Default: title case with underscores replaced
    return category_id.replace('_', ' ').title()


@app.route('/api/create-category', methods=['POST'])
def create_category():
    """Create a new category file"""
    try:
        data = request.json or {}
        category_name = data.get('name', '').strip()
        category_content = data.get('content', {})

        if not category_name:
            return jsonify({
                'success': False,
                'message': 'Category name is required'
            }), 400

        # Validate category name (alphanumeric with underscores only)
        if not re.match(r'^[a-z0-9_]+$', category_name):
            return jsonify({
                'success': False,
                'message': 'Category name must be lowercase alphanumeric with underscores only'
            }), 400

        # Check if category already exists
        category_path = os.path.join('config', f'{category_name}.json')
        if os.path.exists(category_path):
            return jsonify({
                'success': False,
                'message': 'Category already exists'
            }), 409

        # Create category file
        with open(category_path, 'w', encoding='utf-8') as f:
            json.dump(category_content, f, indent=2, ensure_ascii=False)

        return jsonify({
            'success': True,
            'message': f'Category "{category_name}" created successfully',
            'category': {
                'id': category_name,
                'filename': f'{category_name}.json',
                'display_name': format_category_name(category_name)
            }
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500


@app.route('/api/delete-category/<category_id>', methods=['DELETE'])
def delete_category(category_id):
    """Delete a category file"""
    try:
        # Prevent deletion of system files
        if category_id in ['dmp_structure', 'quick_comments', 'category_comments']:
            return jsonify({
                'success': False,
                'message': 'Cannot delete system files'
            }), 403

        category_path = os.path.join('config', f'{category_id}.json')

        if not os.path.exists(category_path):
            return jsonify({
                'success': False,
                'message': 'Category not found'
            }), 404

        # Create backup before deleting
        backup_name = f'{category_id}_backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json'
        backup_path = os.path.join('config', backup_name)

        # Copy to backup
        import shutil
        shutil.copy2(category_path, backup_path)

        # Delete the category file
        os.remove(category_path)

        return jsonify({
            'success': True,
            'message': f'Category "{category_id}" deleted successfully (backup created)',
            'backup_file': backup_name
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/save_quick_comments', methods=['POST'])
def save_quick_comments():
    """Save quick comments"""
    try:
        data = request.json or {}
        quick_comments = data.get('quick_comments', [])
        
        quick_comments_path = os.path.join('config', 'quick_comments.json')
        
        with open(quick_comments_path, 'w', encoding='utf-8') as f:
            json.dump({
                "quick_comments": quick_comments
            }, f, indent=2, ensure_ascii=False)
        
        return jsonify({
            'success': True,
            'message': 'Quick comments saved successfully'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error saving quick comments: {str(e)}'
        })

@app.route('/load_quick_comments', methods=['GET'])
def load_quick_comments():
    """Load quick comments"""
    try:
        quick_comments_path = os.path.join('config', 'quick_comments.json')
        
        if not os.path.exists(quick_comments_path):
            return jsonify({
                'success': True,
                'quick_comments': []
            })
        
        with open(quick_comments_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if data is None:
                data = {}
        
        return jsonify({
            'success': True,
            'quick_comments': data.get('quick_comments', [])
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error loading quick comments: {str(e)}'
        })

@app.route('/list_categories', methods=['GET'])
def list_categories():
    """List all available category files"""
    try:
        config_dir = 'config'
        categories = []

        if os.path.exists(config_dir):
            for filename in os.listdir(config_dir):
                # Skip backup files, dmp_structure, and quick_comments
                if (filename.endswith('.json') and
                    filename not in ['dmp_structure.json', 'quick_comments.json'] and
                    'backup' not in filename.lower()):
                    file_base = filename[:-5]
                    category_name = file_base.replace('_', ' ').title()
                    categories.append({
                        'file': file_base,
                        'name': category_name
                    })

        return jsonify({
            'success': True,
            'categories': categories
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error listing categories: {str(e)}'
        })

@app.route('/create_category', methods=['POST'])
def create_category_legacy():
    """Create a new category file (legacy endpoint)"""
    try:
        data = request.json or {}
        name = data.get('name', '').strip()
        
        if not name:
            return jsonify({
                'success': False,
                'message': 'Category name is required'
            })
        
        file_name = name.lower().replace(' ', '_')
        category_path = os.path.join('config', f'{file_name}.json')
        
        if os.path.exists(category_path):
            return jsonify({
                'success': False,
                'message': 'Category with this name already exists'
            })
        
        category_data = {}
        
        with open(category_path, 'w', encoding='utf-8') as f:
            json.dump(category_data, f, indent=2, ensure_ascii=False)
        
        return jsonify({
            'success': True,
            'message': f'Category "{name}" created successfully',
            'file': file_name
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error creating category: {str(e)}'
        })

@app.route('/delete_category', methods=['POST'])
def delete_category_legacy():
    """Delete a category file (legacy endpoint)"""
    try:
        data = request.json or {}
        file = data.get('file', '').strip()
        
        if not file:
            return jsonify({
                'success': False,
                'message': 'File name is required'
            })
        
        category_path = os.path.join('config', f'{file}.json')
        
        if not os.path.exists(category_path):
            return jsonify({
                'success': False,
                'message': 'Category file does not exist'
            })
        
        os.remove(category_path)
        
        return jsonify({
            'success': True,
            'message': f'Category "{file}" deleted successfully'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error deleting category: {str(e)}'
        })

@app.route('/config/<filename>')
def serve_config(filename):
    """Serve config files"""
    try:
        config_path = os.path.join('config', filename)
        if not os.path.exists(config_path):
            return jsonify({'error': 'File not found'}), 404
        
        with open(config_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if data is None:
                data = {}
            return data
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/test_categories')
def test_categories():
    """Test endpoint to debug category loading"""
    try:
        config_dir = 'config'
        result = {
            'config_dir_exists': os.path.exists(config_dir),
            'files_in_config': [],
            'categories_found': [],
            'load_categories_result': None,
            'load_categories_errors': []
        }
        
        # List all files in config directory
        if os.path.exists(config_dir):
            all_files = os.listdir(config_dir)
            result['files_in_config'] = all_files
            
            # Process each JSON file
            for filename in all_files:
                if filename.endswith('.json'):
                    file_path = os.path.join(config_dir, filename)
                    file_info = {
                        'filename': filename,
                        'file_base': filename[:-5],
                        'excluded': filename in ['dmp_structure.json', 'quick_comments.json'],
                        'file_exists': os.path.exists(file_path),
                        'file_size': os.path.getsize(file_path) if os.path.exists(file_path) else 0,
                        'json_valid': False,
                        'json_error': None,
                        'contains_category_data': False,
                        'category_keys': []
                    }
                    
                    # Try to load and validate JSON
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                            file_info['json_valid'] = True
                            file_info['category_keys'] = list(data.keys())
                            
                            # Check if it contains category data (dict with DMP sections)
                            for key, value in data.items():
                                if not key.startswith('_') and isinstance(value, dict):
                                    file_info['contains_category_data'] = True
                                    break
                                    
                    except Exception as e:
                        file_info['json_error'] = str(e)
                    
                    result['categories_found'].append(file_info)
        
        # Test the actual list_categories logic
        try:
            categories = []
            if os.path.exists(config_dir):
                for filename in os.listdir(config_dir):
                    if filename.endswith('.json') and filename not in ['dmp_structure.json', 'quick_comments.json']:
                        file_base = filename[:-5]
                        category_name = file_base.replace('_', ' ').title()
                        categories.append({
                            'file': file_base,
                            'name': category_name
                        })
            
            result['load_categories_result'] = {
                'success': True,
                'categories': categories,
                'count': len(categories)
            }
        except Exception as e:
            result['load_categories_errors'].append(str(e))
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            'error': str(e),
            'status': 'failed'
        })

@app.route('/progress/<session_id>')
def progress_stream(session_id):
    """
    Server-Sent Events (SSE) endpoint for real-time progress updates

    Client connects to this endpoint and receives progress updates as they occur.
    Format: data: {"message": "...", "progress": 0-100, "status": "processing|complete|error"}
    """
    def generate():
        """Generator function that yields SSE-formatted progress updates"""
        # Initial connection message
        yield f"data: {json.dumps({'message': 'Connected', 'progress': 0, 'status': 'connected'})}\n\n"

        last_progress = -1
        timeout_count = 0
        max_timeout = 300  # 5 minutes max wait (300 * 1 second checks)

        while timeout_count < max_timeout:
            with progress_lock:
                if session_id in progress_state:
                    state = progress_state[session_id]
                    current_progress = state.get('progress', 0)

                    # Only send update if progress changed
                    if current_progress != last_progress:
                        message_data = {
                            'message': state.get('message', ''),
                            'progress': current_progress,
                            'status': state.get('status', 'processing')
                        }
                        yield f"data: {json.dumps(message_data)}\n\n"
                        last_progress = current_progress

                    # Check if processing is complete
                    if state.get('status') in ['complete', 'error']:
                        # Send final message
                        final_data = {
                            'message': state.get('message', ''),
                            'progress': 100 if state.get('status') == 'complete' else current_progress,
                            'status': state.get('status'),
                            'redirect': state.get('redirect')
                        }
                        yield f"data: {json.dumps(final_data)}\n\n"

                        # Cleanup state after short delay
                        time.sleep(1)
                        if session_id in progress_state:
                            del progress_state[session_id]
                        break

            # Check for updates every second
            time.sleep(1)
            timeout_count += 1

        # Timeout reached without completion
        if timeout_count >= max_timeout:
            yield f"data: {json.dumps({'message': 'Processing timeout', 'progress': 0, 'status': 'error'})}\n\n"
            with progress_lock:
                if session_id in progress_state:
                    del progress_state[session_id]

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',  # Disable nginx buffering
            'Connection': 'keep-alive'
        }
    )

@app.route('/health')
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'upload_folder': app.config['UPLOAD_FOLDER'],
        'output_folder': app.config['OUTPUT_FOLDER'],
        'cache_folder': app.config['CACHE_FOLDER'],
        'dmp_folder': app.config['DMP_FOLDER'],
        'reviews_folder': app.config['REVIEWS_FOLDER'],
        'allowed_extensions': list(app.config['ALLOWED_EXTENSIONS']),
        'max_content_length': app.config['MAX_CONTENT_LENGTH']
    })

# ============================================================
# Extraction: skip-terms API
# ============================================================

@app.route('/api/extraction/skip-terms', methods=['GET'])
def get_skip_terms():
    """Return the list of skip terms used during DMP extraction."""
    try:
        mgr = SkipTermsManager()
        return jsonify({'success': True, 'terms': mgr.load()})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/extraction/skip-terms', methods=['POST'])
def add_skip_term():
    """Add a new skip term. Body: {term: str}"""
    try:
        data = request.json or {}
        term = data.get('term', '').strip()
        if not term:
            return jsonify({'success': False, 'message': 'Term is required'}), 400
        mgr = SkipTermsManager()
        terms = mgr.add(term)
        return jsonify({'success': True, 'terms': terms})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/extraction/skip-terms', methods=['DELETE'])
def delete_skip_term():
    """Remove a skip term. Body: {term: str}"""
    try:
        data = request.json or {}
        term = data.get('term', '')
        mgr = SkipTermsManager()
        terms = mgr.remove(term)
        return jsonify({'success': True, 'terms': terms})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# ============================================================
# Extractor Debug Mode API
# ============================================================

@app.route('/api/settings/extractor', methods=['GET'])
def get_extractor():
    """Return currently active extractor identifier."""
    return jsonify({
        'success': True,
        'extractor_name': EXTRACTOR_NAME,
        'available': ['v4'],
    })

@app.route('/api/settings/extractor', methods=['POST'])
def update_extractor():
    """Switch the active extractor. Extend 'available' list when adding new extractors."""
    global EXTRACTOR_NAME
    try:
        data = request.json or {}
        name = data.get('extractor_name')
        available = ['v4']
        if name not in available:
            return jsonify({'success': False, 'message': f'Unknown extractor. Available: {available}'}), 400

        EXTRACTOR_NAME = name

        saved = {}
        if os.path.exists(_GENERAL_SETTINGS_PATH):
            try:
                with open(_GENERAL_SETTINGS_PATH, 'r', encoding='utf-8') as f:
                    saved = json.load(f)
            except Exception:
                saved = {}
        saved['extractor_name'] = EXTRACTOR_NAME
        with open(_GENERAL_SETTINGS_PATH, 'w', encoding='utf-8') as f:
            json.dump(saved, f, indent=2, ensure_ascii=False)

        return jsonify({'success': True, 'message': f'Extractor set to {name}', 'extractor_name': name})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

# Keep old endpoint as alias for backwards compatibility with any existing UI references
@app.route('/api/settings/extractor-debug', methods=['GET', 'POST'])
def extractor_debug_alias():
    """Deprecated alias — use /api/settings/extractor instead."""
    if request.method == 'GET':
        return get_extractor()
    return jsonify({'success': False, 'message': 'Use /api/settings/extractor to switch extractors'}), 410


# ============================================================
# Unified Settings Page
# ============================================================

# Settings modules registry - developers add new modules here
SETTINGS_MODULES = [
    {'id': 'general', 'name': 'General', 'icon': 'sliders-h', 'badge': None},
    {'id': 'comments', 'name': 'Comments', 'icon': 'comments', 'badge': None},
    {'id': 'ai', 'name': 'AI Assistant', 'icon': 'robot', 'badge': None},
    {'id': 'extraction', 'name': 'Ekstrakcja', 'icon': 'filter', 'badge': None},
]

@app.route('/settings')
def settings_page():
    """Unified settings page with modular architecture"""
    return render_template('settings.html', modules=SETTINGS_MODULES,
                           app_version='0.9.1',
                           extraction_success_rate='94.1%')

@app.route('/template_editor')
@app.route('/ai-settings')
def settings_redirect():
    """Redirect old settings pages to unified settings"""
    if request.path == '/ai-settings':
        return redirect('/settings#ai')
    return redirect('/settings#comments')

@app.route('/api/settings/general', methods=['GET'])
def get_general_settings():
    """Get general settings"""
    max_upload_mb = app.config.get('MAX_CONTENT_LENGTH', 16 * 1024 * 1024) // (1024 * 1024)
    return jsonify({
        'success': True,
        'settings': {
            'max_upload_mb': max_upload_mb
        }
    })

@app.route('/api/settings/general', methods=['POST'])
def update_general_settings():
    """Update general settings and persist to config/settings.json"""
    try:
        data = request.json or {}
        if 'max_upload_mb' in data:
            mb = int(data['max_upload_mb'])
            if mb < 1 or mb > 128:
                return jsonify({'success': False, 'message': 'File size must be 1-128 MB'}), 400
            app.config['MAX_CONTENT_LENGTH'] = mb * 1024 * 1024
            # Persist to disk
            saved = {}
            if os.path.exists(_GENERAL_SETTINGS_PATH):
                try:
                    with open(_GENERAL_SETTINGS_PATH, 'r', encoding='utf-8') as f:
                        saved = json.load(f)
                except Exception:
                    saved = {}
            saved['max_upload_mb'] = mb
            with open(_GENERAL_SETTINGS_PATH, 'w', encoding='utf-8') as f:
                json.dump(saved, f, indent=2)
        return jsonify({'success': True, 'message': 'Settings updated'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/settings/cache-count', methods=['GET'])
def get_cache_count():
    """Get number of cached files"""
    cache_dir = app.config['CACHE_FOLDER']
    count = 0
    if os.path.exists(cache_dir):
        count = len([f for f in os.listdir(cache_dir) if f.endswith('.json')])
    return jsonify({'success': True, 'count': count})

@app.route('/api/settings/clear-cache', methods=['POST'])
def clear_cache():
    """Clear all cached extraction results"""
    try:
        cache_dir = app.config['CACHE_FOLDER']
        deleted = 0
        if os.path.exists(cache_dir):
            for f in os.listdir(cache_dir):
                if f.endswith('.json'):
                    os.remove(os.path.join(cache_dir, f))
                    deleted += 1
        return jsonify({'success': True, 'deleted': deleted})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

# ============================================================
# AI Module Routes
# ============================================================

@app.route('/api/ai/config', methods=['GET'])
def get_ai_config():
    """Get AI configuration (without API keys)"""
    config = ai_assistant.get_config(hide_keys=True)
    return jsonify({'success': True, 'config': config})

@app.route('/api/ai/config', methods=['POST'])
def update_ai_config():
    """Update AI configuration"""
    try:
        data = request.json or {}
        ai_assistant.update_settings(data)
        return jsonify({'success': True, 'message': 'Ustawienia zapisane'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/ai/test-connection', methods=['POST'])
def test_ai_connection():
    """Test connection to AI API"""
    success, message = ai_assistant.test_connection()
    return jsonify({'success': success, 'message': message})

@app.route('/api/ai/models', methods=['GET'])
def get_ai_models():
    """Get list of available models from current AI provider"""
    try:
        result = ai_assistant.list_available_models()
        return jsonify(result)
    except Exception as e:
        return jsonify({
            'success': False,
            'models': [],
            'error': str(e)
        }), 500

@app.route('/api/ai/toggle', methods=['POST'])
def toggle_ai():
    """Enable/disable AI module"""
    data = request.json or {}
    if data.get('enabled'):
        ai_assistant.enable()
    else:
        ai_assistant.disable()
    return jsonify({'success': True, 'enabled': ai_assistant.is_enabled()})

@app.route('/api/ai/suggest', methods=['POST'])
def ai_suggest_feedback():
    """Generate AI suggestions for DMP review"""
    if not ai_assistant.is_enabled():
        return jsonify({'success': False, 'message': 'Moduł AI jest wyłączony'})

    try:
        data = request.json or {}
        cache_id = data.get('cache_id')
        section_id = data.get('section_id')  # Optional - for single section

        if not cache_id:
            return jsonify({'success': False, 'message': 'Brak cache_id'})

        # Validate cache_id format (should be UUID)
        if not re.match(r'^[a-f0-9]{8}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{12}$', cache_id):
            return jsonify({'success': False, 'message': 'Nieprawidłowy format cache_id'})

        # Load DMP from cache
        cache_path = os.path.join(app.config['CACHE_FOLDER'], f"cache_{cache_id}.json")
        if not os.path.exists(cache_path):
            return jsonify({'success': False, 'message': 'Cache nie znaleziony'})

        with open(cache_path, 'r', encoding='utf-8') as f:
            dmp_content = json.load(f)

        # Load available comments from categories
        available_comments = load_all_category_comments()

        if section_id:
            # Single section suggestion
            section_content = dmp_content.get(section_id, {})
            if isinstance(section_content, dict):
                content_text = "\n".join(section_content.get('paragraphs', []))
            else:
                content_text = str(section_content)

            section_comments = get_comments_for_section(section_id, available_comments)

            suggestions = ai_assistant.generate_section_suggestion(
                section_id=section_id,
                content=content_text,
                available_comments=section_comments
            )
        else:
            # All sections suggestion
            suggestions = ai_assistant.generate_review_suggestions(
                dmp_content=dmp_content,
                available_comments=available_comments
            )

        return jsonify({'success': True, 'suggestions': suggestions})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/ai/learn', methods=['POST'])
def ai_learn_from_feedback():
    """Learn from user's saved feedback"""
    try:
        data = request.json or {}
        section_id = data.get('section_id')
        dmp_content = data.get('dmp_content', '')
        feedback_text = data.get('feedback_text', '')
        used_comments = data.get('used_comments', [])

        if not section_id:
            return jsonify({'success': False, 'message': 'section_id is required'}), 400

        ai_assistant.learn_from_saved_feedback(
            section_id=section_id,
            dmp_content=dmp_content,
            feedback_text=feedback_text,
            used_comments=used_comments
        )

        return jsonify({'success': True, 'message': 'Wzorce zapisane'})

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/ai/knowledge', methods=['GET'])
def get_knowledge_base():
    """Get knowledge base"""
    return jsonify({
        'success': True,
        'knowledge_base': ai_assistant.get_knowledge_base()
    })

@app.route('/api/ai/knowledge/<section_id>/<issue_id>', methods=['PUT'])
def update_knowledge_entry(section_id, issue_id):
    """Update knowledge base entry"""
    try:
        updates = request.json or {}
        success = ai_assistant.update_knowledge_entry(section_id, issue_id, updates)
        return jsonify({'success': success})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/ai/knowledge/<section_id>/<issue_id>', methods=['DELETE'])
def delete_knowledge_entry(section_id, issue_id):
    """Delete knowledge base entry"""
    try:
        success = ai_assistant.delete_knowledge_entry(section_id, issue_id)
        return jsonify({'success': success})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/ai/statistics', methods=['GET'])
def get_ai_statistics():
    """Get AI usage statistics"""
    return jsonify({
        'success': True,
        'statistics': ai_assistant.get_statistics()
    })

def load_all_category_comments():
    """Load all comments from category files"""
    comments = {}
    config_dir = 'config'

    if not os.path.exists(config_dir):
        return comments

    for category_name in collect_category_base_names(config_dir, CATEGORY_SYSTEM_FILES):
        _, file_path = resolve_category_file(config_dir, category_name, 'pl')
        if not file_path:
            continue

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if data and isinstance(data, dict):
                    comments[category_name] = data
        except Exception:
            continue

    return comments

def get_comments_for_section(section_id, all_comments):
    """Get comments for a specific section"""
    section_comments = []

    for category, sections in all_comments.items():
        if isinstance(sections, dict) and section_id in sections:
            section_data = sections[section_id]
            if isinstance(section_data, list):
                for i, comment in enumerate(section_data):
                    section_comments.append({
                        'id': f"{category}_{section_id}_{i:03d}",
                        'text': comment if isinstance(comment, str) else str(comment),
                        'category': category
                    })

    return section_comments

# ============================================================
# End AI Module Routes
# ============================================================

@app.errorhandler(413)
def too_large(e):
    return jsonify({
        'success': False,
        'message': 'File too large. Maximum size is 16MB.'
    }), 413

@app.errorhandler(404)
def not_found(e):
    return jsonify({
        'success': False,
        'message': 'Resource not found.'
    }), 404

@app.errorhandler(500)
def internal_error(e):
    return jsonify({
        'success': False,
        'message': 'Internal server error. Please try again.'
    }), 500

# Favicon route to fix 404
@app.route('/favicon.ico')
def favicon():
    ico_path = os.path.join(app.root_path, 'static', 'images', 'favicon.ico')
    if os.path.exists(ico_path):
        return send_from_directory(os.path.join(app.root_path, 'static', 'images'), 'favicon.ico', mimetype='image/x-icon')
    # fallback to custom PNG if .ico not present
    return send_from_directory(os.path.join(app.root_path, 'static', 'images'), 'dmp-art-favicon (Niestandardowe).png', mimetype='image/png')

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
