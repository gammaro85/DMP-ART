"""Create a simple test DMP DOCX file for testing"""

import os
from docx import Document

# Create test DMP document
doc = Document()

# Add title
doc.add_heading('PLAN ZARZĄDZANIA DANYMI', level=1)

# Section 1
doc.add_heading('1. Opis danych oraz pozyskiwanie lub ponowne wykorzystanie istniejących danych', level=2)

doc.add_heading('1.1 Jakie dane będą zebrane lub wytworzone?', level=3)
doc.add_paragraph(
    'W projekcie zostaną zebrane dane jakościowe z 60 pogłębionych wywiadów '
    'z prezesami sądów, wiceprezesami, dyrektorami sądów oraz inspektorami.'
)
doc.add_paragraph(
    'Dane będą obejmowały transkrypcje wywiadów, notatki badawcze oraz '
    'dokumenty organizacyjne dotyczące funkcjonowania sądów.'
)

doc.add_heading('1.2 Jakie dane (np. rodzaj, format, ilość) będą zebrane lub wytworzone?', level=3)
doc.add_paragraph('Dane będą zapisane w formatach:')
doc.add_paragraph('- .mp3 / .wav (nagrania audio)')
doc.add_paragraph('- .docx / .pdf (transkrypcje)')
doc.add_paragraph('- .xlsx / .csv (dane statystyczne)')
doc.add_paragraph('Szacowana objętość danych: 20-30 GB.')

# Section 2
doc.add_heading('2. Dokumentacja i jakość danych', level=2)

doc.add_heading('2.1 Jakie metadane i dokumenty (np. metodologia lub sposób zbierania i porządkowania danych) będą towarzyszyć danym?', level=3)
doc.add_paragraph('Każdy wywiad będzie miał metadane:')
doc.add_paragraph('- data i miejsce wywiadu')
doc.add_paragraph('- kategoria respondenta')
doc.add_paragraph('- kod identyfikacyjny')
doc.add_paragraph('- status anonimizacji')

doc.add_heading('2.2 Jakie środki kontroli jakości danych będą stosowane?', level=3)
doc.add_paragraph('Transkrypcje będą weryfikowane przez dwóch niezależnych badaczy.')
doc.add_paragraph('Dane będą podlegały triangulacji: wywiady, dokumenty, statystyki.')

# Save
output_path = os.path.join('tests', 'fixtures', 'test_dmp_simple.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)

print(f"Created test DMP: {output_path}")
print(f"File size: {os.path.getsize(output_path)} bytes")
