# verify_implementation.py
"""
Quick verification script to ensure all enhanced components are working
"""
import os
import sys
import importlib.util

def check_file_exists(filepath, description):
    """Check if a file exists and report status"""
    if os.path.exists(filepath):
        print(f"✅ {description}: {filepath}")
        return True
    else:
        print(f"❌ {description}: {filepath} - NOT FOUND")
        return False

def check_python_imports():
    """Check if required Python modules can be imported"""
    required_modules = [
        ('flask', 'Flask web framework'),
        ('docx', 'python-docx for DOCX processing'),
        ('PyPDF2', 'PyPDF2 for PDF processing'),
        ('json', 'JSON processing'),
        ('re', 'Regular expressions'),
        ('os', 'Operating system interface'),
        ('uuid', 'UUID generation'),
        ('zipfile', 'ZIP file handling')
    ]
    
    print("\nChecking Python module imports:")
    print("-" * 40)
    
    all_imports_ok = True
    
    for module_name, description in required_modules:
        try:
            if module_name == 'docx':
                import docx
            elif module_name == 'PyPDF2':
                import PyPDF2
            else:
                importlib.import_module(module_name)
            print(f"✅ {description}")
        except ImportError as e:
            print(f"❌ {description} - Import failed: {str(e)}")
            all_imports_ok = False
    
    return all_imports_ok

def check_enhanced_extractor():
    """Check if the enhanced extractor has the expected methods"""
    try:
        from utils.extractor import DMPExtractor
        
        extractor = DMPExtractor()
        
        # Check for enhanced methods
        expected_methods = [
            'validate_docx_file',
            'extract_table_content',
            'clean_table_delimiters',
            'process_docx',
            'process_pdf',
            'identify_key_phrases',
            'detect_section_from_text',
            'detect_subsection_from_text'
        ]
        
        print("\nChecking enhanced extractor methods:")
        print("-" * 45)
        
        all_methods_ok = True
        
        for method_name in expected_methods:
            if hasattr(extractor, method_name):
                print(f"✅ {method_name}")
            else:
                print(f"❌ {method_name} - NOT FOUND")
                all_methods_ok = False
        
        return all_methods_ok
        
    except Exception as e:
        print(f"❌ Enhanced extractor check failed: {str(e)}")
        return False

def check_test_file():
    """Check if the test file is available"""
    test_file = "DMP_sheng4 final.docx"
    
    print(f"\nChecking test file:")
    print("-" * 20)
    
    if os.path.exists(test_file):
        file_size = os.path.getsize(test_file)
        print(f"✅ Test file found: {test_file}")
        print(f"   File size: {file_size:,} bytes")
        
        # Try to open with python-docx
        try:
            from docx import Document
            doc = Document(test_file)
            print(f"   Paragraphs: {len(doc.paragraphs)}")
            print(f"   Tables: {len(doc.tables)}")
            return True
        except Exception as e:
            print(f"❌ Cannot open test file: {str(e)}")
            return False
    else:
        print(f"⚠️  Test file not found: {test_file}")
        print("   This is optional - you can still test with other DOCX files")
        return True  # Not critical for basic functionality

def main():
    """Run all verification checks"""
    
    print("DMP ART Enhanced Implementation Verification")
    print("=" * 50)
    
    # Check core files
    print("\nChecking core files:")
    print("-" * 20)
    
    core_files_ok = True
    core_files = [
        ("app.py", "Enhanced Flask application"),
        ("utils/extractor.py", "Enhanced DMP extractor"),
        ("templates/review.html", "Enhanced review template"),
        ("templates/index.html", "Home page template"),
        ("static/css/style.css", "CSS styles"),
        ("static/js/script.js", "JavaScript functionality"),
        ("requirements.txt", "Python dependencies")
    ]
    
    for filepath, description in core_files:
        if not check_file_exists(filepath, description):
            core_files_ok = False
    
    # Check configuration files
    print("\nChecking configuration files:")
    print("-" * 30)
    
    config_files_ok = True
    config_files = [
        ("config/dmp_structure.json", "DMP structure configuration"),
        ("config/key_phrases.json", "Key phrases configuration")
    ]
    
    for filepath, description in config_files:
        if not check_file_exists(filepath, description):
            config_files_ok = False
    
    # Check directories
    print("\nChecking required directories:")
    print("-" * 32)
    
    directories_ok = True
    required_dirs = ["uploads", "outputs", "utils", "templates", "static", "config"]
    
    for directory in required_dirs:
        if os.path.exists(directory):
            print(f"✅ Directory: {directory}")
        else:
            print(f"❌ Directory: {directory} - NOT FOUND")
            print(f"   Creating directory: {directory}")
            try:
                os.makedirs(directory, exist_ok=True)
                print(f"✅ Directory created: {directory}")
            except Exception as e:
                print(f"❌ Failed to create directory: {str(e)}")
                directories_ok = False
    
    # Check Python imports
    imports_ok = check_python_imports()
    
    # Check enhanced extractor
    extractor_ok = check_enhanced_extractor()
    
    # Check test file
    test_file_ok = check_test_file()
    
    # Summary
    print("\n" + "=" * 50)
    print("VERIFICATION SUMMARY")
    print("=" * 50)
    
    checks = [
        ("Core files", core_files_ok),
        ("Configuration files", config_files_ok),
        ("Required directories", directories_ok),
        ("Python imports", imports_ok),
        ("Enhanced extractor", extractor_ok),
        ("Test file", test_file_ok)
    ]
    
    passed = sum(1 for _, ok in checks if ok)
    total = len(checks)
    
    for check_name, ok in checks:
        status = "✅ PASSED" if ok else "❌ FAILED"
        print(f"{check_name}: {status}")
    
    print(f"\nOverall: {passed}/{total} checks passed")
    
    if passed == total:
        print("\n🎉 All checks passed! Your enhanced DMP ART implementation is ready.")
        print("\nNext steps:")
        print("1. Run: python app.py")
        print("2. Open: http://localhost:5000")
        print("3. Test with a DOCX file")
    elif passed >= total - 1:
        print("\n⚠️  Almost ready! Only minor issues detected.")
        print("The application should work, but please fix the failed checks.")
    else:
        print("\n❌ Multiple issues detected. Please resolve them before proceeding.")
        print("\nCommon solutions:")
        print("- Install missing dependencies: pip install -r requirements.txt")
        print("- Ensure all enhanced files are in the correct locations")
        print("- Create missing directories and configuration files")
    
    return passed == total

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)